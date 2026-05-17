# """
# Intelligent Resume Screening and Automated Interview Notification System
# MCA Final Year Project
# """

# from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
# import sqlite3
# import os
# import json
# import re
# import smtplib
# import io
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from datetime import datetime
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# app.secret_key = "resume_screener_mca_2024"

# # Config
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# # ─────────────────────────────────────────────
# # DATABASE SETUP
# # ─────────────────────────────────────────────

# def get_db():
#     conn = sqlite3.connect(DB_PATH)
#     conn.row_factory = sqlite3.Row
#     return conn

# def init_db():
#     conn = get_db()
#     c = conn.cursor()
    
#     c.executescript("""
#     CREATE TABLE IF NOT EXISTS job_roles (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         title TEXT NOT NULL,
#         description TEXT,
#         core_skills TEXT NOT NULL,
#         tools TEXT NOT NULL,
#         project_keywords TEXT NOT NULL,
#         internship_keywords TEXT NOT NULL,
#         experience_keywords TEXT NOT NULL,
#         core_weight REAL DEFAULT 0.40,
#         tools_weight REAL DEFAULT 0.25,
#         projects_weight REAL DEFAULT 0.15,
#         internship_weight REAL DEFAULT 0.10,
#         experience_weight REAL DEFAULT 0.10,
#         min_threshold INTEGER DEFAULT 50,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     CREATE TABLE IF NOT EXISTS candidates (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         job_role_id INTEGER NOT NULL,
#         candidate_type TEXT DEFAULT 'experience',
#         batch_id TEXT,
#         name TEXT,
#         email TEXT,
#         phone TEXT,
#         raw_text TEXT,
#         skills_found TEXT,
#         tools_found TEXT,
#         projects_found TEXT,
#         internship_found TEXT,
#         experience_found TEXT,
#         internship_years TEXT,
#         experience_years TEXT,
#         core_score REAL DEFAULT 0,
#         tools_score REAL DEFAULT 0,
#         projects_score REAL DEFAULT 0,
#         internship_score REAL DEFAULT 0,
#         experience_score REAL DEFAULT 0,
#         total_score REAL DEFAULT 0,
#         status TEXT DEFAULT 'pending',
#         rejection_reason TEXT,
#         email_sent INTEGER DEFAULT 0,
#         email_sent_at TIMESTAMP,
#         filename TEXT,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS upload_batches (
#         id TEXT PRIMARY KEY,
#         job_role_id INTEGER,
#         candidate_type TEXT DEFAULT 'experience',
#         label TEXT,
#         total_resumes INTEGER DEFAULT 0,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS email_settings (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         smtp_host TEXT DEFAULT 'smtp.gmail.com',
#         smtp_port INTEGER DEFAULT 587,
#         sender_email TEXT,
#         sender_password TEXT,
#         email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
#         email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
#     );

#     INSERT OR IGNORE INTO email_settings (id) VALUES (1);
#     """)

#     # Add columns if upgrading from old DB
#     for col_sql in [
#         "ALTER TABLE candidates ADD COLUMN batch_id TEXT",
#         "ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'",
#         "ALTER TABLE candidates ADD COLUMN internship_years TEXT",
#         "ALTER TABLE candidates ADD COLUMN experience_years TEXT",
#     ]:
#         try:
#             c.execute(col_sql)
#             conn.commit()
#         except:
#             pass

#     try:
#         c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
#             id TEXT PRIMARY KEY,
#             job_role_id INTEGER,
#             candidate_type TEXT DEFAULT 'experience',
#             label TEXT,
#             total_resumes INTEGER DEFAULT 0,
#             uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#             FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#         )""")
#         conn.commit()
#     except:
#         pass

#     conn.commit()
#     conn.close()

# # ─────────────────────────────────────────────
# # KEYWORD DATABASE FOR ALL IT ROLES
# # ─────────────────────────────────────────────

# DEFAULT_ROLES = [
#     {
#         "title": "Java Full Stack Developer",
#         "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
#         "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
#         "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
#         "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
#         "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
#         "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
#     },
#     {
#         "title": "Python Full Stack Developer",
#         "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
#         "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
#         "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
#         "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
#         "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
#         "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
#     },
#     {
#         "title": "MERN Stack Developer",
#         "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
#         "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
#         "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
#         "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
#         "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
#         "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
#     },
#     {
#         "title": "Data Analyst",
#         "description": "Analyzes data to derive business insights using statistical and visualization tools.",
#         "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
#         "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
#         "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
#         "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
#         "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
#     },
#     {
#         "title": "Data Scientist",
#         "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
#         "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
#         "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
#         "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
#         "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
#     },
#     {
#         "title": "AI/ML Engineer",
#         "description": "Designs, develops, and deploys machine learning models and AI systems.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
#         "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
#         "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
#         "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
#         "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
#     },
#     {
#         "title": "DevOps Engineer",
#         "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
#         "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
#         "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
#         "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
#         "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
#         "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
#     },
#     {
#         "title": "Cloud Engineer",
#         "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
#         "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
#         "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
#         "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
#         "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
#         "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
#     },
#     {
#         "title": "Software Tester",
#         "description": "Ensures software quality through manual and automated testing methodologies.",
#         "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
#         "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
#         "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
#         "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
#         "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
#     },
#     {
#         "title": "Database Developer",
#         "description": "Designs, develops, and optimizes databases for scalability and performance.",
#         "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
#         "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
#         "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
#         "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
#         "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
#     }
# ]

# # ─────────────────────────────────────────────
# # RESUME PARSING
# # ─────────────────────────────────────────────

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(filepath):
#     text = ""
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#     except Exception as e:
#         print(f"PDF error: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ""
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX error: {e}")
#     return text

# def extract_text(filepath):
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_name(text):
#     lines = [l.strip() for l in text.split("\n") if l.strip()]
#     for line in lines[:5]:
#         if re.search(r"[@:/.()]|\d{5,}", line):
#             continue
#         if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
#             continue
#         words = line.split()
#         if 1 < len(words) <= 5:
#             return line.strip()
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
#     return match.group() if match else ""

# def extract_phone(text):
#     match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
#     return match.group().strip() if match else ""

# def extract_internship_years(text):
#     patterns = [
#         r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
#         r'internship[^.]*?(\d+)\s*month',
#         r'(\d+)\s*month[^.]*?internship',
#         r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             val = match.group(1)
#             if 'month' in pattern:
#                 months = float(val)
#                 years = round(months / 12, 1)
#                 return f"{val} months ({years} years)"
#             return f"{val} years"
#     intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
#     if intern_section:
#         try:
#             start = int(intern_section.group(1))
#             end_str = intern_section.group(2)
#             end = datetime.now().year if end_str == 'present' else int(end_str)
#             years = end - start
#             return f"{years} year(s)"
#         except:
#             pass
#     return ""

# def extract_experience_years(text):
#     patterns = [
#         r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
#         r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             return f"{match.group(1)} years"
#     return ""

# def find_keywords_in_text(text, keywords):
#     text_lower = text.lower()
#     found = []
#     for kw in keywords:
#         pattern = r'\b' + re.escape(kw.lower()) + r'\b'
#         if re.search(pattern, text_lower):
#             found.append(kw)
#     return found

# def has_internship_in_resume(text):
#     patterns = [
#         r'\binternship\b', r'\bintern\b', r'\btrainee\b',
#         r'\bindustry training\b', r'\bsummer training\b', r'\bproject trainee\b',
#     ]
#     text_lower = text.lower()
#     for p in patterns:
#         if re.search(p, text_lower):
#             return True
#     return False

# # ─────────────────────────────────────────────
# # SCORING ALGORITHM
# # ─────────────────────────────────────────────

# def calculate_score_fresher(candidate_data, role, screening_options):
#     require_internship = screening_options.get('require_internship', True)
#     core_total = len(json.loads(role["core_skills"]))
#     proj_total = len(json.loads(role["project_keywords"]))
#     intern_total = len(json.loads(role["internship_keywords"]))

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score = section_pct(candidate_data["skills_found"], core_total)
#     proj_score  = section_pct(candidate_data["projects_found"], proj_total)

#     if require_internship:
#         intern_score = section_pct(candidate_data["internship_found"], intern_total)
#         total = core_score * 0.55 + proj_score * 0.30 + intern_score * 0.15
#     else:
#         intern_score = 0
#         total = core_score * 0.65 + proj_score * 0.35

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": 0,
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": 0,
#         "total_score": round(total, 1)
#     }

# def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
#     weights = {
#         "core": role["core_weight"],
#         "tools": role["tools_weight"],
#         "projects": role["projects_weight"],
#         "internship": role["internship_weight"],
#         "experience": role["experience_weight"]
#     }

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score   = section_pct(candidate_data["skills_found"],      len(json.loads(role["core_skills"])))
#     tools_score  = section_pct(candidate_data["tools_found"],       len(json.loads(role["tools"])))
#     proj_score   = section_pct(candidate_data["projects_found"],    len(json.loads(role["project_keywords"])))
#     intern_score = section_pct(candidate_data["internship_found"],  len(json.loads(role["internship_keywords"])))
#     exp_score    = section_pct(candidate_data["experience_found"],  len(json.loads(role["experience_keywords"])))

#     total = (
#         core_score   * weights["core"] +
#         tools_score  * weights["tools"] +
#         proj_score   * weights["projects"] +
#         intern_score * weights["internship"] +
#         exp_score    * weights["experience"]
#     )

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": round(exp_score, 1),
#         "total_score": round(total, 1)
#     }

# def calculate_score(candidate_data, role):
#     return calculate_score_experience(candidate_data, role)

# def check_experience_years_range(exp_years_str, min_years, max_years):
#     if not exp_years_str: return False
#     match = re.search(r'(\d+\.?\d*)', exp_years_str)
#     if not match: return False
#     years = float(match.group(1))
#     if min_years is not None and years < min_years: return False
#     if max_years is not None and years > max_years: return False
#     return True

# def generate_rejection_reason(scores, threshold, candidate_type='experience'):
#     reasons = []
#     if scores["core_score"] < 20:
#         reasons.append("insufficient core technical skills")
#     if candidate_type == 'fresher':
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#     else:
#         if scores["tools_score"] < 20:
#             reasons.append("limited relevant tool experience")
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#         if scores["internship_score"] < 10 and scores["experience_score"] < 10:
#             reasons.append("no relevant internship or work experience")
#     if not reasons:
#         reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
#     return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

# def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
#     if screening_options is None:
#         screening_options = {}

#     text = extract_text(filepath)
#     if not text:
#         return None

#     name  = extract_name(text)
#     email = extract_email(text)
#     phone = extract_phone(text)

#     core_skills = json.loads(role["core_skills"])
#     tools       = json.loads(role["tools"])
#     proj_kw     = json.loads(role["project_keywords"])
#     intern_kw   = json.loads(role["internship_keywords"])
#     exp_kw      = json.loads(role["experience_keywords"])

#     skills_found      = find_keywords_in_text(text, core_skills)
#     tools_found       = find_keywords_in_text(text, tools)
#     projects_found    = find_keywords_in_text(text, proj_kw)
#     internship_found  = find_keywords_in_text(text, intern_kw)
#     experience_found  = find_keywords_in_text(text, exp_kw)

#     internship_years  = extract_internship_years(text)
#     experience_years  = extract_experience_years(text)

#     candidate_data = {
#         "skills_found": skills_found,
#         "tools_found": tools_found,
#         "projects_found": projects_found,
#         "internship_found": internship_found,
#         "experience_found": experience_found
#     }

#     threshold = role["min_threshold"]

#     if candidate_type == 'fresher':
#         scores = calculate_score_fresher(candidate_data, role, screening_options)
#         status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#         rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
#     else:
#         min_years = screening_options.get('min_years')
#         max_years = screening_options.get('max_years')
#         scores = calculate_score_experience(candidate_data, role, min_years, max_years)

#         if (min_years is not None or max_years is not None) and experience_years:
#             if not check_experience_years_range(experience_years, min_years, max_years):
#                 status = "rejected"
#                 if min_years and max_years:
#                     yr_range = f"{min_years}-{max_years} years"
#                 elif min_years:
#                     yr_range = f"{min_years}+ years"
#                 else:
#                     yr_range = f"up to {max_years} years"
#                 rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
#             else:
#                 status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#                 rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
#         else:
#             status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#             rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)

#     return {
#         "name": name, "email": email, "phone": phone,
#         "raw_text": text[:2000],
#         "skills_found":     json.dumps(skills_found),
#         "tools_found":      json.dumps(tools_found),
#         "projects_found":   json.dumps(projects_found),
#         "internship_found": json.dumps(internship_found),
#         "experience_found": json.dumps(experience_found),
#         "internship_years": internship_years,
#         "experience_years": experience_years,
#         "status": status,
#         "rejection_reason": rejection_reason,
#         **scores
#     }

# # ─────────────────────────────────────────────
# # EMAIL INTEGRATION
# # ─────────────────────────────────────────────

# def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
#     subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
#     body    = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)

#     msg = MIMEMultipart("alternative")
#     msg["Subject"] = subject
#     msg["From"]    = settings["sender_email"]
#     msg["To"]      = candidate_email

#     html_body = f"""
#     <html><body>
#     <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
#       <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
#         <h2 style="color:white; margin:0;">Interview Invitation</h2>
#       </div>
#       <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
#         {body.replace(chr(10), '<br>')}
#       </div>
#     </div>
#     </body></html>
#     """

#     msg.attach(MIMEText(body, "plain"))
#     msg.attach(MIMEText(html_body, "html"))

#     try:
#         server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
#         server.ehlo()
#         server.starttls()
#         server.login(settings["sender_email"], settings["sender_password"])
#         server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
#         server.quit()
#         return True, "Email sent successfully"
#     except Exception as e:
#         return False, str(e)

# # ─────────────────────────────────────────────
# # ROUTES
# # ─────────────────────────────────────────────

# @app.route("/")
# def index():
#     conn = get_db()
#     roles             = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     total_candidates  = conn.execute("SELECT COUNT(*) as c FROM candidates").fetchone()["c"]
#     shortlisted       = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted'").fetchone()["c"]
#     rejected          = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='rejected'").fetchone()["c"]
#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("index.html", roles=roles,
#                            total=total_candidates, shortlisted=shortlisted, rejected=rejected,
#                            popup_notifications=popup_notifications)

# # ── JOB ROLES ──

# @app.route("/roles")
# def roles():
#     conn  = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     conn.close()
#     return render_template("roles.html", roles=roles)

# @app.route("/roles/new", methods=["GET", "POST"])
# def new_role():
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn = get_db()
#         conn.execute("""
#             INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                 internship_keywords, experience_keywords, core_weight, tools_weight,
#                 projects_weight, internship_weight, experience_weight, min_threshold)
#             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50))
#         ))
#         conn.commit(); conn.close()
#         flash("Job role created successfully!", "success")
#         return redirect(url_for("roles"))
#     return render_template("new_role.html")

# @app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
# def edit_role(role_id):
#     conn = get_db()
#     role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn.execute("""
#             UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
#                 project_keywords=?, internship_keywords=?, experience_keywords=?,
#                 core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
#                 experience_weight=?, min_threshold=?
#             WHERE id=?
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50)),
#             role_id
#         ))
#         conn.commit()
#         flash("Job role updated!", "success")
#         return redirect(url_for("roles"))
#     conn.close()
#     return render_template("edit_role.html", role=role)

# @app.route("/roles/<int:role_id>/delete", methods=["POST"])
# def delete_role(role_id):
#     conn = get_db()
#     conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
#     conn.commit(); conn.close()
#     flash("Role deleted.", "info")
#     return redirect(url_for("roles"))

# # ── UPLOAD & SCREEN ──

# @app.route("/upload", methods=["GET", "POST"])
# def upload():
#     conn  = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
#     conn.close()

#     if request.method == "POST":
#         role_id        = int(request.form.get("role_id", 0))
#         files          = request.files.getlist("resumes")
#         candidate_type = request.form.get("candidate_type", "experience")

#         require_internship = request.form.get("require_internship") == "on"
#         require_projects   = request.form.get("require_projects")   == "on"
#         min_years_raw      = request.form.get("min_years", "").strip()
#         max_years_raw      = request.form.get("max_years", "").strip()

#         screening_options = {
#             'require_internship': require_internship,
#             'require_projects':   require_projects,
#             'min_years': float(min_years_raw) if min_years_raw else None,
#             'max_years': float(max_years_raw) if max_years_raw else None,
#         }

#         if not role_id:
#             flash("Please select a job role.", "danger")
#             return redirect(url_for("upload"))
#         if not files or all(f.filename == "" for f in files):
#             flash("Please upload at least one file.", "danger")
#             return redirect(url_for("upload"))

#         conn  = get_db()
#         role  = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()

#         batch_id         = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + str(role_id) + "_" + candidate_type
#         batch_label_parts = [role['title'], candidate_type.title()]
#         if candidate_type == 'experience' and (min_years_raw or max_years_raw):
#             batch_label_parts.append(f"{min_years_raw or '0'}-{max_years_raw or '∞'} yrs")
#         batch_label = " | ".join(batch_label_parts)

#         processed = 0
#         errors    = 0

#         for f in files:
#             if f and allowed_file(f.filename):
#                 filename  = secure_filename(f.filename)
#                 ts        = datetime.now().strftime("%Y%m%d%H%M%S%f")
#                 save_name = f"{ts}_{filename}"
#                 filepath  = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
#                 f.save(filepath)

#                 result = screen_resume(filepath, role, candidate_type, screening_options)
#                 if result:
#                     conn.execute("""
#                         INSERT INTO candidates (job_role_id, candidate_type, batch_id, name, email, phone, raw_text,
#                             skills_found, tools_found, projects_found, internship_found, experience_found,
#                             internship_years, experience_years,
#                             core_score, tools_score, projects_score, internship_score, experience_score,
#                             total_score, status, rejection_reason, filename, uploaded_at)
#                         VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
#                     """, (
#                         role_id, candidate_type, batch_id,
#                         result["name"], result["email"], result["phone"], result["raw_text"],
#                         result["skills_found"], result["tools_found"], result["projects_found"],
#                         result["internship_found"], result["experience_found"],
#                         result.get("internship_years", ""), result.get("experience_years", ""),
#                         result["core_score"], result["tools_score"], result["projects_score"],
#                         result["internship_score"], result["experience_score"],
#                         result["total_score"], result["status"], result["rejection_reason"], save_name,
#                         datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#                     ))
#                     processed += 1
#                 else:
#                     errors += 1
#             else:
#                 errors += 1

#         if processed > 0:
#             now_local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#             conn.execute("""
#                 INSERT OR REPLACE INTO upload_batches (id, job_role_id, candidate_type, label, total_resumes, uploaded_at)
#                 VALUES (?,?,?,?,?,?)
#             """, (batch_id, role_id, candidate_type, batch_label, processed, now_local))

#         conn.commit(); conn.close()

#         flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
#         if processed > 0:
#             notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Screening Complete!',
#                 'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
#             }]
#         return redirect(url_for("results", batch_id=batch_id))

#     return render_template("upload.html", roles=roles)

# # ── RESULTS ──
# # ✅ FIX: All filters (role, type, status) now work correctly whether a batch
# #         is selected from the sidebar OR "All Results" is chosen.
# #         Previously role_id and type_filter were inside an `else` block under
# #         `if batch_id`, so they were completely ignored when a batch was active.

# @app.route("/results")
# def results():
#     role_id       = request.args.get("role_id", type=int)
#     status_filter = request.args.get("status", "all")
#     batch_id      = request.args.get("batch_id")
#     type_filter   = request.args.get("candidate_type", "all")

#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()

#     batches = conn.execute("""
#         SELECT ub.*, j.title as role_title
#         FROM upload_batches ub
#         JOIN job_roles j ON ub.job_role_id=j.id
#         ORDER BY ub.uploaded_at DESC
#     """).fetchall()

#     query  = "SELECT c.*, j.title as role_title, j.min_threshold FROM candidates c JOIN job_roles j ON c.job_role_id=j.id"
#     params = []
#     where  = []

#     # ── batch narrows the base set ──────────────────────────────────────────
#     if batch_id:
#         where.append("c.batch_id=?")
#         params.append(batch_id)

#     # ── these filters ALWAYS apply on top (batch selected or not) ───────────
#     if role_id:
#         where.append("c.job_role_id=?")
#         params.append(role_id)
#     if type_filter != "all":
#         where.append("c.candidate_type=?")
#         params.append(type_filter)
#     if status_filter != "all":
#         where.append("c.status=?")
#         params.append(status_filter)
#     # ────────────────────────────────────────────────────────────────────────

#     if where:
#         query += " WHERE " + " AND ".join(where)
#     query += " ORDER BY c.total_score DESC"

#     candidates = conn.execute(query, params).fetchall()
#     conn.close()

#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("results.html", candidates=candidates, roles=roles,
#                            selected_role=role_id, status_filter=status_filter,
#                            batches=batches, selected_batch=batch_id,
#                            type_filter=type_filter,
#                            popup_notifications=popup_notifications)

# @app.route("/candidate/<int:cid>")
# def candidate_detail(cid):
#     conn = get_db()
#     c = conn.execute("""
#         SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
#                j.project_keywords, j.internship_keywords, j.experience_keywords
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
#     conn.close()

#     if not c:
#         flash("Candidate not found.", "danger")
#         return redirect(url_for("results"))

#     parsed = dict(c)
#     for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
#                   "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
#         try:
#             parsed[field] = json.loads(c[field] or "[]")
#         except:
#             parsed[field] = []

#     return render_template("candidate_detail.html", c=parsed)

# @app.route("/candidate/<int:cid>/delete", methods=["POST"])
# def delete_candidate(cid):
#     conn = get_db()
#     conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
#     conn.commit(); conn.close()
#     flash("Candidate deleted.", "info")
#     return redirect(url_for("results"))

# # ── EMAIL ──

# @app.route("/send_emails", methods=["POST"])
# def send_emails():
#     candidate_ids = request.form.getlist("candidate_ids")
#     if not candidate_ids:
#         flash("No candidates selected.", "danger")
#         return redirect(url_for("results"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured. Please configure SMTP settings first.", "danger")
#         conn.close()
#         return redirect(url_for("email_settings_page"))

#     sent_count = 0
#     fail_count = 0

#     for cid in candidate_ids:
#         c = conn.execute("""
#             SELECT c.name, c.email, j.title as role_title
#             FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#         """, (cid,)).fetchone()
#         if c and c["email"]:
#             success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#             if success:
#                 conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                              (datetime.now(), cid))
#                 sent_count += 1
#             else:
#                 fail_count += 1

#     conn.commit(); conn.close()

#     if sent_count:
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Emails Sent!',
#             'message': f'Successfully sent {sent_count} interview invitation email(s).'
#         }]
#     if fail_count:
#         existing = session.get('popup_notifications', [])
#         existing.append({'type': 'error', 'title': 'Email Failed',
#                          'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'})
#         session['popup_notifications'] = existing

#     return redirect(url_for("results"))

# @app.route("/send_email_single/<int:cid>", methods=["POST"])
# def send_email_single(cid):
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
#     c        = conn.execute("""
#         SELECT c.name, c.email, j.title as role_title
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured.", "danger")
#     elif not c or not c["email"]:
#         flash("Candidate email not found.", "danger")
#     else:
#         success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#         if success:
#             conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                          (datetime.now(), cid))
#             conn.commit()
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Email Sent!',
#                 'message': f'Interview invitation sent to {c["name"]}.'
#             }]
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Email Failed',
#                 'message': f'Could not send email: {msg}'
#             }]

#     conn.close()
#     return redirect(url_for("candidate_detail", cid=cid))

# @app.route("/email-settings", methods=["GET", "POST"])
# def email_settings_page():
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if request.method == "POST":
#         conn.execute("""
#             UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
#                 sender_password=?, email_subject=?, email_body=? WHERE id=1
#         """, (
#             request.form["smtp_host"], int(request.form["smtp_port"]),
#             request.form["sender_email"], request.form["sender_password"],
#             request.form["email_subject"], request.form["email_body"]
#         ))
#         conn.commit()
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Settings Saved!',
#             'message': 'Email settings have been saved successfully.'
#         }]
#         return redirect(url_for("email_settings_page"))

#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("email_settings.html", settings=settings,
#                            popup_notifications=popup_notifications)

# # ── LOGIN ──

# @app.route("/login", methods=["GET", "POST"])
# def login():
#     if request.method == "POST":
#         username = request.form.get("username", "")
#         password = request.form.get("password", "")
#         if username == "admin" and password == "admin123":
#             session['logged_in'] = True
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Login Successful!',
#                 'message': f'Welcome back, {username}!'
#             }]
#             return redirect(url_for("index"))
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Login Failed',
#                 'message': 'Invalid username or password.'
#             }]
#             return redirect(url_for("login"))
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("login.html", popup_notifications=popup_notifications)

# # ── SEED DEFAULT ROLES ──

# @app.route("/seed-roles")
# def seed_roles():
#     conn     = get_db()
#     existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
#     if existing == 0:
#         for role in DEFAULT_ROLES:
#             conn.execute("""
#                 INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                     internship_keywords, experience_keywords)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (
#                 role["title"], role["description"],
#                 json.dumps(role["core_skills"]), json.dumps(role["tools"]),
#                 json.dumps(role["project_keywords"]), json.dumps(role["internship_keywords"]),
#                 json.dumps(role["experience_keywords"])
#             ))
#         conn.commit()
#         flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
#     else:
#         flash("Job roles already exist.", "info")
#     conn.close()
#     return redirect(url_for("roles"))

# # ── API ──

# @app.route("/api/dashboard")
# def api_dashboard():
#     conn = get_db()
#     rows = conn.execute("""
#         SELECT j.title, COUNT(c.id) as total,
#                SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
#                SUM(CASE WHEN c.status='rejected'    THEN 1 ELSE 0 END) as rejected
#         FROM job_roles j LEFT JOIN candidates c ON j.id=c.job_role_id
#         GROUP BY j.id
#     """).fetchall()
#     conn.close()
#     return jsonify([dict(r) for r in rows])

# # ─────────────────────────────────────────────
# # CUSTOM JINJA2 FILTERS
# # ─────────────────────────────────────────────

# @app.template_filter("fromjson")
# def fromjson_filter(value):
#     try:
#         return json.loads(value)
#     except:
#         return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)

# """
# Intelligent Resume Screening and Automated Interview Notification System
# MCA Final Year Project
# """

# from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
# import sqlite3
# import os
# import json
# import re
# import smtplib
# import io
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from datetime import datetime
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# app.secret_key = "resume_screener_mca_2024"

# # Config
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# # ─────────────────────────────────────────────
# # DATABASE SETUP
# # ─────────────────────────────────────────────

# def get_db():
#     conn = sqlite3.connect(DB_PATH)
#     conn.row_factory = sqlite3.Row
#     return conn

# def init_db():
#     conn = get_db()
#     c = conn.cursor()
    
#     c.executescript("""
#     CREATE TABLE IF NOT EXISTS job_roles (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         title TEXT NOT NULL,
#         description TEXT,
#         core_skills TEXT NOT NULL,
#         tools TEXT NOT NULL,
#         project_keywords TEXT NOT NULL,
#         internship_keywords TEXT NOT NULL,
#         experience_keywords TEXT NOT NULL,
#         core_weight REAL DEFAULT 0.40,
#         tools_weight REAL DEFAULT 0.25,
#         projects_weight REAL DEFAULT 0.15,
#         internship_weight REAL DEFAULT 0.10,
#         experience_weight REAL DEFAULT 0.10,
#         min_threshold INTEGER DEFAULT 50,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     CREATE TABLE IF NOT EXISTS candidates (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         job_role_id INTEGER NOT NULL,
#         candidate_type TEXT DEFAULT 'experience',  -- 'fresher' or 'experience'
#         batch_id TEXT,                              -- FIX #3: batch grouping
#         name TEXT,
#         email TEXT,
#         phone TEXT,
#         raw_text TEXT,
#         skills_found TEXT,
#         tools_found TEXT,
#         projects_found TEXT,
#         internship_found TEXT,
#         experience_found TEXT,
#         internship_years TEXT,
#         experience_years TEXT,
#         core_score REAL DEFAULT 0,
#         tools_score REAL DEFAULT 0,
#         projects_score REAL DEFAULT 0,
#         internship_score REAL DEFAULT 0,
#         experience_score REAL DEFAULT 0,
#         total_score REAL DEFAULT 0,
#         status TEXT DEFAULT 'pending',
#         rejection_reason TEXT,
#         email_sent INTEGER DEFAULT 0,
#         email_sent_at TIMESTAMP,
#         filename TEXT,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS upload_batches (
#         id TEXT PRIMARY KEY,
#         job_role_id INTEGER,
#         candidate_type TEXT DEFAULT 'experience',
#         label TEXT,
#         total_resumes INTEGER DEFAULT 0,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS email_settings (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         smtp_host TEXT DEFAULT 'smtp.gmail.com',
#         smtp_port INTEGER DEFAULT 587,
#         sender_email TEXT,
#         sender_password TEXT,
#         email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
#         email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
#     );

#     INSERT OR IGNORE INTO email_settings (id) VALUES (1);
#     """)

#     # Add batch_id column if upgrading from old DB
#     try:
#         c.execute("ALTER TABLE candidates ADD COLUMN batch_id TEXT")
#         conn.commit()
#     except:
#         pass
#     try:
#         c.execute("ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'")
#         conn.commit()
#     except:
#         pass
#     try:
#         c.execute("ALTER TABLE candidates ADD COLUMN internship_years TEXT")
#         conn.commit()
#     except:
#         pass
#     try:
#         c.execute("ALTER TABLE candidates ADD COLUMN experience_years TEXT")
#         conn.commit()
#     except:
#         pass
#     try:
#         c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
#             id TEXT PRIMARY KEY,
#             job_role_id INTEGER,
#             candidate_type TEXT DEFAULT 'experience',
#             label TEXT,
#             total_resumes INTEGER DEFAULT 0,
#             uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#             FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#         )""")
#         conn.commit()
#     except:
#         pass

#     conn.commit()
#     conn.close()

# # ─────────────────────────────────────────────
# # KEYWORD DATABASE FOR ALL IT ROLES
# # ─────────────────────────────────────────────

# DEFAULT_ROLES = [
#     {
#         "title": "Java Full Stack Developer",
#         "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
#         "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
#         "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
#         "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
#         "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
#         "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
#     },
#     {
#         "title": "Python Full Stack Developer",
#         "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
#         "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
#         "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
#         "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
#         "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
#         "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
#     },
#     {
#         "title": "MERN Stack Developer",
#         "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
#         "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
#         "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
#         "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
#         "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
#         "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
#     },
#     {
#         "title": "Data Analyst",
#         "description": "Analyzes data to derive business insights using statistical and visualization tools.",
#         "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
#         "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
#         "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
#         "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
#         "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
#     },
#     {
#         "title": "Data Scientist",
#         "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
#         "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
#         "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
#         "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
#         "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
#     },
#     {
#         "title": "AI/ML Engineer",
#         "description": "Designs, develops, and deploys machine learning models and AI systems.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
#         "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
#         "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
#         "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
#         "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
#     },
#     {
#         "title": "DevOps Engineer",
#         "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
#         "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
#         "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
#         "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
#         "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
#         "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
#     },
#     {
#         "title": "Cloud Engineer",
#         "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
#         "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
#         "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
#         "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
#         "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
#         "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
#     },
#     {
#         "title": "Software Tester",
#         "description": "Ensures software quality through manual and automated testing methodologies.",
#         "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
#         "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
#         "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
#         "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
#         "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
#     },
#     {
#         "title": "Database Developer",
#         "description": "Designs, develops, and optimizes databases for scalability and performance.",
#         "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
#         "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
#         "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
#         "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
#         "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
#     }
# ]

# # ─────────────────────────────────────────────
# # RESUME PARSING
# # ─────────────────────────────────────────────

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(filepath):
#     text = ""
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#     except Exception as e:
#         print(f"PDF error: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ""
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX error: {e}")
#     return text

# def extract_text(filepath):
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_name(text):
#     lines = [l.strip() for l in text.split("\n") if l.strip()]
#     for line in lines[:5]:
#         if re.search(r"[@:/.()]|\d{5,}", line):
#             continue
#         if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
#             continue
#         words = line.split()
#         if 1 < len(words) <= 5:
#             return line.strip()
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
#     return match.group() if match else ""

# def extract_phone(text):
#     match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
#     return match.group().strip() if match else ""

# def extract_internship_years(text):
#     """Extract internship duration from resume text."""
#     patterns = [
#         r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
#         r'internship[^.]*?(\d+)\s*month',
#         r'(\d+)\s*month[^.]*?internship',
#         r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             val = match.group(1)
#             if 'month' in pattern:
#                 months = float(val)
#                 years = round(months / 12, 1)
#                 return f"{val} months ({years} years)"
#             return f"{val} years"
#     # Check for internship section with dates
#     intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
#     if intern_section:
#         try:
#             start = int(intern_section.group(1))
#             end_str = intern_section.group(2)
#             end = datetime.now().year if end_str == 'present' else int(end_str)
#             years = end - start
#             return f"{years} year(s)"
#         except:
#             pass
#     return ""

# def extract_experience_years(text):
#     """Extract work experience years from resume."""
#     patterns = [
#         r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
#         r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             return f"{match.group(1)} years"
#     return ""

# def find_keywords_in_text(text, keywords):
#     text_lower = text.lower()
#     found = []
#     for kw in keywords:
#         pattern = r'\b' + re.escape(kw.lower()) + r'\b'
#         if re.search(pattern, text_lower):
#             found.append(kw)
#     return found

# def has_internship_in_resume(text):
#     """Check if resume contains internship experience."""
#     patterns = [
#         r'\binternship\b',
#         r'\bintern\b',
#         r'\btrainee\b',
#         r'\bindustry training\b',
#         r'\bsummer training\b',
#         r'\bproject trainee\b',
#     ]
#     text_lower = text.lower()
#     for p in patterns:
#         if re.search(p, text_lower):
#             return True
#     return False

# # ─────────────────────────────────────────────
# # SCORING ALGORITHM
# # ─────────────────────────────────────────────

# def calculate_score_fresher(candidate_data, role, screening_options):
#     """
#     Fresher scoring:
#     - Core Skills: configurable weight
#     - Projects: configurable weight
#     - Internship: configurable (can be disabled)
#     Only these 3 sections matter for freshers.
#     """
#     require_internship = screening_options.get('require_internship', True)
    
#     core_total = len(json.loads(role["core_skills"]))
#     proj_total = len(json.loads(role["project_keywords"]))
#     intern_total = len(json.loads(role["internship_keywords"]))
    
#     def section_pct(found, total):
#         if not total:
#             return 0
#         return min((len(found) / total) * 100, 100)
    
#     core_score = section_pct(candidate_data["skills_found"], core_total)
#     proj_score = section_pct(candidate_data["projects_found"], proj_total)
    
#     if require_internship:
#         intern_score = section_pct(candidate_data["internship_found"], intern_total)
#         # Weights: core 55%, projects 30%, internship 15%
#         total = core_score * 0.55 + proj_score * 0.30 + intern_score * 0.15
#     else:
#         intern_score = 0
#         # Weights: core 65%, projects 35%
#         total = core_score * 0.65 + proj_score * 0.35
    
#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": 0,
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": 0,
#         "total_score": round(total, 1)
#     }

# def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
#     """
#     Experience candidate scoring with year range filter.
#     """
#     weights = {
#         "core": role["core_weight"],
#         "tools": role["tools_weight"],
#         "projects": role["projects_weight"],
#         "internship": role["internship_weight"],
#         "experience": role["experience_weight"]
#     }

#     def section_pct(found, total):
#         if not total:
#             return 0
#         return min((len(found) / total) * 100, 100)

#     core_score = section_pct(candidate_data["skills_found"], len(json.loads(role["core_skills"])))
#     tools_score = section_pct(candidate_data["tools_found"], len(json.loads(role["tools"])))
#     proj_score = section_pct(candidate_data["projects_found"], len(json.loads(role["project_keywords"])))
#     intern_score = section_pct(candidate_data["internship_found"], len(json.loads(role["internship_keywords"])))
#     exp_score = section_pct(candidate_data["experience_found"], len(json.loads(role["experience_keywords"])))

#     total = (
#         core_score * weights["core"] +
#         tools_score * weights["tools"] +
#         proj_score * weights["projects"] +
#         intern_score * weights["internship"] +
#         exp_score * weights["experience"]
#     )

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": round(exp_score, 1),
#         "total_score": round(total, 1)
#     }

# def calculate_score(candidate_data, role):
#     return calculate_score_experience(candidate_data, role)

# def check_experience_years_range(exp_years_str, min_years, max_years):
#     """Check if candidate's experience years fall in range."""
#     if not exp_years_str:
#         return False
#     match = re.search(r'(\d+\.?\d*)', exp_years_str)
#     if not match:
#         return False
#     years = float(match.group(1))
#     if min_years is not None and years < min_years:
#         return False
#     if max_years is not None and years > max_years:
#         return False
#     return True

# def generate_rejection_reason(scores, threshold, candidate_type='experience'):
#     reasons = []
#     if scores["core_score"] < 20:
#         reasons.append("insufficient core technical skills")
#     if candidate_type == 'fresher':
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#     else:
#         if scores["tools_score"] < 20:
#             reasons.append("limited relevant tool experience")
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#         if scores["internship_score"] < 10 and scores["experience_score"] < 10:
#             reasons.append("no relevant internship or work experience")
    
#     if not reasons:
#         reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
    
#     return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

# def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
#     """Full pipeline: extract → parse → score → decide."""
#     if screening_options is None:
#         screening_options = {}
    
#     text = extract_text(filepath)
#     if not text:
#         return None
    
#     name = extract_name(text)
#     email = extract_email(text)
#     phone = extract_phone(text)
    
#     core_skills = json.loads(role["core_skills"])
#     tools = json.loads(role["tools"])
#     proj_kw = json.loads(role["project_keywords"])
#     intern_kw = json.loads(role["internship_keywords"])
#     exp_kw = json.loads(role["experience_keywords"])
    
#     skills_found = find_keywords_in_text(text, core_skills)
#     tools_found = find_keywords_in_text(text, tools)
#     projects_found = find_keywords_in_text(text, proj_kw)
#     internship_found = find_keywords_in_text(text, intern_kw)
#     experience_found = find_keywords_in_text(text, exp_kw)
    
#     internship_years = extract_internship_years(text)
#     experience_years = extract_experience_years(text)
    
#     candidate_data = {
#         "skills_found": skills_found,
#         "tools_found": tools_found,
#         "projects_found": projects_found,
#         "internship_found": internship_found,
#         "experience_found": experience_found
#     }
    
#     threshold = role["min_threshold"]
    
#     if candidate_type == 'fresher':
#         scores = calculate_score_fresher(candidate_data, role, screening_options)
#         status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#         rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
#     else:
#         # Experience candidate: check year range
#         min_years = screening_options.get('min_years')
#         max_years = screening_options.get('max_years')
        
#         scores = calculate_score_experience(candidate_data, role, min_years, max_years)
        
#         # Year range filter: if specified and candidate doesn't match, reject
#         if (min_years is not None or max_years is not None) and experience_years:
#             if not check_experience_years_range(experience_years, min_years, max_years):
#                 status = "rejected"
#                 yr_range = ""
#                 if min_years and max_years:
#                     yr_range = f"{min_years}-{max_years} years"
#                 elif min_years:
#                     yr_range = f"{min_years}+ years"
#                 elif max_years:
#                     yr_range = f"up to {max_years} years"
#                 rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
#             else:
#                 status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#                 rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
#         else:
#             status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#             rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
    
#     return {
#         "name": name,
#         "email": email,
#         "phone": phone,
#         "raw_text": text[:2000],
#         "skills_found": json.dumps(skills_found),
#         "tools_found": json.dumps(tools_found),
#         "projects_found": json.dumps(projects_found),
#         "internship_found": json.dumps(internship_found),
#         "experience_found": json.dumps(experience_found),
#         "internship_years": internship_years,
#         "experience_years": experience_years,
#         "status": status,
#         "rejection_reason": rejection_reason,
#         **scores
#     }

# # ─────────────────────────────────────────────
# # EMAIL INTEGRATION
# # ─────────────────────────────────────────────

# def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
#     subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
#     body = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)
    
#     msg = MIMEMultipart("alternative")
#     msg["Subject"] = subject
#     msg["From"] = settings["sender_email"]
#     msg["To"] = candidate_email
    
#     html_body = f"""
#     <html><body>
#     <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
#       <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
#         <h2 style="color:white; margin:0;">Interview Invitation</h2>
#       </div>
#       <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
#         {body.replace(chr(10), '<br>')}
#       </div>
#     </div>
#     </body></html>
#     """
    
#     msg.attach(MIMEText(body, "plain"))
#     msg.attach(MIMEText(html_body, "html"))
    
#     try:
#         server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
#         server.ehlo()
#         server.starttls()
#         server.login(settings["sender_email"], settings["sender_password"])
#         server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
#         server.quit()
#         return True, "Email sent successfully"
#     except Exception as e:
#         return False, str(e)

# # ─────────────────────────────────────────────
# # ROUTES
# # ─────────────────────────────────────────────

# @app.route("/")
# def index():
#     conn = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     total_candidates = conn.execute("SELECT COUNT(*) as c FROM candidates").fetchone()["c"]
#     shortlisted = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted'").fetchone()["c"]
#     rejected = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='rejected'").fetchone()["c"]
#     conn.close()
    
#     # FIX #2: Get popup notifications from session
#     popup_notifications = session.pop('popup_notifications', [])
    
#     return render_template("index.html", roles=roles,
#                            total=total_candidates, shortlisted=shortlisted, rejected=rejected,
#                            popup_notifications=popup_notifications)

# # ── JOB ROLES ──

# @app.route("/roles")
# def roles():
#     conn = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     conn.close()
#     return render_template("roles.html", roles=roles)

# @app.route("/roles/new", methods=["GET", "POST"])
# def new_role():
#     if request.method == "POST":
#         def parse_list(field):
#             raw = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
        
#         conn = get_db()
#         conn.execute("""
#             INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                 internship_keywords, experience_keywords, core_weight, tools_weight,
#                 projects_weight, internship_weight, experience_weight, min_threshold)
#             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
#         """, (
#             request.form["title"],
#             request.form.get("description", ""),
#             parse_list("core_skills"),
#             parse_list("tools"),
#             parse_list("project_keywords"),
#             parse_list("internship_keywords"),
#             parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50))
#         ))
#         conn.commit()
#         conn.close()
#         flash("Job role created successfully!", "success")
#         return redirect(url_for("roles"))
    
#     return render_template("new_role.html")

# @app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
# def edit_role(role_id):
#     conn = get_db()
#     role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
    
#     if request.method == "POST":
#         def parse_list(field):
#             raw = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
        
#         conn.execute("""
#             UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
#                 project_keywords=?, internship_keywords=?, experience_keywords=?,
#                 core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
#                 experience_weight=?, min_threshold=?
#             WHERE id=?
#         """, (
#             request.form["title"],
#             request.form.get("description", ""),
#             parse_list("core_skills"),
#             parse_list("tools"),
#             parse_list("project_keywords"),
#             parse_list("internship_keywords"),
#             parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50)),
#             role_id
#         ))
#         conn.commit()
#         flash("Job role updated!", "success")
#         return redirect(url_for("roles"))
    
#     conn.close()
#     return render_template("edit_role.html", role=role)

# @app.route("/roles/<int:role_id>/delete", methods=["POST"])
# def delete_role(role_id):
#     conn = get_db()
#     conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
#     conn.commit()
#     conn.close()
#     flash("Role deleted.", "info")
#     return redirect(url_for("roles"))

# # ── UPLOAD & SCREEN ──
# # FIX #1 + #4 + #5: Separate Fresher/Experience screening with options

# @app.route("/upload", methods=["GET", "POST"])
# def upload():
#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
#     conn.close()
    
#     if request.method == "POST":
#         role_id = int(request.form.get("role_id", 0))
#         files = request.files.getlist("resumes")
#         candidate_type = request.form.get("candidate_type", "experience")  # 'fresher' or 'experience'
        
#         # Fresher-specific options
#         require_internship = request.form.get("require_internship") == "on"
#         require_projects = request.form.get("require_projects") == "on"
        
#         # Experience-specific options
#         min_years = request.form.get("min_years", "").strip()
#         max_years = request.form.get("max_years", "").strip()
        
#         screening_options = {
#             'require_internship': require_internship,
#             'require_projects': require_projects,
#             'min_years': float(min_years) if min_years else None,
#             'max_years': float(max_years) if max_years else None,
#         }
        
#         if not role_id:
#             flash("Please select a job role.", "danger")
#             return redirect(url_for("upload"))
#         if not files or all(f.filename == "" for f in files):
#             flash("Please upload at least one file.", "danger")
#             return redirect(url_for("upload"))
        
#         conn = get_db()
#         role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
        
#         # FIX #3: Create a unique batch ID for this upload session
#         batch_id = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + str(role_id) + "_" + candidate_type
#         batch_label_parts = [role['title'], candidate_type.title()]
#         if candidate_type == 'experience' and (min_years or max_years):
#             yr_range = f"{min_years or '0'}-{max_years or '∞'} yrs"
#             batch_label_parts.append(yr_range)
#         batch_label = " | ".join(batch_label_parts)
        
#         processed = 0
#         errors = 0
        
#         for f in files:
#             if f and allowed_file(f.filename):
#                 filename = secure_filename(f.filename)
#                 ts = datetime.now().strftime("%Y%m%d%H%M%S%f")
#                 save_name = f"{ts}_{filename}"
#                 filepath = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
#                 f.save(filepath)
                
#                 result = screen_resume(filepath, role, candidate_type, screening_options)
                
#                 if result:
#                     conn.execute("""
#                         INSERT INTO candidates (job_role_id, candidate_type, batch_id, name, email, phone, raw_text,
#                             skills_found, tools_found, projects_found, internship_found, experience_found,
#                             internship_years, experience_years,
#                             core_score, tools_score, projects_score, internship_score, experience_score,
#                             total_score, status, rejection_reason, filename)
#                         VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
#                     """, (
#                         role_id, candidate_type, batch_id,
#                         result["name"], result["email"], result["phone"], result["raw_text"],
#                         result["skills_found"], result["tools_found"], result["projects_found"],
#                         result["internship_found"], result["experience_found"],
#                         result.get("internship_years", ""), result.get("experience_years", ""),
#                         result["core_score"], result["tools_score"], result["projects_score"],
#                         result["internship_score"], result["experience_score"],
#                         result["total_score"], result["status"], result["rejection_reason"], save_name
#                     ))
#                     processed += 1
#                 else:
#                     errors += 1
#             else:
#                 errors += 1
        
#         # Save batch info
#         if processed > 0:
#             conn.execute("""
#                 INSERT OR REPLACE INTO upload_batches (id, job_role_id, candidate_type, label, total_resumes)
#                 VALUES (?,?,?,?,?)
#             """, (batch_id, role_id, candidate_type, batch_label, processed))
        
#         conn.commit()
#         conn.close()
        
#         flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
#         # FIX #2: Store popup notification in session
#         if processed > 0:
#             notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Screening Complete!',
#                 'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
#             }]
#         return redirect(url_for("results", batch_id=batch_id))
    
#     return render_template("upload.html", roles=roles)

# # ── RESULTS ──
# # FIX #3: Batch-separated results

# @app.route("/results")
# def results():
#     role_id = request.args.get("role_id", type=int)
#     status_filter = request.args.get("status", "all")
#     batch_id = request.args.get("batch_id")
#     type_filter = request.args.get("candidate_type", "all")
    
#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
    
#     # Get all batches for the sidebar
#     batches = conn.execute("""
#         SELECT ub.*, j.title as role_title 
#         FROM upload_batches ub 
#         JOIN job_roles j ON ub.job_role_id=j.id 
#         ORDER BY ub.uploaded_at DESC
#     """).fetchall()
    
#     query = "SELECT c.*, j.title as role_title, j.min_threshold FROM candidates c JOIN job_roles j ON c.job_role_id=j.id"
#     params = []
#     where = []
    
#     if batch_id:
#         where.append("c.batch_id=?")
#         params.append(batch_id)
#     else:
#         if role_id:
#             where.append("c.job_role_id=?")
#             params.append(role_id)
#         if type_filter != "all":
#             where.append("c.candidate_type=?")
#             params.append(type_filter)
    
#     if status_filter != "all":
#         where.append("c.status=?")
#         params.append(status_filter)
#     if where:
#         query += " WHERE " + " AND ".join(where)
#     query += " ORDER BY c.total_score DESC"
    
#     candidates = conn.execute(query, params).fetchall()
#     conn.close()
    
#     # FIX #2: Get popup notifications from session
#     popup_notifications = session.pop('popup_notifications', [])
    
#     return render_template("results.html", candidates=candidates, roles=roles,
#                            selected_role=role_id, status_filter=status_filter,
#                            batches=batches, selected_batch=batch_id,
#                            type_filter=type_filter,
#                            popup_notifications=popup_notifications)

# @app.route("/candidate/<int:cid>")
# def candidate_detail(cid):
#     conn = get_db()
#     c = conn.execute("""
#         SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
#                j.project_keywords, j.internship_keywords, j.experience_keywords
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
#     conn.close()
    
#     if not c:
#         flash("Candidate not found.", "danger")
#         return redirect(url_for("results"))
    
#     parsed = dict(c)
#     for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
#                   "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
#         try:
#             parsed[field] = json.loads(c[field] or "[]")
#         except:
#             parsed[field] = []
    
#     return render_template("candidate_detail.html", c=parsed)

# @app.route("/candidate/<int:cid>/delete", methods=["POST"])
# def delete_candidate(cid):
#     conn = get_db()
#     conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
#     conn.commit()
#     conn.close()
#     flash("Candidate deleted.", "info")
#     return redirect(url_for("results"))

# # ── EMAIL ──

# @app.route("/send_emails", methods=["POST"])
# def send_emails():
#     candidate_ids = request.form.getlist("candidate_ids")
#     if not candidate_ids:
#         flash("No candidates selected.", "danger")
#         return redirect(url_for("results"))
    
#     conn = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
    
#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured. Please configure SMTP settings first.", "danger")
#         conn.close()
#         return redirect(url_for("email_settings_page"))
    
#     sent_count = 0
#     fail_count = 0
    
#     for cid in candidate_ids:
#         c = conn.execute("""
#             SELECT c.name, c.email, j.title as role_title
#             FROM candidates c JOIN job_roles j ON c.job_role_id=j.id
#             WHERE c.id=?
#         """, (cid,)).fetchone()
        
#         if c and c["email"]:
#             success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#             if success:
#                 conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                              (datetime.now(), cid))
#                 sent_count += 1
#             else:
#                 fail_count += 1
    
#     conn.commit()
#     conn.close()
    
#     # FIX #2: Popup notification
#     if sent_count:
#         session['popup_notifications'] = [{
#             'type': 'success',
#             'title': 'Emails Sent!',
#             'message': f'Successfully sent {sent_count} interview invitation email(s).'
#         }]
#     if fail_count:
#         existing = session.get('popup_notifications', [])
#         existing.append({
#             'type': 'error',
#             'title': 'Email Failed',
#             'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'
#         })
#         session['popup_notifications'] = existing
    
#     return redirect(url_for("results"))

# @app.route("/send_email_single/<int:cid>", methods=["POST"])
# def send_email_single(cid):
#     conn = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
#     c = conn.execute("""
#         SELECT c.name, c.email, j.title as role_title
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
    
#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured.", "danger")
#     elif not c or not c["email"]:
#         flash("Candidate email not found.", "danger")
#     else:
#         success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#         if success:
#             conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                          (datetime.now(), cid))
#             conn.commit()
#             # FIX #2: Popup
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Email Sent!',
#                 'message': f'Interview invitation sent to {c["name"]}.'
#             }]
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error',
#                 'title': 'Email Failed',
#                 'message': f'Could not send email: {msg}'
#             }]
    
#     conn.close()
#     return redirect(url_for("candidate_detail", cid=cid))

# @app.route("/email-settings", methods=["GET", "POST"])
# def email_settings_page():
#     conn = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
    
#     if request.method == "POST":
#         conn.execute("""
#             UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
#                 sender_password=?, email_subject=?, email_body=? WHERE id=1
#         """, (
#             request.form["smtp_host"],
#             int(request.form["smtp_port"]),
#             request.form["sender_email"],
#             request.form["sender_password"],
#             request.form["email_subject"],
#             request.form["email_body"]
#         ))
#         conn.commit()
#         # FIX #2: Popup notification for email settings saved
#         session['popup_notifications'] = [{
#             'type': 'success',
#             'title': 'Settings Saved!',
#             'message': 'Email settings have been saved successfully.'
#         }]
#         return redirect(url_for("email_settings_page"))
    
#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("email_settings.html", settings=settings, popup_notifications=popup_notifications)

# # ── LOGIN (FIX #2: Popup after login) ──

# @app.route("/login", methods=["GET", "POST"])
# def login():
#     if request.method == "POST":
#         username = request.form.get("username", "")
#         password = request.form.get("password", "")
#         # Simple demo auth - replace with real auth
#         if username == "admin" and password == "admin123":
#             session['logged_in'] = True
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Login Successful!',
#                 'message': f'Welcome back, {username}!'
#             }]
#             return redirect(url_for("index"))
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error',
#                 'title': 'Login Failed',
#                 'message': 'Invalid username or password.'
#             }]
#             return redirect(url_for("login"))
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("login.html", popup_notifications=popup_notifications)

# # ── SEED DEFAULT ROLES ──

# @app.route("/seed-roles")
# def seed_roles():
#     conn = get_db()
#     existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
#     if existing == 0:
#         for role in DEFAULT_ROLES:
#             conn.execute("""
#                 INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                     internship_keywords, experience_keywords)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (
#                 role["title"], role["description"],
#                 json.dumps(role["core_skills"]),
#                 json.dumps(role["tools"]),
#                 json.dumps(role["project_keywords"]),
#                 json.dumps(role["internship_keywords"]),
#                 json.dumps(role["experience_keywords"])
#             ))
#         conn.commit()
#         flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
#     else:
#         flash("Job roles already exist.", "info")
#     conn.close()
#     return redirect(url_for("roles"))

# # ── API ENDPOINTS ──

# @app.route("/api/dashboard")
# def api_dashboard():
#     conn = get_db()
#     roles = conn.execute("""
#         SELECT j.title, COUNT(c.id) as total,
#                SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
#                SUM(CASE WHEN c.status='rejected' THEN 1 ELSE 0 END) as rejected
#         FROM job_roles j LEFT JOIN candidates c ON j.id=c.job_role_id
#         GROUP BY j.id
#     """).fetchall()
#     conn.close()
#     return jsonify([dict(r) for r in roles])

# # ─────────────────────────────────────────────
# # CUSTOM JINJA2 FILTERS
# # ─────────────────────────────────────────────

# @app.template_filter("fromjson")
# def fromjson_filter(value):
#     try:
#         return json.loads(value)
#     except:
#         return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)

# """
# Intelligent Resume Screening and Automated Interview Notification System
# MCA Final Year Project
# """

# from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
# import sqlite3
# import os
# import json
# import re
# import smtplib
# import io
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from datetime import datetime
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# app.secret_key = "resume_screener_mca_2024"

# # Config
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# # ─────────────────────────────────────────────
# # DATABASE SETUP
# # ─────────────────────────────────────────────

# def get_db():
#     conn = sqlite3.connect(DB_PATH)
#     conn.row_factory = sqlite3.Row
#     return conn

# def init_db():
#     conn = get_db()
#     c = conn.cursor()
    
#     c.executescript("""
#     CREATE TABLE IF NOT EXISTS job_roles (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         title TEXT NOT NULL,
#         description TEXT,
#         core_skills TEXT NOT NULL,
#         tools TEXT NOT NULL,
#         project_keywords TEXT NOT NULL,
#         internship_keywords TEXT NOT NULL,
#         experience_keywords TEXT NOT NULL,
#         core_weight REAL DEFAULT 0.40,
#         tools_weight REAL DEFAULT 0.25,
#         projects_weight REAL DEFAULT 0.15,
#         internship_weight REAL DEFAULT 0.10,
#         experience_weight REAL DEFAULT 0.10,
#         min_threshold INTEGER DEFAULT 50,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     CREATE TABLE IF NOT EXISTS candidates (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         job_role_id INTEGER NOT NULL,
#         candidate_type TEXT DEFAULT 'experience',
#         batch_id TEXT,
#         name TEXT,
#         email TEXT,
#         phone TEXT,
#         raw_text TEXT,
#         skills_found TEXT,
#         tools_found TEXT,
#         projects_found TEXT,
#         internship_found TEXT,
#         experience_found TEXT,
#         internship_years TEXT,
#         experience_years TEXT,
#         core_score REAL DEFAULT 0,
#         tools_score REAL DEFAULT 0,
#         projects_score REAL DEFAULT 0,
#         internship_score REAL DEFAULT 0,
#         experience_score REAL DEFAULT 0,
#         total_score REAL DEFAULT 0,
#         status TEXT DEFAULT 'pending',
#         rejection_reason TEXT,
#         email_sent INTEGER DEFAULT 0,
#         email_sent_at TIMESTAMP,
#         filename TEXT,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS upload_batches (
#         id TEXT PRIMARY KEY,
#         job_role_id INTEGER,
#         candidate_type TEXT DEFAULT 'experience',
#         label TEXT,
#         total_resumes INTEGER DEFAULT 0,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS email_settings (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         smtp_host TEXT DEFAULT 'smtp.gmail.com',
#         smtp_port INTEGER DEFAULT 587,
#         sender_email TEXT,
#         sender_password TEXT,
#         email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
#         email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
#     );

#     INSERT OR IGNORE INTO email_settings (id) VALUES (1);
#     """)

#     # Add columns if upgrading from old DB
#     for col_sql in [
#         "ALTER TABLE candidates ADD COLUMN batch_id TEXT",
#         "ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'",
#         "ALTER TABLE candidates ADD COLUMN internship_years TEXT",
#         "ALTER TABLE candidates ADD COLUMN experience_years TEXT",
#     ]:
#         try:
#             c.execute(col_sql)
#             conn.commit()
#         except:
#             pass

#     try:
#         c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
#             id TEXT PRIMARY KEY,
#             job_role_id INTEGER,
#             candidate_type TEXT DEFAULT 'experience',
#             label TEXT,
#             total_resumes INTEGER DEFAULT 0,
#             uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#             FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#         )""")
#         conn.commit()
#     except:
#         pass

#     conn.commit()
#     conn.close()

# # ─────────────────────────────────────────────
# # KEYWORD DATABASE FOR ALL IT ROLES
# # ─────────────────────────────────────────────

# DEFAULT_ROLES = [
#     {
#         "title": "Java Full Stack Developer",
#         "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
#         "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
#         "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
#         "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
#         "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
#         "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
#     },
#     {
#         "title": "Python Full Stack Developer",
#         "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
#         "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
#         "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
#         "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
#         "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
#         "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
#     },
#     {
#         "title": "MERN Stack Developer",
#         "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
#         "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
#         "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
#         "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
#         "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
#         "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
#     },
#     {
#         "title": "Data Analyst",
#         "description": "Analyzes data to derive business insights using statistical and visualization tools.",
#         "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
#         "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
#         "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
#         "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
#         "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
#     },
#     {
#         "title": "Data Scientist",
#         "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
#         "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
#         "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
#         "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
#         "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
#     },
#     {
#         "title": "AI/ML Engineer",
#         "description": "Designs, develops, and deploys machine learning models and AI systems.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
#         "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
#         "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
#         "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
#         "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
#     },
#     {
#         "title": "DevOps Engineer",
#         "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
#         "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
#         "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
#         "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
#         "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
#         "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
#     },
#     {
#         "title": "Cloud Engineer",
#         "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
#         "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
#         "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
#         "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
#         "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
#         "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
#     },
#     {
#         "title": "Software Tester",
#         "description": "Ensures software quality through manual and automated testing methodologies.",
#         "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
#         "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
#         "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
#         "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
#         "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
#     },
#     {
#         "title": "Database Developer",
#         "description": "Designs, develops, and optimizes databases for scalability and performance.",
#         "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
#         "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
#         "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
#         "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
#         "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
#     }
# ]

# # ─────────────────────────────────────────────
# # RESUME PARSING
# # ─────────────────────────────────────────────

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(filepath):
#     text = ""
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#     except Exception as e:
#         print(f"PDF error: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ""
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX error: {e}")
#     return text

# def extract_text(filepath):
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_name(text):
#     lines = [l.strip() for l in text.split("\n") if l.strip()]
#     for line in lines[:5]:
#         if re.search(r"[@:/.()]|\d{5,}", line):
#             continue
#         if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
#             continue
#         words = line.split()
#         if 1 < len(words) <= 5:
#             return line.strip()
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
#     return match.group() if match else ""

# def extract_phone(text):
#     match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
#     return match.group().strip() if match else ""

# def extract_internship_years(text):
#     patterns = [
#         r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
#         r'internship[^.]*?(\d+)\s*month',
#         r'(\d+)\s*month[^.]*?internship',
#         r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             val = match.group(1)
#             if 'month' in pattern:
#                 months = float(val)
#                 years = round(months / 12, 1)
#                 return f"{val} months ({years} years)"
#             return f"{val} years"
#     intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
#     if intern_section:
#         try:
#             start = int(intern_section.group(1))
#             end_str = intern_section.group(2)
#             end = datetime.now().year if end_str == 'present' else int(end_str)
#             years = end - start
#             return f"{years} year(s)"
#         except:
#             pass
#     return ""

# def extract_experience_years(text):
#     patterns = [
#         r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
#         r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             return f"{match.group(1)} years"
#     return ""

# def find_keywords_in_text(text, keywords):
#     text_lower = text.lower()
#     found = []
#     for kw in keywords:
#         pattern = r'\b' + re.escape(kw.lower()) + r'\b'
#         if re.search(pattern, text_lower):
#             found.append(kw)
#     return found

# def has_internship_in_resume(text):
#     patterns = [
#         r'\binternship\b', r'\bintern\b', r'\btrainee\b',
#         r'\bindustry training\b', r'\bsummer training\b', r'\bproject trainee\b',
#     ]
#     text_lower = text.lower()
#     for p in patterns:
#         if re.search(p, text_lower):
#             return True
#     return False

# # ─────────────────────────────────────────────
# # SCORING ALGORITHM
# # ─────────────────────────────────────────────

# def calculate_score_fresher(candidate_data, role, screening_options):
#     require_internship = screening_options.get('require_internship', True)
#     core_total = len(json.loads(role["core_skills"]))
#     proj_total = len(json.loads(role["project_keywords"]))
#     intern_total = len(json.loads(role["internship_keywords"]))

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score = section_pct(candidate_data["skills_found"], core_total)
#     proj_score  = section_pct(candidate_data["projects_found"], proj_total)

#     if require_internship:
#         intern_score = section_pct(candidate_data["internship_found"], intern_total)
#         total = core_score * 0.55 + proj_score * 0.30 + intern_score * 0.15
#     else:
#         intern_score = 0
#         total = core_score * 0.65 + proj_score * 0.35

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": 0,
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": 0,
#         "total_score": round(total, 1)
#     }

# def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
#     weights = {
#         "core": role["core_weight"],
#         "tools": role["tools_weight"],
#         "projects": role["projects_weight"],
#         "internship": role["internship_weight"],
#         "experience": role["experience_weight"]
#     }

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score   = section_pct(candidate_data["skills_found"],      len(json.loads(role["core_skills"])))
#     tools_score  = section_pct(candidate_data["tools_found"],       len(json.loads(role["tools"])))
#     proj_score   = section_pct(candidate_data["projects_found"],    len(json.loads(role["project_keywords"])))
#     intern_score = section_pct(candidate_data["internship_found"],  len(json.loads(role["internship_keywords"])))
#     exp_score    = section_pct(candidate_data["experience_found"],  len(json.loads(role["experience_keywords"])))

#     total = (
#         core_score   * weights["core"] +
#         tools_score  * weights["tools"] +
#         proj_score   * weights["projects"] +
#         intern_score * weights["internship"] +
#         exp_score    * weights["experience"]
#     )

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": round(exp_score, 1),
#         "total_score": round(total, 1)
#     }

# def calculate_score(candidate_data, role):
#     return calculate_score_experience(candidate_data, role)

# def check_experience_years_range(exp_years_str, min_years, max_years):
#     if not exp_years_str: return False
#     match = re.search(r'(\d+\.?\d*)', exp_years_str)
#     if not match: return False
#     years = float(match.group(1))
#     if min_years is not None and years < min_years: return False
#     if max_years is not None and years > max_years: return False
#     return True

# def generate_rejection_reason(scores, threshold, candidate_type='experience'):
#     reasons = []
#     if scores["core_score"] < 20:
#         reasons.append("insufficient core technical skills")
#     if candidate_type == 'fresher':
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#     else:
#         if scores["tools_score"] < 20:
#             reasons.append("limited relevant tool experience")
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#         if scores["internship_score"] < 10 and scores["experience_score"] < 10:
#             reasons.append("no relevant internship or work experience")
#     if not reasons:
#         reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
#     return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

# def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
#     if screening_options is None:
#         screening_options = {}

#     text = extract_text(filepath)
#     if not text:
#         return None

#     name  = extract_name(text)
#     email = extract_email(text)
#     phone = extract_phone(text)

#     core_skills = json.loads(role["core_skills"])
#     tools       = json.loads(role["tools"])
#     proj_kw     = json.loads(role["project_keywords"])
#     intern_kw   = json.loads(role["internship_keywords"])
#     exp_kw      = json.loads(role["experience_keywords"])

#     skills_found      = find_keywords_in_text(text, core_skills)
#     tools_found       = find_keywords_in_text(text, tools)
#     projects_found    = find_keywords_in_text(text, proj_kw)
#     internship_found  = find_keywords_in_text(text, intern_kw)
#     experience_found  = find_keywords_in_text(text, exp_kw)

#     internship_years  = extract_internship_years(text)
#     experience_years  = extract_experience_years(text)

#     candidate_data = {
#         "skills_found": skills_found,
#         "tools_found": tools_found,
#         "projects_found": projects_found,
#         "internship_found": internship_found,
#         "experience_found": experience_found
#     }

#     threshold = role["min_threshold"]

#     if candidate_type == 'fresher':
#         scores = calculate_score_fresher(candidate_data, role, screening_options)
#         status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#         rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
#     else:
#         min_years = screening_options.get('min_years')
#         max_years = screening_options.get('max_years')
#         scores = calculate_score_experience(candidate_data, role, min_years, max_years)

#         if (min_years is not None or max_years is not None) and experience_years:
#             if not check_experience_years_range(experience_years, min_years, max_years):
#                 status = "rejected"
#                 if min_years and max_years:
#                     yr_range = f"{min_years}-{max_years} years"
#                 elif min_years:
#                     yr_range = f"{min_years}+ years"
#                 else:
#                     yr_range = f"up to {max_years} years"
#                 rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
#             else:
#                 status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#                 rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
#         else:
#             status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#             rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)

#     return {
#         "name": name, "email": email, "phone": phone,
#         "raw_text": text[:2000],
#         "skills_found":     json.dumps(skills_found),
#         "tools_found":      json.dumps(tools_found),
#         "projects_found":   json.dumps(projects_found),
#         "internship_found": json.dumps(internship_found),
#         "experience_found": json.dumps(experience_found),
#         "internship_years": internship_years,
#         "experience_years": experience_years,
#         "status": status,
#         "rejection_reason": rejection_reason,
#         **scores
#     }

# # ─────────────────────────────────────────────
# # EMAIL INTEGRATION
# # ─────────────────────────────────────────────

# def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
#     subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
#     body    = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)

#     msg = MIMEMultipart("alternative")
#     msg["Subject"] = subject
#     msg["From"]    = settings["sender_email"]
#     msg["To"]      = candidate_email

#     html_body = f"""
#     <html><body>
#     <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
#       <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
#         <h2 style="color:white; margin:0;">Interview Invitation</h2>
#       </div>
#       <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
#         {body.replace(chr(10), '<br>')}
#       </div>
#     </div>
#     </body></html>
#     """

#     msg.attach(MIMEText(body, "plain"))
#     msg.attach(MIMEText(html_body, "html"))

#     try:
#         server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
#         server.ehlo()
#         server.starttls()
#         server.login(settings["sender_email"], settings["sender_password"])
#         server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
#         server.quit()
#         return True, "Email sent successfully"
#     except Exception as e:
#         return False, str(e)

# # ─────────────────────────────────────────────
# # ROUTES
# # ─────────────────────────────────────────────

# @app.route("/")
# def index():
#     conn = get_db()
#     roles             = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     total_candidates  = conn.execute("SELECT COUNT(*) as c FROM candidates").fetchone()["c"]
#     shortlisted       = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted'").fetchone()["c"]
#     rejected          = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='rejected'").fetchone()["c"]
#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("index.html", roles=roles,
#                            total=total_candidates, shortlisted=shortlisted, rejected=rejected,
#                            popup_notifications=popup_notifications)

# # ── JOB ROLES ──

# @app.route("/roles")
# def roles():
#     conn  = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     conn.close()
#     return render_template("roles.html", roles=roles)

# @app.route("/roles/new", methods=["GET", "POST"])
# def new_role():
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn = get_db()
#         conn.execute("""
#             INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                 internship_keywords, experience_keywords, core_weight, tools_weight,
#                 projects_weight, internship_weight, experience_weight, min_threshold)
#             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50))
#         ))
#         conn.commit(); conn.close()
#         flash("Job role created successfully!", "success")
#         return redirect(url_for("roles"))
#     return render_template("new_role.html")

# @app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
# def edit_role(role_id):
#     conn = get_db()
#     role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn.execute("""
#             UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
#                 project_keywords=?, internship_keywords=?, experience_keywords=?,
#                 core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
#                 experience_weight=?, min_threshold=?
#             WHERE id=?
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50)),
#             role_id
#         ))
#         conn.commit()
#         flash("Job role updated!", "success")
#         return redirect(url_for("roles"))
#     conn.close()
#     return render_template("edit_role.html", role=role)

# @app.route("/roles/<int:role_id>/delete", methods=["POST"])
# def delete_role(role_id):
#     conn = get_db()
#     conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
#     conn.commit(); conn.close()
#     flash("Role deleted.", "info")
#     return redirect(url_for("roles"))

# # ── UPLOAD & SCREEN ──

# @app.route("/upload", methods=["GET", "POST"])
# def upload():
#     conn  = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
#     conn.close()

#     if request.method == "POST":
#         role_id        = int(request.form.get("role_id", 0))
#         files          = request.files.getlist("resumes")
#         candidate_type = request.form.get("candidate_type", "experience")

#         require_internship = request.form.get("require_internship") == "on"
#         require_projects   = request.form.get("require_projects")   == "on"
#         min_years_raw      = request.form.get("min_years", "").strip()
#         max_years_raw      = request.form.get("max_years", "").strip()

#         screening_options = {
#             'require_internship': require_internship,
#             'require_projects':   require_projects,
#             'min_years': float(min_years_raw) if min_years_raw else None,
#             'max_years': float(max_years_raw) if max_years_raw else None,
#         }

#         if not role_id:
#             flash("Please select a job role.", "danger")
#             return redirect(url_for("upload"))
#         if not files or all(f.filename == "" for f in files):
#             flash("Please upload at least one file.", "danger")
#             return redirect(url_for("upload"))

#         conn  = get_db()
#         role  = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()

#         batch_id         = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + str(role_id) + "_" + candidate_type
#         batch_label_parts = [role['title'], candidate_type.title()]
#         if candidate_type == 'experience' and (min_years_raw or max_years_raw):
#             batch_label_parts.append(f"{min_years_raw or '0'}-{max_years_raw or '∞'} yrs")
#         batch_label = " | ".join(batch_label_parts)

#         processed = 0
#         errors    = 0

#         for f in files:
#             if f and allowed_file(f.filename):
#                 filename  = secure_filename(f.filename)
#                 ts        = datetime.now().strftime("%Y%m%d%H%M%S%f")
#                 save_name = f"{ts}_{filename}"
#                 filepath  = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
#                 f.save(filepath)

#                 result = screen_resume(filepath, role, candidate_type, screening_options)
#                 if result:
#                     conn.execute("""
#                         INSERT INTO candidates (job_role_id, candidate_type, batch_id, name, email, phone, raw_text,
#                             skills_found, tools_found, projects_found, internship_found, experience_found,
#                             internship_years, experience_years,
#                             core_score, tools_score, projects_score, internship_score, experience_score,
#                             total_score, status, rejection_reason, filename)
#                         VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
#                     """, (
#                         role_id, candidate_type, batch_id,
#                         result["name"], result["email"], result["phone"], result["raw_text"],
#                         result["skills_found"], result["tools_found"], result["projects_found"],
#                         result["internship_found"], result["experience_found"],
#                         result.get("internship_years", ""), result.get("experience_years", ""),
#                         result["core_score"], result["tools_score"], result["projects_score"],
#                         result["internship_score"], result["experience_score"],
#                         result["total_score"], result["status"], result["rejection_reason"], save_name
#                     ))
#                     processed += 1
#                 else:
#                     errors += 1
#             else:
#                 errors += 1

#         if processed > 0:
#             conn.execute("""
#                 INSERT OR REPLACE INTO upload_batches (id, job_role_id, candidate_type, label, total_resumes)
#                 VALUES (?,?,?,?,?)
#             """, (batch_id, role_id, candidate_type, batch_label, processed))

#         conn.commit(); conn.close()

#         flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
#         if processed > 0:
#             notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Screening Complete!',
#                 'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
#             }]
#         return redirect(url_for("results", batch_id=batch_id))

#     return render_template("upload.html", roles=roles)

# # ── RESULTS ──
# # ✅ FIX: All filters (role, type, status) now work correctly whether a batch
# #         is selected from the sidebar OR "All Results" is chosen.
# #         Previously role_id and type_filter were inside an `else` block under
# #         `if batch_id`, so they were completely ignored when a batch was active.

# @app.route("/results")
# def results():
#     role_id       = request.args.get("role_id", type=int)
#     status_filter = request.args.get("status", "all")
#     batch_id      = request.args.get("batch_id")
#     type_filter   = request.args.get("candidate_type", "all")

#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()

#     batches = conn.execute("""
#         SELECT ub.*, j.title as role_title
#         FROM upload_batches ub
#         JOIN job_roles j ON ub.job_role_id=j.id
#         ORDER BY ub.uploaded_at DESC
#     """).fetchall()

#     query  = "SELECT c.*, j.title as role_title, j.min_threshold FROM candidates c JOIN job_roles j ON c.job_role_id=j.id"
#     params = []
#     where  = []

#     # ── batch narrows the base set ──────────────────────────────────────────
#     if batch_id:
#         where.append("c.batch_id=?")
#         params.append(batch_id)

#     # ── these filters ALWAYS apply on top (batch selected or not) ───────────
#     if role_id:
#         where.append("c.job_role_id=?")
#         params.append(role_id)
#     if type_filter != "all":
#         where.append("c.candidate_type=?")
#         params.append(type_filter)
#     if status_filter != "all":
#         where.append("c.status=?")
#         params.append(status_filter)
#     # ────────────────────────────────────────────────────────────────────────

#     if where:
#         query += " WHERE " + " AND ".join(where)
#     query += " ORDER BY c.total_score DESC"

#     candidates = conn.execute(query, params).fetchall()
#     conn.close()

#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("results.html", candidates=candidates, roles=roles,
#                            selected_role=role_id, status_filter=status_filter,
#                            batches=batches, selected_batch=batch_id,
#                            type_filter=type_filter,
#                            popup_notifications=popup_notifications)

# @app.route("/candidate/<int:cid>")
# def candidate_detail(cid):
#     conn = get_db()
#     c = conn.execute("""
#         SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
#                j.project_keywords, j.internship_keywords, j.experience_keywords
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
#     conn.close()

#     if not c:
#         flash("Candidate not found.", "danger")
#         return redirect(url_for("results"))

#     parsed = dict(c)
#     for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
#                   "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
#         try:
#             parsed[field] = json.loads(c[field] or "[]")
#         except:
#             parsed[field] = []

#     return render_template("candidate_detail.html", c=parsed)

# @app.route("/candidate/<int:cid>/delete", methods=["POST"])
# def delete_candidate(cid):
#     conn = get_db()
#     conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
#     conn.commit(); conn.close()
#     flash("Candidate deleted.", "info")
#     return redirect(url_for("results"))

# # ── EMAIL ──

# @app.route("/send_emails", methods=["POST"])
# def send_emails():
#     candidate_ids = request.form.getlist("candidate_ids")
#     if not candidate_ids:
#         flash("No candidates selected.", "danger")
#         return redirect(url_for("results"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured. Please configure SMTP settings first.", "danger")
#         conn.close()
#         return redirect(url_for("email_settings_page"))

#     sent_count = 0
#     fail_count = 0

#     for cid in candidate_ids:
#         c = conn.execute("""
#             SELECT c.name, c.email, j.title as role_title
#             FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#         """, (cid,)).fetchone()
#         if c and c["email"]:
#             success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#             if success:
#                 conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                              (datetime.now(), cid))
#                 sent_count += 1
#             else:
#                 fail_count += 1

#     conn.commit(); conn.close()

#     if sent_count:
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Emails Sent!',
#             'message': f'Successfully sent {sent_count} interview invitation email(s).'
#         }]
#     if fail_count:
#         existing = session.get('popup_notifications', [])
#         existing.append({'type': 'error', 'title': 'Email Failed',
#                          'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'})
#         session['popup_notifications'] = existing

#     return redirect(url_for("results"))

# @app.route("/send_email_single/<int:cid>", methods=["POST"])
# def send_email_single(cid):
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
#     c        = conn.execute("""
#         SELECT c.name, c.email, j.title as role_title
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured.", "danger")
#     elif not c or not c["email"]:
#         flash("Candidate email not found.", "danger")
#     else:
#         success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#         if success:
#             conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                          (datetime.now(), cid))
#             conn.commit()
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Email Sent!',
#                 'message': f'Interview invitation sent to {c["name"]}.'
#             }]
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Email Failed',
#                 'message': f'Could not send email: {msg}'
#             }]

#     conn.close()
#     return redirect(url_for("candidate_detail", cid=cid))

# @app.route("/email-settings", methods=["GET", "POST"])
# def email_settings_page():
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if request.method == "POST":
#         conn.execute("""
#             UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
#                 sender_password=?, email_subject=?, email_body=? WHERE id=1
#         """, (
#             request.form["smtp_host"], int(request.form["smtp_port"]),
#             request.form["sender_email"], request.form["sender_password"],
#             request.form["email_subject"], request.form["email_body"]
#         ))
#         conn.commit()
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Settings Saved!',
#             'message': 'Email settings have been saved successfully.'
#         }]
#         return redirect(url_for("email_settings_page"))

#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("email_settings.html", settings=settings,
#                            popup_notifications=popup_notifications)

# # ── LOGIN ──

# @app.route("/login", methods=["GET", "POST"])
# def login():
#     if request.method == "POST":
#         username = request.form.get("username", "")
#         password = request.form.get("password", "")
#         if username == "admin" and password == "admin123":
#             session['logged_in'] = True
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Login Successful!',
#                 'message': f'Welcome back, {username}!'
#             }]
#             return redirect(url_for("index"))
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Login Failed',
#                 'message': 'Invalid username or password.'
#             }]
#             return redirect(url_for("login"))
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("login.html", popup_notifications=popup_notifications)

# # ── SEED DEFAULT ROLES ──

# @app.route("/seed-roles")
# def seed_roles():
#     conn     = get_db()
#     existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
#     if existing == 0:
#         for role in DEFAULT_ROLES:
#             conn.execute("""
#                 INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                     internship_keywords, experience_keywords)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (
#                 role["title"], role["description"],
#                 json.dumps(role["core_skills"]), json.dumps(role["tools"]),
#                 json.dumps(role["project_keywords"]), json.dumps(role["internship_keywords"]),
#                 json.dumps(role["experience_keywords"])
#             ))
#         conn.commit()
#         flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
#     else:
#         flash("Job roles already exist.", "info")
#     conn.close()
#     return redirect(url_for("roles"))

# # ── API ──

# @app.route("/api/dashboard")
# def api_dashboard():
#     conn = get_db()
#     rows = conn.execute("""
#         SELECT j.title, COUNT(c.id) as total,
#                SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
#                SUM(CASE WHEN c.status='rejected'    THEN 1 ELSE 0 END) as rejected
#         FROM job_roles j LEFT JOIN candidates c ON j.id=c.job_role_id
#         GROUP BY j.id
#     """).fetchall()
#     conn.close()
#     return jsonify([dict(r) for r in rows])

# # ─────────────────────────────────────────────
# # CUSTOM JINJA2 FILTERS
# # ─────────────────────────────────────────────

# @app.template_filter("fromjson")
# def fromjson_filter(value):
#     try:
#         return json.loads(value)
#     except:
#         return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)


# """
# Intelligent Resume Screening and Automated Interview Notification System
# MCA Final Year Project
# """

# from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
# import sqlite3
# import os
# import json
# import re
# import smtplib
# import io
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from datetime import datetime
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename

# app = Flask(__name__)
# app.secret_key = "resume_screener_mca_2024"

# # Config
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# # ─────────────────────────────────────────────
# # DATABASE SETUP
# # ─────────────────────────────────────────────

# def get_db():
#     conn = sqlite3.connect(DB_PATH)
#     conn.row_factory = sqlite3.Row
#     return conn

# def init_db():
#     conn = get_db()
#     c = conn.cursor()
    
#     c.executescript("""
#     CREATE TABLE IF NOT EXISTS job_roles (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         title TEXT NOT NULL,
#         description TEXT,
#         core_skills TEXT NOT NULL,
#         tools TEXT NOT NULL,
#         project_keywords TEXT NOT NULL,
#         internship_keywords TEXT NOT NULL,
#         experience_keywords TEXT NOT NULL,
#         core_weight REAL DEFAULT 0.40,
#         tools_weight REAL DEFAULT 0.25,
#         projects_weight REAL DEFAULT 0.15,
#         internship_weight REAL DEFAULT 0.10,
#         experience_weight REAL DEFAULT 0.10,
#         min_threshold INTEGER DEFAULT 50,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     CREATE TABLE IF NOT EXISTS candidates (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         job_role_id INTEGER NOT NULL,
#         candidate_type TEXT DEFAULT 'experience',
#         batch_id TEXT,
#         name TEXT,
#         email TEXT,
#         phone TEXT,
#         raw_text TEXT,
#         skills_found TEXT,
#         tools_found TEXT,
#         projects_found TEXT,
#         internship_found TEXT,
#         experience_found TEXT,
#         internship_years TEXT,
#         experience_years TEXT,
#         core_score REAL DEFAULT 0,
#         tools_score REAL DEFAULT 0,
#         projects_score REAL DEFAULT 0,
#         internship_score REAL DEFAULT 0,
#         experience_score REAL DEFAULT 0,
#         total_score REAL DEFAULT 0,
#         status TEXT DEFAULT 'pending',
#         rejection_reason TEXT,
#         email_sent INTEGER DEFAULT 0,
#         email_sent_at TIMESTAMP,
#         filename TEXT,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS upload_batches (
#         id TEXT PRIMARY KEY,
#         job_role_id INTEGER,
#         candidate_type TEXT DEFAULT 'experience',
#         label TEXT,
#         total_resumes INTEGER DEFAULT 0,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS email_settings (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         smtp_host TEXT DEFAULT 'smtp.gmail.com',
#         smtp_port INTEGER DEFAULT 587,
#         sender_email TEXT,
#         sender_password TEXT,
#         email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
#         email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
#     );

# CREATE TABLE IF NOT EXISTS users (
#     id INTEGER PRIMARY KEY AUTOINCREMENT,
#     name TEXT,
#     email TEXT UNIQUE,
#     password TEXT,
#     created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
# );
                    
#     INSERT OR IGNORE INTO email_settings (id) VALUES (1);
#     """)

#     # Add columns if upgrading from old DB
#     for col_sql in [
#         "ALTER TABLE candidates ADD COLUMN batch_id TEXT",
#         "ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'",
#         "ALTER TABLE candidates ADD COLUMN internship_years TEXT",
#         "ALTER TABLE candidates ADD COLUMN experience_years TEXT",
#     ]:
#         try:
#             c.execute(col_sql)
#             conn.commit()
#         except:
#             pass

#     try:
#         c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
#             id TEXT PRIMARY KEY,
#             job_role_id INTEGER,
#             candidate_type TEXT DEFAULT 'experience',
#             label TEXT,
#             total_resumes INTEGER DEFAULT 0,
#             uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#             FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#         )""")
#         conn.commit()
#     except:
#         pass

#     conn.commit()
#     conn.close()

# # ─────────────────────────────────────────────
# # KEYWORD DATABASE FOR ALL IT ROLES
# # ─────────────────────────────────────────────

# DEFAULT_ROLES = [
#     {
#         "title": "Java Full Stack Developer",
#         "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
#         "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
#         "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
#         "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
#         "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
#         "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
#     },
#     {
#         "title": "Python Full Stack Developer",
#         "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
#         "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
#         "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
#         "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
#         "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
#         "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
#     },
#     {
#         "title": "MERN Stack Developer",
#         "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
#         "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
#         "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
#         "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
#         "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
#         "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
#     },
#     {
#         "title": "Data Analyst",
#         "description": "Analyzes data to derive business insights using statistical and visualization tools.",
#         "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
#         "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
#         "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
#         "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
#         "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
#     },
#     {
#         "title": "Data Scientist",
#         "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
#         "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
#         "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
#         "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
#         "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
#     },
#     {
#         "title": "AI/ML Engineer",
#         "description": "Designs, develops, and deploys machine learning models and AI systems.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
#         "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
#         "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
#         "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
#         "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
#     },
#     {
#         "title": "DevOps Engineer",
#         "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
#         "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
#         "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
#         "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
#         "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
#         "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
#     },
#     {
#         "title": "Cloud Engineer",
#         "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
#         "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
#         "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
#         "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
#         "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
#         "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
#     },
#     {
#         "title": "Software Tester",
#         "description": "Ensures software quality through manual and automated testing methodologies.",
#         "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
#         "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
#         "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
#         "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
#         "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
#     },
#     {
#         "title": "Database Developer",
#         "description": "Designs, develops, and optimizes databases for scalability and performance.",
#         "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
#         "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
#         "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
#         "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
#         "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
#     }
# ]

# # ─────────────────────────────────────────────
# # RESUME PARSING
# # ─────────────────────────────────────────────

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(filepath):
#     text = ""
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#     except Exception as e:
#         print(f"PDF error: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ""
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX error: {e}")
#     return text

# def extract_text(filepath):
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_name(text):
#     lines = [l.strip() for l in text.split("\n") if l.strip()]
#     for line in lines[:5]:
#         if re.search(r"[@:/.()]|\d{5,}", line):
#             continue
#         if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
#             continue
#         words = line.split()
#         if 1 < len(words) <= 5:
#             return line.strip()
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
#     return match.group() if match else ""

# def extract_phone(text):
#     match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
#     return match.group().strip() if match else ""

# def extract_internship_years(text):
#     patterns = [
#         r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
#         r'internship[^.]*?(\d+)\s*month',
#         r'(\d+)\s*month[^.]*?internship',
#         r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             val = match.group(1)
#             if 'month' in pattern:
#                 months = float(val)
#                 years = round(months / 12, 1)
#                 return f"{val} months ({years} years)"
#             return f"{val} years"
#     intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
#     if intern_section:
#         try:
#             start = int(intern_section.group(1))
#             end_str = intern_section.group(2)
#             end = datetime.now().year if end_str == 'present' else int(end_str)
#             years = end - start
#             return f"{years} year(s)"
#         except:
#             pass
#     return ""

# def extract_experience_years(text):
#     patterns = [
#         r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
#         r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             return f"{match.group(1)} years"
#     return ""

# def find_keywords_in_text(text, keywords):
#     text_lower = text.lower()
#     found = []
#     for kw in keywords:
#         pattern = r'\b' + re.escape(kw.lower()) + r'\b'
#         if re.search(pattern, text_lower):
#             found.append(kw)
#     return found

# def has_internship_in_resume(text):
#     patterns = [
#         r'\binternship\b', r'\bintern\b', r'\btrainee\b',
#         r'\bindustry training\b', r'\bsummer training\b', r'\bproject trainee\b',
#     ]
#     text_lower = text.lower()
#     for p in patterns:
#         if re.search(p, text_lower):
#             return True
#     return False

# # ─────────────────────────────────────────────
# # SCORING ALGORITHM
# # ─────────────────────────────────────────────

# def calculate_score_fresher(candidate_data, role, screening_options):
#     require_internship = screening_options.get('require_internship', True)
#     core_total   = len(json.loads(role["core_skills"]))
#     tools_total  = len(json.loads(role["tools"]))
#     proj_total   = len(json.loads(role["project_keywords"]))
#     intern_total = len(json.loads(role["internship_keywords"]))

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score  = section_pct(candidate_data["skills_found"],  core_total)
#     tools_score = section_pct(candidate_data["tools_found"],   tools_total)
#     proj_score  = section_pct(candidate_data["projects_found"], proj_total)

#     if require_internship:
#         intern_score = section_pct(candidate_data["internship_found"], intern_total)
#         total = core_score * 0.45 + tools_score * 0.20 + proj_score * 0.20 + intern_score * 0.15
#     else:
#         intern_score = 0
#         total = core_score * 0.50 + tools_score * 0.25 + proj_score * 0.25

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": 0,
#         "total_score": round(total, 1)
#     }

# def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
#     weights = {
#         "core": role["core_weight"],
#         "tools": role["tools_weight"],
#         "projects": role["projects_weight"],
#         "internship": role["internship_weight"],
#         "experience": role["experience_weight"]
#     }

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score   = section_pct(candidate_data["skills_found"],      len(json.loads(role["core_skills"])))
#     tools_score  = section_pct(candidate_data["tools_found"],       len(json.loads(role["tools"])))
#     proj_score   = section_pct(candidate_data["projects_found"],    len(json.loads(role["project_keywords"])))
#     intern_score = section_pct(candidate_data["internship_found"],  len(json.loads(role["internship_keywords"])))
#     exp_score    = section_pct(candidate_data["experience_found"],  len(json.loads(role["experience_keywords"])))

#     total = (
#         core_score   * weights["core"] +
#         tools_score  * weights["tools"] +
#         proj_score   * weights["projects"] +
#         intern_score * weights["internship"] +
#         exp_score    * weights["experience"]
#     )

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": round(exp_score, 1),
#         "total_score": round(total, 1)
#     }

# def calculate_score(candidate_data, role):
#     return calculate_score_experience(candidate_data, role)

# def check_experience_years_range(exp_years_str, min_years, max_years):
#     if not exp_years_str: return False
#     match = re.search(r'(\d+\.?\d*)', exp_years_str)
#     if not match: return False
#     years = float(match.group(1))
#     if min_years is not None and years < min_years: return False
#     if max_years is not None and years > max_years: return False
#     return True

# def generate_rejection_reason(scores, threshold, candidate_type='experience'):
#     reasons = []
#     if scores["core_score"] < 20:
#         reasons.append("insufficient core technical skills")
#     if candidate_type == 'fresher':
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#     else:
#         if scores["tools_score"] < 20:
#             reasons.append("limited relevant tool experience")
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#         if scores["internship_score"] < 10 and scores["experience_score"] < 10:
#             reasons.append("no relevant internship or work experience")
#     if not reasons:
#         reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
#     return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

# def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
#     if screening_options is None:
#         screening_options = {}

#     text = extract_text(filepath)
#     if not text:
#         return None

#     name  = extract_name(text)
#     email = extract_email(text)
#     phone = extract_phone(text)

#     core_skills = json.loads(role["core_skills"])
#     tools       = json.loads(role["tools"])
#     proj_kw     = json.loads(role["project_keywords"])
#     intern_kw   = json.loads(role["internship_keywords"])
#     exp_kw      = json.loads(role["experience_keywords"])

#     skills_found      = find_keywords_in_text(text, core_skills)
#     tools_found       = find_keywords_in_text(text, tools)
#     projects_found    = find_keywords_in_text(text, proj_kw)
#     internship_found  = find_keywords_in_text(text, intern_kw)
#     experience_found  = find_keywords_in_text(text, exp_kw)

#     internship_years  = extract_internship_years(text)
#     experience_years  = extract_experience_years(text)

#     candidate_data = {
#         "skills_found": skills_found,
#         "tools_found": tools_found,
#         "projects_found": projects_found,
#         "internship_found": internship_found,
#         "experience_found": experience_found
#     }

#     threshold = role["min_threshold"]

#     if candidate_type == 'fresher':
#         scores = calculate_score_fresher(candidate_data, role, screening_options)
#         status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#         rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
#     else:
#         min_years = screening_options.get('min_years')
#         max_years = screening_options.get('max_years')
#         scores = calculate_score_experience(candidate_data, role, min_years, max_years)

#         if (min_years is not None or max_years is not None) and experience_years:
#             if not check_experience_years_range(experience_years, min_years, max_years):
#                 status = "rejected"
#                 if min_years and max_years:
#                     yr_range = f"{min_years}-{max_years} years"
#                 elif min_years:
#                     yr_range = f"{min_years}+ years"
#                 else:
#                     yr_range = f"up to {max_years} years"
#                 rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
#             else:
#                 status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#                 rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
#         else:
#             status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#             rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)

#     return {
#         "name": name, "email": email, "phone": phone,
#         "raw_text": text[:2000],
#         "skills_found":     json.dumps(skills_found),
#         "tools_found":      json.dumps(tools_found),
#         "projects_found":   json.dumps(projects_found),
#         "internship_found": json.dumps(internship_found),
#         "experience_found": json.dumps(experience_found),
#         "internship_years": internship_years,
#         "experience_years": experience_years,
#         "status": status,
#         "rejection_reason": rejection_reason,
#         **scores
#     }

# # ─────────────────────────────────────────────
# # EMAIL INTEGRATION
# # ─────────────────────────────────────────────

# def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
#     subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
#     body    = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)

#     msg = MIMEMultipart("alternative")
#     msg["Subject"] = subject
#     msg["From"]    = settings["sender_email"]
#     msg["To"]      = candidate_email

#     html_body = f"""
#     <html><body>
#     <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
#       <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
#         <h2 style="color:white; margin:0;">Interview Invitation</h2>
#       </div>
#       <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
#         {body.replace(chr(10), '<br>')}
#       </div>
#     </div>
#     </body></html>
#     """

#     msg.attach(MIMEText(body, "plain"))
#     msg.attach(MIMEText(html_body, "html"))

#     try:
#         server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
#         server.ehlo()
#         server.starttls()
#         server.login(settings["sender_email"], settings["sender_password"])
#         server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
#         server.quit()
#         return True, "Email sent successfully"
#     except Exception as e:
#         return False, str(e)

# # ─────────────────────────────────────────────
# # ROUTES
# # ─────────────────────────────────────────────

# @app.route("/")
# def index():
#     conn = get_db()
#     roles             = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     total_candidates  = conn.execute("SELECT COUNT(*) as c FROM candidates").fetchone()["c"]
#     shortlisted       = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted'").fetchone()["c"]
#     rejected          = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='rejected'").fetchone()["c"]
#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("index.html", roles=roles,
#                            total=total_candidates, shortlisted=shortlisted, rejected=rejected,
#                            popup_notifications=popup_notifications)

# # ── JOB ROLES ──

# @app.route("/roles")
# def roles():
#     conn  = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     conn.close()
#     return render_template("roles.html", roles=roles)

# @app.route("/roles/new", methods=["GET", "POST"])
# def new_role():
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn = get_db()
#         conn.execute("""
#             INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                 internship_keywords, experience_keywords, core_weight, tools_weight,
#                 projects_weight, internship_weight, experience_weight, min_threshold)
#             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50))
#         ))
#         conn.commit(); conn.close()
#         flash("Job role created successfully!", "success")
#         return redirect(url_for("roles"))
#     return render_template("new_role.html")

# @app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
# def edit_role(role_id):
#     conn = get_db()
#     role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn.execute("""
#             UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
#                 project_keywords=?, internship_keywords=?, experience_keywords=?,
#                 core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
#                 experience_weight=?, min_threshold=?
#             WHERE id=?
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50)),
#             role_id
#         ))
#         conn.commit()
#         flash("Job role updated!", "success")
#         return redirect(url_for("roles"))
#     conn.close()
#     return render_template("edit_role.html", role=role)

# @app.route("/roles/<int:role_id>/delete", methods=["POST"])
# def delete_role(role_id):
#     conn = get_db()
#     conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
#     conn.commit(); conn.close()
#     flash("Role deleted.", "info")
#     return redirect(url_for("roles"))

# # ── UPLOAD & SCREEN ──

# @app.route("/upload", methods=["GET", "POST"])
# def upload():
#     conn  = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
#     conn.close()

#     if request.method == "POST":
#         role_id        = int(request.form.get("role_id", 0))
#         files          = request.files.getlist("resumes")
#         candidate_type = request.form.get("candidate_type", "experience")

#         require_internship = request.form.get("require_internship") == "on"
#         require_projects   = request.form.get("require_projects")   == "on"
#         min_years_raw      = request.form.get("min_years", "").strip()
#         max_years_raw      = request.form.get("max_years", "").strip()

#         screening_options = {
#             'require_internship': require_internship,
#             'require_projects':   require_projects,
#             'min_years': float(min_years_raw) if min_years_raw else None,
#             'max_years': float(max_years_raw) if max_years_raw else None,
#         }

#         if not role_id:
#             flash("Please select a job role.", "danger")
#             return redirect(url_for("upload"))
#         if not files or all(f.filename == "" for f in files):
#             flash("Please upload at least one file.", "danger")
#             return redirect(url_for("upload"))

#         conn  = get_db()
#         role  = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()

#         batch_id         = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + str(role_id) + "_" + candidate_type
#         batch_label_parts = [role['title'], candidate_type.title()]
#         if candidate_type == 'experience' and (min_years_raw or max_years_raw):
#             batch_label_parts.append(f"{min_years_raw or '0'}-{max_years_raw or '∞'} yrs")
#         batch_label = " | ".join(batch_label_parts)

#         processed = 0
#         errors    = 0

#         for f in files:
#             if f and allowed_file(f.filename):
#                 filename  = secure_filename(f.filename)
#                 ts        = datetime.now().strftime("%Y%m%d%H%M%S%f")
#                 save_name = f"{ts}_{filename}"
#                 filepath  = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
#                 f.save(filepath)

#                 result = screen_resume(filepath, role, candidate_type, screening_options)
#                 if result:
#                     conn.execute("""
#                         INSERT INTO candidates (job_role_id, candidate_type, batch_id, name, email, phone, raw_text,
#                             skills_found, tools_found, projects_found, internship_found, experience_found,
#                             internship_years, experience_years,
#                             core_score, tools_score, projects_score, internship_score, experience_score,
#                             total_score, status, rejection_reason, filename, uploaded_at)
#                         VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
#                     """, (
#                         role_id, candidate_type, batch_id,
#                         result["name"], result["email"], result["phone"], result["raw_text"],
#                         result["skills_found"], result["tools_found"], result["projects_found"],
#                         result["internship_found"], result["experience_found"],
#                         result.get("internship_years", ""), result.get("experience_years", ""),
#                         result["core_score"], result["tools_score"], result["projects_score"],
#                         result["internship_score"], result["experience_score"],
#                         result["total_score"], result["status"], result["rejection_reason"], save_name,
#                         datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#                     ))
#                     processed += 1
#                 else:
#                     errors += 1
#             else:
#                 errors += 1

#         if processed > 0:
#             now_local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#             conn.execute("""
#                 INSERT OR REPLACE INTO upload_batches (id, job_role_id, candidate_type, label, total_resumes, uploaded_at)
#                 VALUES (?,?,?,?,?,?)
#             """, (batch_id, role_id, candidate_type, batch_label, processed, now_local))

#         conn.commit(); conn.close()

#         flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
#         if processed > 0:
#             notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Screening Complete!',
#                 'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
#             }]
#         return redirect(url_for("results", batch_id=batch_id))

#     return render_template("upload.html", roles=roles)

# # ── RESULTS ──
# # ✅ FIX: All filters (role, type, status) now work correctly whether a batch
# #         is selected from the sidebar OR "All Results" is chosen.
# #         Previously role_id and type_filter were inside an `else` block under
# #         `if batch_id`, so they were completely ignored when a batch was active.

# @app.route("/results")
# def results():
#     role_id       = request.args.get("role_id", type=int)
#     status_filter = request.args.get("status", "all")
#     batch_id      = request.args.get("batch_id")
#     type_filter   = request.args.get("candidate_type", "all")

#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()

#     batches = conn.execute("""
#         SELECT ub.*, j.title as role_title
#         FROM upload_batches ub
#         JOIN job_roles j ON ub.job_role_id=j.id
#         ORDER BY ub.uploaded_at DESC
#     """).fetchall()

#     query  = "SELECT c.*, j.title as role_title, j.min_threshold FROM candidates c JOIN job_roles j ON c.job_role_id=j.id"
#     params = []
#     where  = []

#     # ── batch narrows the base set ──────────────────────────────────────────
#     if batch_id:
#         where.append("c.batch_id=?")
#         params.append(batch_id)

#     # ── these filters ALWAYS apply on top (batch selected or not) ───────────
#     if role_id:
#         where.append("c.job_role_id=?")
#         params.append(role_id)
#     if type_filter != "all":
#         where.append("c.candidate_type=?")
#         params.append(type_filter)
#     if status_filter != "all":
#         where.append("c.status=?")
#         params.append(status_filter)
#     # ────────────────────────────────────────────────────────────────────────

#     if where:
#         query += " WHERE " + " AND ".join(where)
#     query += " ORDER BY c.total_score DESC"

#     candidates = conn.execute(query, params).fetchall()
#     conn.close()

#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("results.html", candidates=candidates, roles=roles,
#                            selected_role=role_id, status_filter=status_filter,
#                            batches=batches, selected_batch=batch_id,
#                            type_filter=type_filter,
#                            popup_notifications=popup_notifications)

# @app.route("/candidate/<int:cid>")
# def candidate_detail(cid):
#     conn = get_db()
#     c = conn.execute("""
#         SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
#                j.project_keywords, j.internship_keywords, j.experience_keywords
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
#     conn.close()

#     if not c:
#         flash("Candidate not found.", "danger")
#         return redirect(url_for("results"))

#     parsed = dict(c)
#     for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
#                   "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
#         try:
#             parsed[field] = json.loads(c[field] or "[]")
#         except:
#             parsed[field] = []

#     return render_template("candidate_detail.html", c=parsed)

# @app.route("/candidate/<int:cid>/delete", methods=["POST"])
# def delete_candidate(cid):
#     conn = get_db()
#     conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
#     conn.commit(); conn.close()
#     flash("Candidate deleted.", "info")
#     return redirect(url_for("results"))

# # ── EMAIL ──

# @app.route("/send_emails", methods=["POST"])
# def send_emails():
#     candidate_ids = request.form.getlist("candidate_ids")
#     if not candidate_ids:
#         flash("No candidates selected.", "danger")
#         return redirect(url_for("results"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured. Please configure SMTP settings first.", "danger")
#         conn.close()
#         return redirect(url_for("email_settings_page"))

#     sent_count = 0
#     fail_count = 0

#     for cid in candidate_ids:
#         c = conn.execute("""
#             SELECT c.name, c.email, j.title as role_title
#             FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#         """, (cid,)).fetchone()
#         if c and c["email"]:
#             success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#             if success:
#                 conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                              (datetime.now(), cid))
#                 sent_count += 1
#             else:
#                 fail_count += 1

#     conn.commit(); conn.close()

#     if sent_count:
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Emails Sent!',
#             'message': f'Successfully sent {sent_count} interview invitation email(s).'
#         }]
#     if fail_count:
#         existing = session.get('popup_notifications', [])
#         existing.append({'type': 'error', 'title': 'Email Failed',
#                          'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'})
#         session['popup_notifications'] = existing

#     return redirect(url_for("results"))

# @app.route("/send_email_single/<int:cid>", methods=["POST"])
# def send_email_single(cid):
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
#     c        = conn.execute("""
#         SELECT c.name, c.email, j.title as role_title
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured.", "danger")
#     elif not c or not c["email"]:
#         flash("Candidate email not found.", "danger")
#     else:
#         success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#         if success:
#             conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                          (datetime.now(), cid))
#             conn.commit()
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Email Sent!',
#                 'message': f'Interview invitation sent to {c["name"]}.'
#             }]
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Email Failed',
#                 'message': f'Could not send email: {msg}'
#             }]

#     conn.close()
#     return redirect(url_for("candidate_detail", cid=cid))

# @app.route("/email-settings", methods=["GET", "POST"])
# def email_settings_page():
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if request.method == "POST":
#         conn.execute("""
#             UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
#                 sender_password=?, email_subject=?, email_body=? WHERE id=1
#         """, (
#             request.form["smtp_host"], int(request.form["smtp_port"]),
#             request.form["sender_email"], request.form["sender_password"],
#             request.form["email_subject"], request.form["email_body"]
#         ))
#         conn.commit()
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Settings Saved!',
#             'message': 'Email settings have been saved successfully.'
#         }]
#         return redirect(url_for("email_settings_page"))

#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("email_settings.html", settings=settings,
#                            popup_notifications=popup_notifications)

# @app.route("/signup", methods=["GET","POST"])
# def signup():

#     if request.method == "POST":
#         name = request.form["name"]
#         email = request.form["email"]
#         password = request.form["password"]

#         conn = get_db()

#         existing = conn.execute(
#             "SELECT * FROM users WHERE email=?", (email,)
#         ).fetchone()

#         if existing:
#             flash("Email already registered", "danger")
#             return redirect(url_for("signup"))

#         conn.execute(
#             "INSERT INTO users (name,email,password) VALUES (?,?,?)",
#             (name,email,password)
#         )

#         conn.commit()
#         conn.close()

#         flash("Signup successful. Please login.", "success")
#         return redirect(url_for("login"))

#     return render_template("signup.html")

# # ── LOGIN ──

# # @app.route("/login", methods=["GET", "POST"])
# # def login():
# #     if request.method == "POST":
# #         username = request.form.get("username", "")
# #         password = request.form.get("password", "")
# #         if username == "admin" and password == "admin123":
# #             session['logged_in'] = True
# #             session['popup_notifications'] = [{
# #                 'type': 'success', 'title': 'Login Successful!',
# #                 'message': f'Welcome back, {username}!'
# #             }]
# #             return redirect(url_for("index"))
# #         else:
# #             session['popup_notifications'] = [{
# #                 'type': 'error', 'title': 'Login Failed',
# #                 'message': 'Invalid username or password.'
# #             }]
# #             return redirect(url_for("login"))
# #     popup_notifications = session.pop('popup_notifications', [])
# #     return render_template("login.html", popup_notifications=popup_notifications)

# @app.route("/login", methods=["GET","POST"])
# def login():

#     if request.method == "POST":

#         email = request.form["email"]
#         password = request.form["password"]

#         conn = get_db()

#         user = conn.execute(
#             "SELECT * FROM users WHERE email=? AND password=?",
#             (email,password)
#         ).fetchone()

#         conn.close()

#         if user:
#             session["user_id"] = user["id"]
#             session["user_name"] = user["name"]

#             return redirect(url_for("index"))

#         else:
#             flash("Invalid email or password","danger")
#             return redirect(url_for("login"))

#     return render_template("login.html")

# @app.route("/logout")
# def logout():
#     session.clear()
#     return redirect(url_for("login"))

# # ── SEED DEFAULT ROLES ──

# @app.route("/seed-roles")
# def seed_roles():
#     conn     = get_db()
#     existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
#     if existing == 0:
#         for role in DEFAULT_ROLES:
#             conn.execute("""
#                 INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                     internship_keywords, experience_keywords)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (
#                 role["title"], role["description"],
#                 json.dumps(role["core_skills"]), json.dumps(role["tools"]),
#                 json.dumps(role["project_keywords"]), json.dumps(role["internship_keywords"]),
#                 json.dumps(role["experience_keywords"])
#             ))
#         conn.commit()
#         flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
#     else:
#         flash("Job roles already exist.", "info")
#     conn.close()
#     return redirect(url_for("roles"))

# # ── API ──

# @app.route("/api/dashboard")
# def api_dashboard():
#     conn = get_db()
#     rows = conn.execute("""
#         SELECT j.title, COUNT(c.id) as total,
#                SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
#                SUM(CASE WHEN c.status='rejected'    THEN 1 ELSE 0 END) as rejected
#         FROM job_roles j LEFT JOIN candidates c ON j.id=c.job_role_id
#         GROUP BY j.id
#     """).fetchall()
#     conn.close()
#     return jsonify([dict(r) for r in rows])

# # ─────────────────────────────────────────────
# # CUSTOM JINJA2 FILTERS
# # ─────────────────────────────────────────────

# @app.template_filter("fromjson")
# def fromjson_filter(value):
#     try:
#         return json.loads(value)
#     except:
#         return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)





# """
# Intelligent Resume Screening and Automated Interview Notification System
# MCA Final Year Project
# """

# from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
# import sqlite3
# import os
# import json
# import re
# import smtplib
# import io
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from datetime import datetime
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename
# from werkzeug.security import generate_password_hash, check_password_hash

# app = Flask(__name__)
# app.secret_key = "resume_screener_mca_2024"

# # Config
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# # ─────────────────────────────────────────────
# # DATABASE SETUP
# # ─────────────────────────────────────────────

# def get_db():
#     conn = sqlite3.connect(DB_PATH)
#     conn.row_factory = sqlite3.Row
#     return conn

# def init_db():
#     conn = get_db()
#     c = conn.cursor()
    
#     c.executescript("""
#     CREATE TABLE IF NOT EXISTS job_roles (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         title TEXT NOT NULL,
#         description TEXT,
#         core_skills TEXT NOT NULL,
#         tools TEXT NOT NULL,
#         project_keywords TEXT NOT NULL,
#         internship_keywords TEXT NOT NULL,
#         experience_keywords TEXT NOT NULL,
#         core_weight REAL DEFAULT 0.40,
#         tools_weight REAL DEFAULT 0.25,
#         projects_weight REAL DEFAULT 0.15,
#         internship_weight REAL DEFAULT 0.10,
#         experience_weight REAL DEFAULT 0.10,
#         min_threshold INTEGER DEFAULT 50,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     CREATE TABLE IF NOT EXISTS candidates (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         job_role_id INTEGER NOT NULL,
#         candidate_type TEXT DEFAULT 'experience',
#         batch_id TEXT,
#         name TEXT,
#         email TEXT,
#         phone TEXT,
#         raw_text TEXT,
#         skills_found TEXT,
#         tools_found TEXT,
#         projects_found TEXT,
#         internship_found TEXT,
#         experience_found TEXT,
#         internship_years TEXT,
#         experience_years TEXT,
#         core_score REAL DEFAULT 0,
#         tools_score REAL DEFAULT 0,
#         projects_score REAL DEFAULT 0,
#         internship_score REAL DEFAULT 0,
#         experience_score REAL DEFAULT 0,
#         total_score REAL DEFAULT 0,
#         status TEXT DEFAULT 'pending',
#         rejection_reason TEXT,
#         email_sent INTEGER DEFAULT 0,
#         email_sent_at TIMESTAMP,
#         filename TEXT,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS upload_batches (
#         id TEXT PRIMARY KEY,
#         job_role_id INTEGER,
#         candidate_type TEXT DEFAULT 'experience',
#         label TEXT,
#         total_resumes INTEGER DEFAULT 0,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#     );

#     CREATE TABLE IF NOT EXISTS email_settings (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         smtp_host TEXT DEFAULT 'smtp.gmail.com',
#         smtp_port INTEGER DEFAULT 587,
#         sender_email TEXT,
#         sender_password TEXT,
#         email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
#         email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
#     );

#     CREATE TABLE IF NOT EXISTS users (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         name TEXT NOT NULL,
#         email TEXT UNIQUE NOT NULL,
#         password TEXT NOT NULL,
#         login_count INTEGER DEFAULT 0,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     INSERT OR IGNORE INTO email_settings (id) VALUES (1);
#     """)

#     # Add columns if upgrading from old DB
#     for col_sql in [
#         "ALTER TABLE candidates ADD COLUMN batch_id TEXT",
#         "ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'",
#         "ALTER TABLE candidates ADD COLUMN internship_years TEXT",
#         "ALTER TABLE candidates ADD COLUMN experience_years TEXT",
#         "ALTER TABLE users ADD COLUMN login_count INTEGER DEFAULT 0",
#     ]:
#         try:
#             c.execute(col_sql)
#             conn.commit()
#         except:
#             pass

#     try:
#         c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
#             id TEXT PRIMARY KEY,
#             job_role_id INTEGER,
#             candidate_type TEXT DEFAULT 'experience',
#             label TEXT,
#             total_resumes INTEGER DEFAULT 0,
#             uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#             FOREIGN KEY (job_role_id) REFERENCES job_roles(id)
#         )""")
#         conn.commit()
#     except:
#         pass

#     conn.commit()
#     conn.close()

# # ─────────────────────────────────────────────
# # KEYWORD DATABASE FOR ALL IT ROLES
# # ─────────────────────────────────────────────

# DEFAULT_ROLES = [
#     {
#         "title": "Java Full Stack Developer",
#         "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
#         "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
#         "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
#         "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
#         "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
#         "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
#     },
#     {
#         "title": "Python Full Stack Developer",
#         "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
#         "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
#         "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
#         "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
#         "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
#         "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
#     },
#     {
#         "title": "MERN Stack Developer",
#         "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
#         "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
#         "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
#         "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
#         "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
#         "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
#     },
#     {
#         "title": "Data Analyst",
#         "description": "Analyzes data to derive business insights using statistical and visualization tools.",
#         "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
#         "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
#         "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
#         "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
#         "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
#     },
#     {
#         "title": "Data Scientist",
#         "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
#         "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
#         "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
#         "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
#         "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
#     },
#     {
#         "title": "AI/ML Engineer",
#         "description": "Designs, develops, and deploys machine learning models and AI systems.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
#         "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
#         "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
#         "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
#         "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
#     },
#     {
#         "title": "DevOps Engineer",
#         "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
#         "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
#         "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
#         "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
#         "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
#         "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
#     },
#     {
#         "title": "Cloud Engineer",
#         "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
#         "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
#         "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
#         "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
#         "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
#         "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
#     },
#     {
#         "title": "Software Tester",
#         "description": "Ensures software quality through manual and automated testing methodologies.",
#         "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
#         "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
#         "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
#         "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
#         "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
#     },
#     {
#         "title": "Database Developer",
#         "description": "Designs, develops, and optimizes databases for scalability and performance.",
#         "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
#         "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
#         "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
#         "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
#         "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
#     }
# ]

# # ─────────────────────────────────────────────
# # RESUME PARSING
# # ─────────────────────────────────────────────

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(filepath):
#     text = ""
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#     except Exception as e:
#         print(f"PDF error: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ""
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX error: {e}")
#     return text

# def extract_text(filepath):
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_name(text):
#     lines = [l.strip() for l in text.split("\n") if l.strip()]
#     for line in lines[:5]:
#         if re.search(r"[@:/.()]|\d{5,}", line):
#             continue
#         if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
#             continue
#         words = line.split()
#         if 1 < len(words) <= 5:
#             return line.strip()
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
#     return match.group() if match else ""

# def extract_phone(text):
#     match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
#     return match.group().strip() if match else ""

# def extract_internship_years(text):
#     patterns = [
#         r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
#         r'internship[^.]*?(\d+)\s*month',
#         r'(\d+)\s*month[^.]*?internship',
#         r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             val = match.group(1)
#             if 'month' in pattern:
#                 months = float(val)
#                 years = round(months / 12, 1)
#                 return f"{val} months ({years} years)"
#             return f"{val} years"
#     intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
#     if intern_section:
#         try:
#             start = int(intern_section.group(1))
#             end_str = intern_section.group(2)
#             end = datetime.now().year if end_str == 'present' else int(end_str)
#             years = end - start
#             return f"{years} year(s)"
#         except:
#             pass
#     return ""

# def extract_experience_years(text):
#     patterns = [
#         r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
#         r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             return f"{match.group(1)} years"
#     return ""

# def find_keywords_in_text(text, keywords):
#     text_lower = text.lower()
#     found = []
#     for kw in keywords:
#         pattern = r'\b' + re.escape(kw.lower()) + r'\b'
#         if re.search(pattern, text_lower):
#             found.append(kw)
#     return found

# def has_internship_in_resume(text):
#     patterns = [
#         r'\binternship\b', r'\bintern\b', r'\btrainee\b',
#         r'\bindustry training\b', r'\bsummer training\b', r'\bproject trainee\b',
#     ]
#     text_lower = text.lower()
#     for p in patterns:
#         if re.search(p, text_lower):
#             return True
#     return False

# # ─────────────────────────────────────────────
# # SCORING ALGORITHM
# # ─────────────────────────────────────────────

# def calculate_score_fresher(candidate_data, role, screening_options):
#     require_internship = screening_options.get('require_internship', True)
#     core_total   = len(json.loads(role["core_skills"]))
#     tools_total  = len(json.loads(role["tools"]))
#     proj_total   = len(json.loads(role["project_keywords"]))
#     intern_total = len(json.loads(role["internship_keywords"]))

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score  = section_pct(candidate_data["skills_found"],  core_total)
#     tools_score = section_pct(candidate_data["tools_found"],   tools_total)
#     proj_score  = section_pct(candidate_data["projects_found"], proj_total)

#     if require_internship:
#         intern_score = section_pct(candidate_data["internship_found"], intern_total)
#         total = core_score * 0.45 + tools_score * 0.20 + proj_score * 0.20 + intern_score * 0.15
#     else:
#         intern_score = 0
#         total = core_score * 0.50 + tools_score * 0.25 + proj_score * 0.25

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": 0,
#         "total_score": round(total, 1)
#     }

# def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
#     weights = {
#         "core": role["core_weight"],
#         "tools": role["tools_weight"],
#         "projects": role["projects_weight"],
#         "internship": role["internship_weight"],
#         "experience": role["experience_weight"]
#     }

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score   = section_pct(candidate_data["skills_found"],      len(json.loads(role["core_skills"])))
#     tools_score  = section_pct(candidate_data["tools_found"],       len(json.loads(role["tools"])))
#     proj_score   = section_pct(candidate_data["projects_found"],    len(json.loads(role["project_keywords"])))
#     intern_score = section_pct(candidate_data["internship_found"],  len(json.loads(role["internship_keywords"])))
#     exp_score    = section_pct(candidate_data["experience_found"],  len(json.loads(role["experience_keywords"])))

#     total = (
#         core_score   * weights["core"] +
#         tools_score  * weights["tools"] +
#         proj_score   * weights["projects"] +
#         intern_score * weights["internship"] +
#         exp_score    * weights["experience"]
#     )

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": round(exp_score, 1),
#         "total_score": round(total, 1)
#     }

# def calculate_score(candidate_data, role):
#     return calculate_score_experience(candidate_data, role)

# def check_experience_years_range(exp_years_str, min_years, max_years):
#     if not exp_years_str: return False
#     match = re.search(r'(\d+\.?\d*)', exp_years_str)
#     if not match: return False
#     years = float(match.group(1))
#     if min_years is not None and years < min_years: return False
#     if max_years is not None and years > max_years: return False
#     return True

# def generate_rejection_reason(scores, threshold, candidate_type='experience'):
#     reasons = []
#     if scores["core_score"] < 20:
#         reasons.append("insufficient core technical skills")
#     if candidate_type == 'fresher':
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#     else:
#         if scores["tools_score"] < 20:
#             reasons.append("limited relevant tool experience")
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#         if scores["internship_score"] < 10 and scores["experience_score"] < 10:
#             reasons.append("no relevant internship or work experience")
#     if not reasons:
#         reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
#     return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

# def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
#     if screening_options is None:
#         screening_options = {}

#     text = extract_text(filepath)
#     if not text:
#         return None

#     name  = extract_name(text)
#     email = extract_email(text)
#     phone = extract_phone(text)

#     core_skills = json.loads(role["core_skills"])
#     tools       = json.loads(role["tools"])
#     proj_kw     = json.loads(role["project_keywords"])
#     intern_kw   = json.loads(role["internship_keywords"])
#     exp_kw      = json.loads(role["experience_keywords"])

#     skills_found      = find_keywords_in_text(text, core_skills)
#     tools_found       = find_keywords_in_text(text, tools)
#     projects_found    = find_keywords_in_text(text, proj_kw)
#     internship_found  = find_keywords_in_text(text, intern_kw)
#     experience_found  = find_keywords_in_text(text, exp_kw)

#     internship_years  = extract_internship_years(text)
#     experience_years  = extract_experience_years(text)

#     candidate_data = {
#         "skills_found": skills_found,
#         "tools_found": tools_found,
#         "projects_found": projects_found,
#         "internship_found": internship_found,
#         "experience_found": experience_found
#     }

#     threshold = role["min_threshold"]

#     if candidate_type == 'fresher':
#         scores = calculate_score_fresher(candidate_data, role, screening_options)
#         status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#         rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
#     else:
#         min_years = screening_options.get('min_years')
#         max_years = screening_options.get('max_years')
#         scores = calculate_score_experience(candidate_data, role, min_years, max_years)

#         if (min_years is not None or max_years is not None) and experience_years:
#             if not check_experience_years_range(experience_years, min_years, max_years):
#                 status = "rejected"
#                 if min_years and max_years:
#                     yr_range = f"{min_years}-{max_years} years"
#                 elif min_years:
#                     yr_range = f"{min_years}+ years"
#                 else:
#                     yr_range = f"up to {max_years} years"
#                 rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
#             else:
#                 status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#                 rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
#         else:
#             status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#             rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)

#     return {
#         "name": name, "email": email, "phone": phone,
#         "raw_text": text[:2000],
#         "skills_found":     json.dumps(skills_found),
#         "tools_found":      json.dumps(tools_found),
#         "projects_found":   json.dumps(projects_found),
#         "internship_found": json.dumps(internship_found),
#         "experience_found": json.dumps(experience_found),
#         "internship_years": internship_years,
#         "experience_years": experience_years,
#         "status": status,
#         "rejection_reason": rejection_reason,
#         **scores
#     }

# # ─────────────────────────────────────────────
# # EMAIL INTEGRATION
# # ─────────────────────────────────────────────

# def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
#     subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
#     body    = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)

#     msg = MIMEMultipart("alternative")
#     msg["Subject"] = subject
#     msg["From"]    = settings["sender_email"]
#     msg["To"]      = candidate_email

#     html_body = f"""
#     <html><body>
#     <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
#       <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
#         <h2 style="color:white; margin:0;">Interview Invitation</h2>
#       </div>
#       <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
#         {body.replace(chr(10), '<br>')}
#       </div>
#     </div>
#     </body></html>
#     """

#     msg.attach(MIMEText(body, "plain"))
#     msg.attach(MIMEText(html_body, "html"))

#     try:
#         server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
#         server.ehlo()
#         server.starttls()
#         server.login(settings["sender_email"], settings["sender_password"])
#         server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
#         server.quit()
#         return True, "Email sent successfully"
#     except Exception as e:
#         return False, str(e)

# # ─────────────────────────────────────────────
# # ROUTES
# # ─────────────────────────────────────────────

# @app.route("/")
# def index():
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     user_id = session["user_id"]
#     conn = get_db()
#     roles             = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     total_candidates  = conn.execute("SELECT COUNT(*) as c FROM candidates").fetchone()["c"]
#     shortlisted       = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted'").fetchone()["c"]
#     rejected          = conn.execute("SELECT COUNT(*) as c FROM candidates WHERE status='rejected'").fetchone()["c"]
#     conn.close()

#     is_new_user = session.pop("is_new_user", False)
#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("index.html", roles=roles,
#                            total=total_candidates, shortlisted=shortlisted, rejected=rejected,
#                            user_name=session.get("user_name", ""),
#                            is_new_user=is_new_user,
#                            popup_notifications=popup_notifications)

# # ── JOB ROLES ──

# @app.route("/roles")
# def roles():
#     conn  = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     conn.close()
#     return render_template("roles.html", roles=roles)

# @app.route("/roles/new", methods=["GET", "POST"])
# def new_role():
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn = get_db()
#         conn.execute("""
#             INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                 internship_keywords, experience_keywords, core_weight, tools_weight,
#                 projects_weight, internship_weight, experience_weight, min_threshold)
#             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50))
#         ))
#         conn.commit(); conn.close()
#         flash("Job role created successfully!", "success")
#         return redirect(url_for("roles"))
#     return render_template("new_role.html")

# @app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
# def edit_role(role_id):
#     conn = get_db()
#     role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn.execute("""
#             UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
#                 project_keywords=?, internship_keywords=?, experience_keywords=?,
#                 core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
#                 experience_weight=?, min_threshold=?
#             WHERE id=?
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50)),
#             role_id
#         ))
#         conn.commit()
#         flash("Job role updated!", "success")
#         return redirect(url_for("roles"))
#     conn.close()
#     return render_template("edit_role.html", role=role)

# @app.route("/roles/<int:role_id>/delete", methods=["POST"])
# def delete_role(role_id):
#     conn = get_db()
#     conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
#     conn.commit(); conn.close()
#     flash("Role deleted.", "info")
#     return redirect(url_for("roles"))

# # ── UPLOAD & SCREEN ──

# @app.route("/upload", methods=["GET", "POST"])
# def upload():
#     conn  = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
#     conn.close()

#     if request.method == "POST":
#         role_id        = int(request.form.get("role_id", 0))
#         files          = request.files.getlist("resumes")
#         candidate_type = request.form.get("candidate_type", "experience")

#         require_internship = request.form.get("require_internship") == "on"
#         require_projects   = request.form.get("require_projects")   == "on"
#         min_years_raw      = request.form.get("min_years", "").strip()
#         max_years_raw      = request.form.get("max_years", "").strip()

#         screening_options = {
#             'require_internship': require_internship,
#             'require_projects':   require_projects,
#             'min_years': float(min_years_raw) if min_years_raw else None,
#             'max_years': float(max_years_raw) if max_years_raw else None,
#         }

#         if not role_id:
#             flash("Please select a job role.", "danger")
#             return redirect(url_for("upload"))
#         if not files or all(f.filename == "" for f in files):
#             flash("Please upload at least one file.", "danger")
#             return redirect(url_for("upload"))

#         conn  = get_db()
#         role  = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()

#         batch_id         = datetime.now().strftime("%Y%m%d%H%M%S") + "_" + str(role_id) + "_" + candidate_type
#         batch_label_parts = [role['title'], candidate_type.title()]
#         if candidate_type == 'experience' and (min_years_raw or max_years_raw):
#             batch_label_parts.append(f"{min_years_raw or '0'}-{max_years_raw or '∞'} yrs")
#         batch_label = " | ".join(batch_label_parts)

#         processed = 0
#         errors    = 0

#         for f in files:
#             if f and allowed_file(f.filename):
#                 filename  = secure_filename(f.filename)
#                 ts        = datetime.now().strftime("%Y%m%d%H%M%S%f")
#                 save_name = f"{ts}_{filename}"
#                 filepath  = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
#                 f.save(filepath)

#                 result = screen_resume(filepath, role, candidate_type, screening_options)
#                 if result:
#                     conn.execute("""
#                         INSERT INTO candidates (job_role_id, candidate_type, batch_id, name, email, phone, raw_text,
#                             skills_found, tools_found, projects_found, internship_found, experience_found,
#                             internship_years, experience_years,
#                             core_score, tools_score, projects_score, internship_score, experience_score,
#                             total_score, status, rejection_reason, filename, uploaded_at)
#                         VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
#                     """, (
#                         role_id, candidate_type, batch_id,
#                         result["name"], result["email"], result["phone"], result["raw_text"],
#                         result["skills_found"], result["tools_found"], result["projects_found"],
#                         result["internship_found"], result["experience_found"],
#                         result.get("internship_years", ""), result.get("experience_years", ""),
#                         result["core_score"], result["tools_score"], result["projects_score"],
#                         result["internship_score"], result["experience_score"],
#                         result["total_score"], result["status"], result["rejection_reason"], save_name,
#                         datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#                     ))
#                     processed += 1
#                 else:
#                     errors += 1
#             else:
#                 errors += 1

#         if processed > 0:
#             now_local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#             conn.execute("""
#                 INSERT OR REPLACE INTO upload_batches (id, job_role_id, candidate_type, label, total_resumes, uploaded_at)
#                 VALUES (?,?,?,?,?,?)
#             """, (batch_id, role_id, candidate_type, batch_label, processed, now_local))

#         conn.commit(); conn.close()

#         flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
#         if processed > 0:
#             notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Screening Complete!',
#                 'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
#             }]
#         return redirect(url_for("results", batch_id=batch_id))

#     return render_template("upload.html", roles=roles)

# # ── RESULTS ──

# @app.route("/results")
# def results():
#     role_id       = request.args.get("role_id", type=int)
#     status_filter = request.args.get("status", "all")
#     batch_id      = request.args.get("batch_id")
#     type_filter   = request.args.get("candidate_type", "all")

#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()

#     batches = conn.execute("""
#         SELECT ub.*, j.title as role_title
#         FROM upload_batches ub
#         JOIN job_roles j ON ub.job_role_id=j.id
#         ORDER BY ub.uploaded_at DESC
#     """).fetchall()

#     query  = "SELECT c.*, j.title as role_title, j.min_threshold FROM candidates c JOIN job_roles j ON c.job_role_id=j.id"
#     params = []
#     where  = []

#     if batch_id:
#         where.append("c.batch_id=?")
#         params.append(batch_id)

#     if role_id:
#         where.append("c.job_role_id=?")
#         params.append(role_id)
#     if type_filter != "all":
#         where.append("c.candidate_type=?")
#         params.append(type_filter)
#     if status_filter != "all":
#         where.append("c.status=?")
#         params.append(status_filter)

#     if where:
#         query += " WHERE " + " AND ".join(where)
#     query += " ORDER BY c.total_score DESC"

#     candidates = conn.execute(query, params).fetchall()
#     conn.close()

#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("results.html", candidates=candidates, roles=roles,
#                            selected_role=role_id, status_filter=status_filter,
#                            batches=batches, selected_batch=batch_id,
#                            type_filter=type_filter,
#                            popup_notifications=popup_notifications)

# @app.route("/candidate/<int:cid>")
# def candidate_detail(cid):
#     conn = get_db()
#     c = conn.execute("""
#         SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
#                j.project_keywords, j.internship_keywords, j.experience_keywords
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
#     conn.close()

#     if not c:
#         flash("Candidate not found.", "danger")
#         return redirect(url_for("results"))

#     parsed = dict(c)
#     for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
#                   "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
#         try:
#             parsed[field] = json.loads(c[field] or "[]")
#         except:
#             parsed[field] = []

#     return render_template("candidate_detail.html", c=parsed)

# @app.route("/candidate/<int:cid>/delete", methods=["POST"])
# def delete_candidate(cid):
#     conn = get_db()
#     conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
#     conn.commit(); conn.close()
#     flash("Candidate deleted.", "info")
#     return redirect(url_for("results"))

# # ── EMAIL ──

# @app.route("/send_emails", methods=["POST"])
# def send_emails():
#     candidate_ids = request.form.getlist("candidate_ids")
#     if not candidate_ids:
#         flash("No candidates selected.", "danger")
#         return redirect(url_for("results"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured. Please configure SMTP settings first.", "danger")
#         conn.close()
#         return redirect(url_for("email_settings_page"))

#     sent_count = 0
#     fail_count = 0

#     for cid in candidate_ids:
#         c = conn.execute("""
#             SELECT c.name, c.email, j.title as role_title
#             FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#         """, (cid,)).fetchone()
#         if c and c["email"]:
#             success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#             if success:
#                 conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                              (datetime.now(), cid))
#                 sent_count += 1
#             else:
#                 fail_count += 1

#     conn.commit(); conn.close()

#     if sent_count:
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Emails Sent!',
#             'message': f'Successfully sent {sent_count} interview invitation email(s).'
#         }]
#     if fail_count:
#         existing = session.get('popup_notifications', [])
#         existing.append({'type': 'error', 'title': 'Email Failed',
#                          'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'})
#         session['popup_notifications'] = existing

#     return redirect(url_for("results"))

# @app.route("/send_email_single/<int:cid>", methods=["POST"])
# def send_email_single(cid):
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
#     c        = conn.execute("""
#         SELECT c.name, c.email, j.title as role_title
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured.", "danger")
#     elif not c or not c["email"]:
#         flash("Candidate email not found.", "danger")
#     else:
#         success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#         if success:
#             conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                          (datetime.now(), cid))
#             conn.commit()
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Email Sent!',
#                 'message': f'Interview invitation sent to {c["name"]}.'
#             }]
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Email Failed',
#                 'message': f'Could not send email: {msg}'
#             }]

#     conn.close()
#     return redirect(url_for("candidate_detail", cid=cid))

# @app.route("/email-settings", methods=["GET", "POST"])
# def email_settings_page():
#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if request.method == "POST":
#         conn.execute("""
#             UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
#                 sender_password=?, email_subject=?, email_body=? WHERE id=1
#         """, (
#             request.form["smtp_host"], int(request.form["smtp_port"]),
#             request.form["sender_email"], request.form["sender_password"],
#             request.form["email_subject"], request.form["email_body"]
#         ))
#         conn.commit()
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Settings Saved!',
#             'message': 'Email settings have been saved successfully.'
#         }]
#         return redirect(url_for("email_settings_page"))

#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("email_settings.html", settings=settings,
#                            popup_notifications=popup_notifications)

# # ── SIGNUP ──

# @app.route("/signup", methods=["GET", "POST"])
# def signup():
#     if request.method == "POST":
#         name     = request.form["name"]
#         email    = request.form["email"]
#         # Hash the password securely before storing
#         password = generate_password_hash(request.form["password"])

#         conn = get_db()
#         try:
#             conn.execute(
#                 "INSERT INTO users (name, email, password) VALUES (?,?,?)",
#                 (name, email, password)
#             )
#             conn.commit()
#             # flash("Account created successfully! Please login.", "success")
#             return redirect(url_for("login"))
#         except Exception:
#             flash("Email already registered!", "danger")
#         finally:
#             conn.close()

#     return render_template("signup.html")

# # ── LOGIN ──

# @app.route("/login", methods=["GET", "POST"])
# def login():
#     if request.method == "POST":
#         email    = request.form["email"]
#         password = request.form["password"]

#         conn = get_db()
#         user = conn.execute(
#             "SELECT * FROM users WHERE email=?", (email,)
#         ).fetchone()

#         # Use check_password_hash to verify against hashed password
#         if user and check_password_hash(user["password"], password):
#             login_count = user["login_count"] if user["login_count"] else 0
#             is_new_user = (login_count == 0)

#             # Increment login count on every successful login
#             conn.execute(
#                 "UPDATE users SET login_count = login_count + 1 WHERE id=?",
#                 (user["id"],)
#             )
#             conn.commit()
#             conn.close()

#             session["user_id"]   = user["id"]
#             session["user_name"] = user["name"]
#             session["is_new_user"] = is_new_user
#             # flash("Login successful!", "success")
#             return redirect(url_for("index"))
#         else:
#             conn.close()
#             flash("Invalid email or password", "danger")

#     return render_template("login.html")

# # ── LOGOUT ──

# @app.route("/logout")
# def logout():
#     session.clear()
#     # flash("Logged out successfully.", "info")
#     return redirect(url_for("login"))

# # ── SEED DEFAULT ROLES ──

# @app.route("/seed-roles")
# def seed_roles():
#     conn     = get_db()
#     existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
#     if existing == 0:
#         for role in DEFAULT_ROLES:
#             conn.execute("""
#                 INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                     internship_keywords, experience_keywords)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (
#                 role["title"], role["description"],
#                 json.dumps(role["core_skills"]), json.dumps(role["tools"]),
#                 json.dumps(role["project_keywords"]), json.dumps(role["internship_keywords"]),
#                 json.dumps(role["experience_keywords"])
#             ))
#         conn.commit()
#         flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
#     else:
#         flash("Job roles already exist.", "info")
#     conn.close()
#     return redirect(url_for("roles"))

# # ── API ──

# @app.route("/api/dashboard")
# def api_dashboard():
#     conn = get_db()
#     rows = conn.execute("""
#         SELECT j.title, COUNT(c.id) as total,
#                SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
#                SUM(CASE WHEN c.status='rejected'    THEN 1 ELSE 0 END) as rejected
#         FROM job_roles j LEFT JOIN candidates c ON j.id=c.job_role_id
#         GROUP BY j.id
#     """).fetchall()
#     conn.close()
#     return jsonify([dict(r) for r in rows])

# # ─────────────────────────────────────────────
# # CUSTOM JINJA2 FILTERS
# # ─────────────────────────────────────────────

# @app.template_filter("fromjson")
# def fromjson_filter(value):
#     try:
#         return json.loads(value)
#     except:
#         return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)



# """
# Intelligent Resume Screening and Automated Interview Notification System
# MCA Final Year Project
# """

# from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
# import sqlite3
# import os
# import json
# import re
# import smtplib
# import io
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from datetime import datetime
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename
# from werkzeug.security import generate_password_hash, check_password_hash

# app = Flask(__name__)
# app.secret_key = "resume_screener_mca_2024"

# # Config
# BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
# DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
# ALLOWED_EXTENSIONS = {"pdf", "docx"}

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
# app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# # ─────────────────────────────────────────────
# # DATABASE SETUP
# # ─────────────────────────────────────────────

# def get_db():
#     conn = sqlite3.connect(DB_PATH)
#     conn.row_factory = sqlite3.Row
#     return conn

# def init_db():
#     conn = get_db()
#     c = conn.cursor()
    
#     c.executescript("""
#     CREATE TABLE IF NOT EXISTS job_roles (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         title TEXT NOT NULL,
#         description TEXT,
#         core_skills TEXT NOT NULL,
#         tools TEXT NOT NULL,
#         project_keywords TEXT NOT NULL,
#         internship_keywords TEXT NOT NULL,
#         experience_keywords TEXT NOT NULL,
#         core_weight REAL DEFAULT 0.40,
#         tools_weight REAL DEFAULT 0.25,
#         projects_weight REAL DEFAULT 0.15,
#         internship_weight REAL DEFAULT 0.10,
#         experience_weight REAL DEFAULT 0.10,
#         min_threshold INTEGER DEFAULT 50,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     CREATE TABLE IF NOT EXISTS candidates (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         user_id INTEGER,
#         job_role_id INTEGER NOT NULL,
#         candidate_type TEXT DEFAULT 'experience',
#         batch_id TEXT,
#         name TEXT,
#         email TEXT,
#         phone TEXT,
#         raw_text TEXT,
#         skills_found TEXT,
#         tools_found TEXT,
#         projects_found TEXT,
#         internship_found TEXT,
#         experience_found TEXT,
#         internship_years TEXT,
#         experience_years TEXT,
#         core_score REAL DEFAULT 0,
#         tools_score REAL DEFAULT 0,
#         projects_score REAL DEFAULT 0,
#         internship_score REAL DEFAULT 0,
#         experience_score REAL DEFAULT 0,
#         total_score REAL DEFAULT 0,
#         status TEXT DEFAULT 'pending',
#         rejection_reason TEXT,
#         email_sent INTEGER DEFAULT 0,
#         email_sent_at TIMESTAMP,
#         filename TEXT,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id),
#         FOREIGN KEY (user_id) REFERENCES users(id)
#     );

#     CREATE TABLE IF NOT EXISTS upload_batches (
#         id TEXT PRIMARY KEY,
#         user_id INTEGER,
#         job_role_id INTEGER,
#         candidate_type TEXT DEFAULT 'experience',
#         label TEXT,
#         total_resumes INTEGER DEFAULT 0,
#         uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#         FOREIGN KEY (job_role_id) REFERENCES job_roles(id),
#         FOREIGN KEY (user_id) REFERENCES users(id)
#     );

#     CREATE TABLE IF NOT EXISTS email_settings (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         smtp_host TEXT DEFAULT 'smtp.gmail.com',
#         smtp_port INTEGER DEFAULT 587,
#         sender_email TEXT,
#         sender_password TEXT,
#         email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
#         email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
#     );

#     CREATE TABLE IF NOT EXISTS users (
#         id INTEGER PRIMARY KEY AUTOINCREMENT,
#         name TEXT NOT NULL,
#         email TEXT UNIQUE NOT NULL,
#         password TEXT NOT NULL,
#         login_count INTEGER DEFAULT 0,
#         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
#     );

#     INSERT OR IGNORE INTO email_settings (id) VALUES (1);
#     """)

#     # ── Migrations for existing databases ──
#     for col_sql in [
#         "ALTER TABLE candidates ADD COLUMN batch_id TEXT",
#         "ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'",
#         "ALTER TABLE candidates ADD COLUMN internship_years TEXT",
#         "ALTER TABLE candidates ADD COLUMN experience_years TEXT",
#         "ALTER TABLE users ADD COLUMN login_count INTEGER DEFAULT 0",
#         # KEY FIX: user_id columns for data separation
#         "ALTER TABLE candidates ADD COLUMN user_id INTEGER REFERENCES users(id)",
#         "ALTER TABLE upload_batches ADD COLUMN user_id INTEGER REFERENCES users(id)",
#     ]:
#         try:
#             c.execute(col_sql)
#             conn.commit()
#         except:
#             pass

#     try:
#         c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
#             id TEXT PRIMARY KEY,
#             user_id INTEGER,
#             job_role_id INTEGER,
#             candidate_type TEXT DEFAULT 'experience',
#             label TEXT,
#             total_resumes INTEGER DEFAULT 0,
#             uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#             FOREIGN KEY (job_role_id) REFERENCES job_roles(id),
#             FOREIGN KEY (user_id) REFERENCES users(id)
#         )""")
#         conn.commit()
#     except:
#         pass

#     conn.commit()
#     conn.close()

# # ─────────────────────────────────────────────
# # KEYWORD DATABASE FOR ALL IT ROLES
# # ─────────────────────────────────────────────

# DEFAULT_ROLES = [
#     {
#         "title": "Java Full Stack Developer",
#         "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
#         "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
#         "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
#         "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
#         "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
#         "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
#     },
#     {
#         "title": "Python Full Stack Developer",
#         "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
#         "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
#         "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
#         "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
#         "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
#         "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
#     },
#     {
#         "title": "MERN Stack Developer",
#         "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
#         "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
#         "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
#         "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
#         "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
#         "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
#     },
#     {
#         "title": "Data Analyst",
#         "description": "Analyzes data to derive business insights using statistical and visualization tools.",
#         "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
#         "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
#         "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
#         "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
#         "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
#     },
#     {
#         "title": "Data Scientist",
#         "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
#         "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
#         "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
#         "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
#         "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
#     },
#     {
#         "title": "AI/ML Engineer",
#         "description": "Designs, develops, and deploys machine learning models and AI systems.",
#         "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
#         "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
#         "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
#         "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
#         "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
#     },
#     {
#         "title": "DevOps Engineer",
#         "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
#         "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
#         "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
#         "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
#         "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
#         "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
#     },
#     {
#         "title": "Cloud Engineer",
#         "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
#         "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
#         "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
#         "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
#         "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
#         "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
#     },
#     {
#         "title": "Software Tester",
#         "description": "Ensures software quality through manual and automated testing methodologies.",
#         "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
#         "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
#         "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
#         "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
#         "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
#     },
#     {
#         "title": "Database Developer",
#         "description": "Designs, develops, and optimizes databases for scalability and performance.",
#         "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
#         "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
#         "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
#         "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
#         "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
#     }
# ]

# # ─────────────────────────────────────────────
# # RESUME PARSING
# # ─────────────────────────────────────────────

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(filepath):
#     text = ""
#     try:
#         with pdfplumber.open(filepath) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text()
#                 if t:
#                     text += t + "\n"
#     except Exception as e:
#         print(f"PDF error: {e}")
#     return text

# def extract_text_from_docx(filepath):
#     text = ""
#     try:
#         doc = docx.Document(filepath)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     except Exception as e:
#         print(f"DOCX error: {e}")
#     return text

# def extract_text(filepath):
#     ext = filepath.rsplit(".", 1)[1].lower()
#     if ext == "pdf":
#         return extract_text_from_pdf(filepath)
#     elif ext == "docx":
#         return extract_text_from_docx(filepath)
#     return ""

# def extract_name(text):
#     lines = [l.strip() for l in text.split("\n") if l.strip()]
#     for line in lines[:5]:
#         if re.search(r"[@:/.()]|\d{5,}", line):
#             continue
#         if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
#             continue
#         words = line.split()
#         if 1 < len(words) <= 5:
#             return line.strip()
#     return "Unknown"

# def extract_email(text):
#     match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
#     return match.group() if match else ""

# def extract_phone(text):
#     match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
#     return match.group().strip() if match else ""

# def extract_internship_years(text):
#     patterns = [
#         r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
#         r'internship[^.]*?(\d+)\s*month',
#         r'(\d+)\s*month[^.]*?internship',
#         r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             val = match.group(1)
#             if 'month' in pattern:
#                 months = float(val)
#                 years = round(months / 12, 1)
#                 return f"{val} months ({years} years)"
#             return f"{val} years"
#     intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
#     if intern_section:
#         try:
#             start = int(intern_section.group(1))
#             end_str = intern_section.group(2)
#             end = datetime.now().year if end_str == 'present' else int(end_str)
#             years = end - start
#             return f"{years} year(s)"
#         except:
#             pass
#     return ""

# def extract_experience_years(text):
#     patterns = [
#         r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
#         r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
#         r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
#         r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
#     ]
#     text_lower = text.lower()
#     for pattern in patterns:
#         match = re.search(pattern, text_lower)
#         if match:
#             return f"{match.group(1)} years"
#     return ""

# def find_keywords_in_text(text, keywords):
#     text_lower = text.lower()
#     found = []
#     for kw in keywords:
#         pattern = r'\b' + re.escape(kw.lower()) + r'\b'
#         if re.search(pattern, text_lower):
#             found.append(kw)
#     return found

# def has_internship_in_resume(text):
#     patterns = [
#         r'\binternship\b', r'\bintern\b', r'\btrainee\b',
#         r'\bindustry training\b', r'\bsummer training\b', r'\bproject trainee\b',
#     ]
#     text_lower = text.lower()
#     for p in patterns:
#         if re.search(p, text_lower):
#             return True
#     return False

# # ─────────────────────────────────────────────
# # SCORING ALGORITHM
# # ─────────────────────────────────────────────

# def calculate_score_fresher(candidate_data, role, screening_options):
#     require_internship = screening_options.get('require_internship', True)
#     core_total   = len(json.loads(role["core_skills"]))
#     tools_total  = len(json.loads(role["tools"]))
#     proj_total   = len(json.loads(role["project_keywords"]))
#     intern_total = len(json.loads(role["internship_keywords"]))

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score  = section_pct(candidate_data["skills_found"],  core_total)
#     tools_score = section_pct(candidate_data["tools_found"],   tools_total)
#     proj_score  = section_pct(candidate_data["projects_found"], proj_total)

#     if require_internship:
#         intern_score = section_pct(candidate_data["internship_found"], intern_total)
#         total = core_score * 0.45 + tools_score * 0.20 + proj_score * 0.20 + intern_score * 0.15
#     else:
#         intern_score = 0
#         total = core_score * 0.50 + tools_score * 0.25 + proj_score * 0.25

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": 0,
#         "total_score": round(total, 1)
#     }

# def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
#     weights = {
#         "core": role["core_weight"],
#         "tools": role["tools_weight"],
#         "projects": role["projects_weight"],
#         "internship": role["internship_weight"],
#         "experience": role["experience_weight"]
#     }

#     def section_pct(found, total):
#         if not total: return 0
#         return min((len(found) / total) * 100, 100)

#     core_score   = section_pct(candidate_data["skills_found"],      len(json.loads(role["core_skills"])))
#     tools_score  = section_pct(candidate_data["tools_found"],       len(json.loads(role["tools"])))
#     proj_score   = section_pct(candidate_data["projects_found"],    len(json.loads(role["project_keywords"])))
#     intern_score = section_pct(candidate_data["internship_found"],  len(json.loads(role["internship_keywords"])))
#     exp_score    = section_pct(candidate_data["experience_found"],  len(json.loads(role["experience_keywords"])))

#     total = (
#         core_score   * weights["core"] +
#         tools_score  * weights["tools"] +
#         proj_score   * weights["projects"] +
#         intern_score * weights["internship"] +
#         exp_score    * weights["experience"]
#     )

#     return {
#         "core_score": round(core_score, 1),
#         "tools_score": round(tools_score, 1),
#         "projects_score": round(proj_score, 1),
#         "internship_score": round(intern_score, 1),
#         "experience_score": round(exp_score, 1),
#         "total_score": round(total, 1)
#     }

# def calculate_score(candidate_data, role):
#     return calculate_score_experience(candidate_data, role)

# def check_experience_years_range(exp_years_str, min_years, max_years):
#     if not exp_years_str: return False
#     match = re.search(r'(\d+\.?\d*)', exp_years_str)
#     if not match: return False
#     years = float(match.group(1))
#     if min_years is not None and years < min_years: return False
#     if max_years is not None and years > max_years: return False
#     return True

# def generate_rejection_reason(scores, threshold, candidate_type='experience'):
#     reasons = []
#     if scores["core_score"] < 20:
#         reasons.append("insufficient core technical skills")
#     if candidate_type == 'fresher':
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#     else:
#         if scores["tools_score"] < 20:
#             reasons.append("limited relevant tool experience")
#         if scores["projects_score"] < 10:
#             reasons.append("no relevant project exposure")
#         if scores["internship_score"] < 10 and scores["experience_score"] < 10:
#             reasons.append("no relevant internship or work experience")
#     if not reasons:
#         reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
#     return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

# def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
#     if screening_options is None:
#         screening_options = {}

#     text = extract_text(filepath)
#     if not text:
#         return None

#     name  = extract_name(text)
#     email = extract_email(text)
#     phone = extract_phone(text)

#     core_skills = json.loads(role["core_skills"])
#     tools       = json.loads(role["tools"])
#     proj_kw     = json.loads(role["project_keywords"])
#     intern_kw   = json.loads(role["internship_keywords"])
#     exp_kw      = json.loads(role["experience_keywords"])

#     skills_found      = find_keywords_in_text(text, core_skills)
#     tools_found       = find_keywords_in_text(text, tools)
#     projects_found    = find_keywords_in_text(text, proj_kw)
#     internship_found  = find_keywords_in_text(text, intern_kw)
#     experience_found  = find_keywords_in_text(text, exp_kw)

#     internship_years  = extract_internship_years(text)
#     experience_years  = extract_experience_years(text)

#     candidate_data = {
#         "skills_found": skills_found,
#         "tools_found": tools_found,
#         "projects_found": projects_found,
#         "internship_found": internship_found,
#         "experience_found": experience_found
#     }

#     threshold = role["min_threshold"]

#     if candidate_type == 'fresher':
#         scores = calculate_score_fresher(candidate_data, role, screening_options)
#         status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#         rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
#     else:
#         min_years = screening_options.get('min_years')
#         max_years = screening_options.get('max_years')
#         scores = calculate_score_experience(candidate_data, role, min_years, max_years)

#         if (min_years is not None or max_years is not None) and experience_years:
#             if not check_experience_years_range(experience_years, min_years, max_years):
#                 status = "rejected"
#                 if min_years and max_years:
#                     yr_range = f"{min_years}-{max_years} years"
#                 elif min_years:
#                     yr_range = f"{min_years}+ years"
#                 else:
#                     yr_range = f"up to {max_years} years"
#                 rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
#             else:
#                 status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#                 rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
#         else:
#             status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
#             rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)

#     return {
#         "name": name, "email": email, "phone": phone,
#         "raw_text": text[:2000],
#         "skills_found":     json.dumps(skills_found),
#         "tools_found":      json.dumps(tools_found),
#         "projects_found":   json.dumps(projects_found),
#         "internship_found": json.dumps(internship_found),
#         "experience_found": json.dumps(experience_found),
#         "internship_years": internship_years,
#         "experience_years": experience_years,
#         "status": status,
#         "rejection_reason": rejection_reason,
#         **scores
#     }

# # ─────────────────────────────────────────────
# # EMAIL INTEGRATION
# # ─────────────────────────────────────────────

# def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
#     subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
#     body    = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)

#     msg = MIMEMultipart("alternative")
#     msg["Subject"] = subject
#     msg["From"]    = settings["sender_email"]
#     msg["To"]      = candidate_email

#     html_body = f"""
#     <html><body>
#     <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
#       <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
#         <h2 style="color:white; margin:0;">Interview Invitation</h2>
#       </div>
#       <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
#         {body.replace(chr(10), '<br>')}
#       </div>
#     </div>
#     </body></html>
#     """

#     msg.attach(MIMEText(body, "plain"))
#     msg.attach(MIMEText(html_body, "html"))

#     try:
#         server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
#         server.ehlo()
#         server.starttls()
#         server.login(settings["sender_email"], settings["sender_password"])
#         server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
#         server.quit()
#         return True, "Email sent successfully"
#     except Exception as e:
#         return False, str(e)

# # ─────────────────────────────────────────────
# # ROUTES
# # ─────────────────────────────────────────────

# @app.route("/")
# def index():
#     # ── Redirect to login if not logged in ──
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     user_id = session["user_id"]
#     conn = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()

#     # ── Only count THIS user's candidates ──
#     total_candidates = conn.execute(
#         "SELECT COUNT(*) as c FROM candidates WHERE user_id=?", (user_id,)
#     ).fetchone()["c"]
#     shortlisted = conn.execute(
#         "SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted' AND user_id=?", (user_id,)
#     ).fetchone()["c"]
#     rejected = conn.execute(
#         "SELECT COUNT(*) as c FROM candidates WHERE status='rejected' AND user_id=?", (user_id,)
#     ).fetchone()["c"]
#     conn.close()

#     is_new_user = session.pop("is_new_user", False)
#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("index.html", roles=roles,
#                            total=total_candidates, shortlisted=shortlisted, rejected=rejected,
#                            user_name=session.get("user_name", ""),
#                            is_new_user=is_new_user,
#                            popup_notifications=popup_notifications)

# # ── JOB ROLES ──

# @app.route("/roles")
# def roles():
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     conn  = get_db()
#     roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
#     conn.close()
#     return render_template("roles.html", roles=roles)

# @app.route("/roles/new", methods=["GET", "POST"])
# def new_role():
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn = get_db()
#         conn.execute("""
#             INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                 internship_keywords, experience_keywords, core_weight, tools_weight,
#                 projects_weight, internship_weight, experience_weight, min_threshold)
#             VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50))
#         ))
#         conn.commit(); conn.close()
#         flash("Job role created successfully!", "success")
#         return redirect(url_for("roles"))
#     return render_template("new_role.html")

# @app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
# def edit_role(role_id):
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     conn = get_db()
#     role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
#     if request.method == "POST":
#         def parse_list(field):
#             raw   = request.form.get(field, "")
#             items = [x.strip() for x in raw.split(",") if x.strip()]
#             return json.dumps(items)
#         conn.execute("""
#             UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
#                 project_keywords=?, internship_keywords=?, experience_keywords=?,
#                 core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
#                 experience_weight=?, min_threshold=?
#             WHERE id=?
#         """, (
#             request.form["title"], request.form.get("description", ""),
#             parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
#             parse_list("internship_keywords"), parse_list("experience_keywords"),
#             float(request.form.get("core_weight", 0.40)),
#             float(request.form.get("tools_weight", 0.25)),
#             float(request.form.get("projects_weight", 0.15)),
#             float(request.form.get("internship_weight", 0.10)),
#             float(request.form.get("experience_weight", 0.10)),
#             int(request.form.get("min_threshold", 50)),
#             role_id
#         ))
#         conn.commit()
#         flash("Job role updated!", "success")
#         return redirect(url_for("roles"))
#     conn.close()
#     return render_template("edit_role.html", role=role)

# @app.route("/roles/<int:role_id>/delete", methods=["POST"])
# def delete_role(role_id):
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     conn = get_db()
#     conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
#     conn.commit(); conn.close()
#     flash("Role deleted.", "info")
#     return redirect(url_for("roles"))

# # ── UPLOAD & SCREEN ──

# @app.route("/upload", methods=["GET", "POST"])
# def upload():
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     conn  = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
#     conn.close()

#     if request.method == "POST":
#         # ── Get logged-in user ──
#         user_id = session["user_id"]

#         role_id        = int(request.form.get("role_id", 0))
#         files          = request.files.getlist("resumes")
#         candidate_type = request.form.get("candidate_type", "experience")

#         require_internship = request.form.get("require_internship") == "on"
#         require_projects   = request.form.get("require_projects")   == "on"
#         min_years_raw      = request.form.get("min_years", "").strip()
#         max_years_raw      = request.form.get("max_years", "").strip()

#         screening_options = {
#             'require_internship': require_internship,
#             'require_projects':   require_projects,
#             'min_years': float(min_years_raw) if min_years_raw else None,
#             'max_years': float(max_years_raw) if max_years_raw else None,
#         }

#         if not role_id:
#             flash("Please select a job role.", "danger")
#             return redirect(url_for("upload"))
#         if not files or all(f.filename == "" for f in files):
#             flash("Please upload at least one file.", "danger")
#             return redirect(url_for("upload"))

#         conn  = get_db()
#         role  = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()

#         # ── Batch ID includes user_id to keep batches user-specific ──
#         batch_id = (datetime.now().strftime("%Y%m%d%H%M%S")
#                     + "_u" + str(user_id)
#                     + "_" + str(role_id)
#                     + "_" + candidate_type)

#         batch_label_parts = [role['title'], candidate_type.title()]
#         if candidate_type == 'experience' and (min_years_raw or max_years_raw):
#             batch_label_parts.append(f"{min_years_raw or '0'}-{max_years_raw or '∞'} yrs")
#         batch_label = " | ".join(batch_label_parts)

#         processed = 0
#         errors    = 0

#         for f in files:
#             if f and allowed_file(f.filename):
#                 filename  = secure_filename(f.filename)
#                 ts        = datetime.now().strftime("%Y%m%d%H%M%S%f")
#                 save_name = f"{ts}_{filename}"
#                 filepath  = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
#                 f.save(filepath)

#                 result = screen_resume(filepath, role, candidate_type, screening_options)
#                 if result:
#                     # ── Save user_id with every candidate record ──
#                     conn.execute("""
#                         INSERT INTO candidates (
#                             user_id, job_role_id, candidate_type, batch_id,
#                             name, email, phone, raw_text,
#                             skills_found, tools_found, projects_found,
#                             internship_found, experience_found,
#                             internship_years, experience_years,
#                             core_score, tools_score, projects_score,
#                             internship_score, experience_score,
#                             total_score, status, rejection_reason,
#                             filename, uploaded_at)
#                         VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
#                     """, (
#                         user_id, role_id, candidate_type, batch_id,
#                         result["name"], result["email"], result["phone"], result["raw_text"],
#                         result["skills_found"], result["tools_found"], result["projects_found"],
#                         result["internship_found"], result["experience_found"],
#                         result.get("internship_years", ""), result.get("experience_years", ""),
#                         result["core_score"], result["tools_score"], result["projects_score"],
#                         result["internship_score"], result["experience_score"],
#                         result["total_score"], result["status"], result["rejection_reason"],
#                         save_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#                     ))
#                     processed += 1
#                 else:
#                     errors += 1
#             else:
#                 errors += 1

#         if processed > 0:
#             now_local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
#             # ── Save user_id with batch too ──
#             conn.execute("""
#                 INSERT OR REPLACE INTO upload_batches
#                     (id, user_id, job_role_id, candidate_type, label, total_resumes, uploaded_at)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (batch_id, user_id, role_id, candidate_type, batch_label, processed, now_local))

#         conn.commit(); conn.close()

#         flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
#         if processed > 0:
#             notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
#             session['popup_notifications'] = [{
#                 'type': 'success',
#                 'title': 'Screening Complete!',
#                 'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
#             }]
#         return redirect(url_for("results", batch_id=batch_id))

#     return render_template("upload.html", roles=roles)

# # ── RESULTS ──

# @app.route("/results")
# def results():
#     # ── Redirect if not logged in ──
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     user_id = session["user_id"]

#     role_id       = request.args.get("role_id", type=int)
#     status_filter = request.args.get("status", "all")
#     batch_id      = request.args.get("batch_id")
#     type_filter   = request.args.get("candidate_type", "all")

#     conn = get_db()
#     roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()

#     # ── Only show THIS user's batches in sidebar ──
#     batches = conn.execute("""
#         SELECT ub.*, j.title as role_title
#         FROM upload_batches ub
#         JOIN job_roles j ON ub.job_role_id=j.id
#         WHERE ub.user_id=?
#         ORDER BY ub.uploaded_at DESC
#     """, (user_id,)).fetchall()

#     # ── Always filter candidates by current user ──
#     query  = """SELECT c.*, j.title as role_title, j.min_threshold
#                 FROM candidates c
#                 JOIN job_roles j ON c.job_role_id=j.id"""
#     params = []
#     where  = ["c.user_id=?"]     # KEY LINE — never show other users' data
#     params.append(user_id)

#     if batch_id:
#         where.append("c.batch_id=?")
#         params.append(batch_id)
#     if role_id:
#         where.append("c.job_role_id=?")
#         params.append(role_id)
#     if type_filter != "all":
#         where.append("c.candidate_type=?")
#         params.append(type_filter)
#     if status_filter != "all":
#         where.append("c.status=?")
#         params.append(status_filter)

#     query += " WHERE " + " AND ".join(where)
#     query += " ORDER BY c.total_score DESC"

#     candidates = conn.execute(query, params).fetchall()
#     conn.close()

#     popup_notifications = session.pop('popup_notifications', [])

#     return render_template("results.html", candidates=candidates, roles=roles,
#                            selected_role=role_id, status_filter=status_filter,
#                            batches=batches, selected_batch=batch_id,
#                            type_filter=type_filter,
#                            popup_notifications=popup_notifications)

# @app.route("/candidate/<int:cid>")
# def candidate_detail(cid):
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     conn = get_db()
#     c = conn.execute("""
#         SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
#                j.project_keywords, j.internship_keywords, j.experience_keywords
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()
#     conn.close()

#     if not c:
#         flash("Candidate not found.", "danger")
#         return redirect(url_for("results"))

#     parsed = dict(c)
#     for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
#                   "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
#         try:
#             parsed[field] = json.loads(c[field] or "[]")
#         except:
#             parsed[field] = []

#     return render_template("candidate_detail.html", c=parsed)

# @app.route("/candidate/<int:cid>/delete", methods=["POST"])
# def delete_candidate(cid):
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     conn = get_db()
#     conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
#     conn.commit(); conn.close()
#     flash("Candidate deleted.", "info")
#     return redirect(url_for("results"))

# # ── EMAIL ──

# @app.route("/send_emails", methods=["POST"])
# def send_emails():
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     candidate_ids = request.form.getlist("candidate_ids")
#     if not candidate_ids:
#         flash("No candidates selected.", "danger")
#         return redirect(url_for("results"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured. Please configure SMTP settings first.", "danger")
#         conn.close()
#         return redirect(url_for("email_settings_page"))

#     sent_count = 0
#     fail_count = 0

#     for cid in candidate_ids:
#         c = conn.execute("""
#             SELECT c.name, c.email, j.title as role_title
#             FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#         """, (cid,)).fetchone()
#         if c and c["email"]:
#             success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#             if success:
#                 conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                              (datetime.now(), cid))
#                 sent_count += 1
#             else:
#                 fail_count += 1

#     conn.commit(); conn.close()

#     if sent_count:
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Emails Sent!',
#             'message': f'Successfully sent {sent_count} interview invitation email(s).'
#         }]
#     if fail_count:
#         existing = session.get('popup_notifications', [])
#         existing.append({'type': 'error', 'title': 'Email Failed',
#                          'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'})
#         session['popup_notifications'] = existing

#     return redirect(url_for("results"))

# @app.route("/send_email_single/<int:cid>", methods=["POST"])
# def send_email_single(cid):
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
#     c        = conn.execute("""
#         SELECT c.name, c.email, j.title as role_title
#         FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
#     """, (cid,)).fetchone()

#     if not settings or not settings["sender_email"]:
#         flash("Email settings not configured.", "danger")
#     elif not c or not c["email"]:
#         flash("Candidate email not found.", "danger")
#     else:
#         success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
#         if success:
#             conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
#                          (datetime.now(), cid))
#             conn.commit()
#             session['popup_notifications'] = [{
#                 'type': 'success', 'title': 'Email Sent!',
#                 'message': f'Interview invitation sent to {c["name"]}.'
#             }]
#         else:
#             session['popup_notifications'] = [{
#                 'type': 'error', 'title': 'Email Failed',
#                 'message': f'Could not send email: {msg}'
#             }]

#     conn.close()
#     return redirect(url_for("candidate_detail", cid=cid))

# @app.route("/email-settings", methods=["GET", "POST"])
# def email_settings_page():
#     if "user_id" not in session:
#         return redirect(url_for("login"))

#     conn     = get_db()
#     settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

#     if request.method == "POST":
#         conn.execute("""
#             UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
#                 sender_password=?, email_subject=?, email_body=? WHERE id=1
#         """, (
#             request.form["smtp_host"], int(request.form["smtp_port"]),
#             request.form["sender_email"], request.form["sender_password"],
#             request.form["email_subject"], request.form["email_body"]
#         ))
#         conn.commit()
#         session['popup_notifications'] = [{
#             'type': 'success', 'title': 'Settings Saved!',
#             'message': 'Email settings have been saved successfully.'
#         }]
#         return redirect(url_for("email_settings_page"))

#     conn.close()
#     popup_notifications = session.pop('popup_notifications', [])
#     return render_template("email_settings.html", settings=settings,
#                            popup_notifications=popup_notifications)

# # ── SIGNUP ──

# @app.route("/signup", methods=["GET", "POST"])
# def signup():
#     if request.method == "POST":
#         name     = request.form["name"]
#         email    = request.form["email"]
#         password = generate_password_hash(request.form["password"])

#         conn = get_db()
#         try:
#             conn.execute(
#                 "INSERT INTO users (name, email, password) VALUES (?,?,?)",
#                 (name, email, password)
#             )
#             conn.commit()
#             return redirect(url_for("login"))
#         except Exception:
#             flash("Email already registered!", "danger")
#         finally:
#             conn.close()

#     return render_template("signup.html")

# # ── LOGIN ──

# @app.route("/login", methods=["GET", "POST"])
# def login():
#     if request.method == "POST":
#         email    = request.form["email"]
#         password = request.form["password"]

#         conn = get_db()
#         user = conn.execute(
#             "SELECT * FROM users WHERE email=?", (email,)
#         ).fetchone()

#         if user and check_password_hash(user["password"], password):
#             login_count = user["login_count"] if user["login_count"] else 0
#             is_new_user = (login_count == 0)

#             conn.execute(
#                 "UPDATE users SET login_count = login_count + 1 WHERE id=?",
#                 (user["id"],)
#             )
#             conn.commit()
#             conn.close()

#             session["user_id"]     = user["id"]
#             session["user_name"]   = user["name"]
#             session["is_new_user"] = is_new_user
#             return redirect(url_for("index"))
#         else:
#             conn.close()
#             flash("Invalid email or password", "danger")

#     return render_template("login.html")

# # ── LOGOUT ──

# @app.route("/logout")
# def logout():
#     session.clear()
#     return redirect(url_for("login"))

# # ── SEED DEFAULT ROLES ──

# @app.route("/seed-roles")
# def seed_roles():
#     if "user_id" not in session:
#         return redirect(url_for("login"))
#     conn     = get_db()
#     existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
#     if existing == 0:
#         for role in DEFAULT_ROLES:
#             conn.execute("""
#                 INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
#                     internship_keywords, experience_keywords)
#                 VALUES (?,?,?,?,?,?,?)
#             """, (
#                 role["title"], role["description"],
#                 json.dumps(role["core_skills"]), json.dumps(role["tools"]),
#                 json.dumps(role["project_keywords"]), json.dumps(role["internship_keywords"]),
#                 json.dumps(role["experience_keywords"])
#             ))
#         conn.commit()
#         flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
#     else:
#         flash("Job roles already exist.", "info")
#     conn.close()
#     return redirect(url_for("roles"))

# # ── API ──

# @app.route("/api/dashboard")
# def api_dashboard():
#     if "user_id" not in session:
#         return jsonify([])
#     user_id = session["user_id"]
#     conn = get_db()
#     # ── Only count this user's candidates per role ──
#     rows = conn.execute("""
#         SELECT j.title,
#                COUNT(c.id) as total,
#                SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
#                SUM(CASE WHEN c.status='rejected'    THEN 1 ELSE 0 END) as rejected
#         FROM job_roles j
#         LEFT JOIN candidates c ON j.id=c.job_role_id AND c.user_id=?
#         GROUP BY j.id
#     """, (user_id,)).fetchall()
#     conn.close()
#     return jsonify([dict(r) for r in rows])

# # ─────────────────────────────────────────────
# # CUSTOM JINJA2 FILTERS
# # ─────────────────────────────────────────────

# @app.template_filter("fromjson")
# def fromjson_filter(value):
#     try:
#         return json.loads(value)
#     except:
#         return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)


"""
Dynamic  Resume Screening and Interview Automation System
"""

from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
import sqlite3
import os
import json
import re
import smtplib
import io
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
import pdfplumber
import docx
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask(__name__)
app.secret_key = "resume_screener_mca_2024"

# Config
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
DB_PATH = os.path.join(BASE_DIR, "resume_screener.db")
ALLOWED_EXTENSIONS = {"pdf", "docx"}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB

# ─────────────────────────────────────────────
# DATABASE SETUP
# ─────────────────────────────────────────────

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    
    c.executescript("""
    CREATE TABLE IF NOT EXISTS job_roles (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT NOT NULL,
        description TEXT,
        core_skills TEXT NOT NULL,
        tools TEXT NOT NULL,
        project_keywords TEXT NOT NULL,
        internship_keywords TEXT NOT NULL,
        experience_keywords TEXT NOT NULL,
        core_weight REAL DEFAULT 0.40,
        tools_weight REAL DEFAULT 0.25,
        projects_weight REAL DEFAULT 0.15,
        internship_weight REAL DEFAULT 0.10,
        experience_weight REAL DEFAULT 0.10,
        min_threshold INTEGER DEFAULT 50,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );

    CREATE TABLE IF NOT EXISTS candidates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        job_role_id INTEGER NOT NULL,
        candidate_type TEXT DEFAULT 'experience',
        batch_id TEXT,
        name TEXT,
        email TEXT,
        phone TEXT,
        raw_text TEXT,
        skills_found TEXT,
        tools_found TEXT,
        projects_found TEXT,
        internship_found TEXT,
        experience_found TEXT,
        internship_years TEXT,
        experience_years TEXT,
        core_score REAL DEFAULT 0,
        tools_score REAL DEFAULT 0,
        projects_score REAL DEFAULT 0,
        internship_score REAL DEFAULT 0,
        experience_score REAL DEFAULT 0,
        total_score REAL DEFAULT 0,
        status TEXT DEFAULT 'pending',
        rejection_reason TEXT,
        email_sent INTEGER DEFAULT 0,
        email_sent_at TIMESTAMP,
        filename TEXT,
        uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (job_role_id) REFERENCES job_roles(id),
        FOREIGN KEY (user_id) REFERENCES users(id)
    );

    CREATE TABLE IF NOT EXISTS upload_batches (
        id TEXT PRIMARY KEY,
        user_id INTEGER,
        job_role_id INTEGER,
        candidate_type TEXT DEFAULT 'experience',
        label TEXT,
        total_resumes INTEGER DEFAULT 0,
        uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (job_role_id) REFERENCES job_roles(id),
        FOREIGN KEY (user_id) REFERENCES users(id)
    );

    CREATE TABLE IF NOT EXISTS email_settings (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        smtp_host TEXT DEFAULT 'smtp.gmail.com',
        smtp_port INTEGER DEFAULT 587,
        sender_email TEXT,
        sender_password TEXT,
        email_subject TEXT DEFAULT 'Interview Invitation - {job_role}',
        email_body TEXT DEFAULT 'Dear {name},\n\nCongratulations! We are pleased to inform you that your application for the position of {job_role} has been shortlisted.\n\nWe would like to invite you for an interview. Our HR team will contact you shortly with the interview schedule.\n\nBest Regards,\nHR Team'
    );

    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL,
        login_count INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    );

    INSERT OR IGNORE INTO email_settings (id) VALUES (1);
    """)

    # ── Migrations for existing databases ──
    for col_sql in [
        "ALTER TABLE candidates ADD COLUMN batch_id TEXT",
        "ALTER TABLE candidates ADD COLUMN candidate_type TEXT DEFAULT 'experience'",
        "ALTER TABLE candidates ADD COLUMN internship_years TEXT",
        "ALTER TABLE candidates ADD COLUMN experience_years TEXT",
        "ALTER TABLE users ADD COLUMN login_count INTEGER DEFAULT 0",
        # KEY FIX: user_id columns for data separation
        "ALTER TABLE candidates ADD COLUMN user_id INTEGER REFERENCES users(id)",
        "ALTER TABLE upload_batches ADD COLUMN user_id INTEGER REFERENCES users(id)",
    ]:
        try:
            c.execute(col_sql)
            conn.commit()
        except:
            pass

    try:
        c.execute("""CREATE TABLE IF NOT EXISTS upload_batches (
            id TEXT PRIMARY KEY,
            user_id INTEGER,
            job_role_id INTEGER,
            candidate_type TEXT DEFAULT 'experience',
            label TEXT,
            total_resumes INTEGER DEFAULT 0,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (job_role_id) REFERENCES job_roles(id),
            FOREIGN KEY (user_id) REFERENCES users(id)
        )""")
        conn.commit()
    except:
        pass

    conn.commit()
    conn.close()

# ─────────────────────────────────────────────
# KEYWORD DATABASE FOR ALL IT ROLES
# ─────────────────────────────────────────────

DEFAULT_ROLES = [
    {
        "title": "Java Full Stack Developer",
        "description": "Develops end-to-end applications using Java backend and modern frontend frameworks.",
        "core_skills": ["Java", "Spring Boot", "Spring MVC", "Hibernate", "JPA", "REST API", "Microservices", "HTML", "CSS", "JavaScript", "React", "Angular"],
        "tools": ["Maven", "Gradle", "Git", "MySQL", "PostgreSQL", "Docker", "Jenkins", "Postman", "IntelliJ IDEA", "Eclipse", "Tomcat", "Redis"],
        "project_keywords": ["spring boot", "microservice", "rest api", "crud", "ecommerce", "banking", "full stack", "java project", "web application"],
        "internship_keywords": ["java", "spring", "backend", "full stack", "software development", "web development"],
        "experience_keywords": ["java developer", "full stack", "spring boot", "backend developer", "software engineer"]
    },
    {
        "title": "Python Full Stack Developer",
        "description": "Builds web applications using Python backend frameworks and modern frontend technologies.",
        "core_skills": ["Python", "Django", "Flask", "FastAPI", "REST API", "HTML", "CSS", "JavaScript", "React", "Bootstrap", "SQLAlchemy"],
        "tools": ["Git", "PostgreSQL", "MySQL", "Redis", "Docker", "Celery", "Nginx", "PyCharm", "VS Code", "Postman", "Heroku"],
        "project_keywords": ["django", "flask", "python web", "rest api", "fastapi", "ecommerce", "blog", "full stack python", "web app"],
        "internship_keywords": ["python", "django", "flask", "web development", "backend", "full stack"],
        "experience_keywords": ["python developer", "django developer", "flask developer", "full stack", "backend python"]
    },
    {
        "title": "MERN Stack Developer",
        "description": "Develops applications using MongoDB, Express.js, React, and Node.js.",
        "core_skills": ["MongoDB", "Express.js", "React", "Node.js", "JavaScript", "HTML", "CSS", "REST API", "JWT", "Redux"],
        "tools": ["Git", "npm", "Postman", "VS Code", "Heroku", "Netlify", "Firebase", "Mongoose", "Axios", "Webpack"],
        "project_keywords": ["mern", "react", "node.js", "mongodb", "express", "full stack javascript", "spa", "web application"],
        "internship_keywords": ["react", "node", "javascript", "mern", "frontend", "backend", "web development"],
        "experience_keywords": ["mern developer", "react developer", "node.js developer", "full stack javascript"]
    },
    {
        "title": "Data Analyst",
        "description": "Analyzes data to derive business insights using statistical and visualization tools.",
        "core_skills": ["Python", "SQL", "Excel", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Power BI", "Tableau", "Statistics", "Data Visualization"],
        "tools": ["MySQL", "PostgreSQL", "Jupyter Notebook", "Google Sheets", "Power BI", "Tableau", "Excel", "VS Code", "Git"],
        "project_keywords": ["data analysis", "dashboard", "visualization", "eda", "exploratory data analysis", "sales analysis", "business intelligence", "sql queries", "reporting"],
        "internship_keywords": ["data analysis", "sql", "python", "excel", "tableau", "power bi", "analytics"],
        "experience_keywords": ["data analyst", "business analyst", "analytics", "reporting analyst", "sql developer"]
    },
    {
        "title": "Data Scientist",
        "description": "Builds predictive models and extracts insights from large datasets using ML/AI techniques.",
        "core_skills": ["Python", "Machine Learning", "Deep Learning", "Statistics", "SQL", "Pandas", "NumPy", "Scikit-learn", "TensorFlow", "Keras", "NLP"],
        "tools": ["Jupyter Notebook", "Google Colab", "Git", "Power BI", "Tableau", "AWS", "Azure", "Docker", "Spark", "Hadoop"],
        "project_keywords": ["machine learning", "prediction", "classification", "regression", "neural network", "nlp", "deep learning", "model", "dataset", "kaggle"],
        "internship_keywords": ["data science", "machine learning", "python", "ml", "ai", "deep learning", "analytics"],
        "experience_keywords": ["data scientist", "machine learning engineer", "ml engineer", "ai developer", "research scientist"]
    },
    {
        "title": "AI/ML Engineer",
        "description": "Designs, develops, and deploys machine learning models and AI systems.",
        "core_skills": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "Scikit-learn", "NLP", "Computer Vision", "MLOps", "REST API"],
        "tools": ["Jupyter", "Docker", "Kubernetes", "AWS SageMaker", "Azure ML", "MLflow", "Kubeflow", "Git", "DVC", "FastAPI"],
        "project_keywords": ["ai", "machine learning", "deep learning", "model deployment", "nlp", "computer vision", "recommendation system", "chatbot", "generative ai"],
        "internship_keywords": ["machine learning", "ai", "deep learning", "python", "tensorflow", "pytorch", "nlp"],
        "experience_keywords": ["ml engineer", "ai engineer", "machine learning engineer", "deep learning engineer", "research engineer"]
    },
    {
        "title": "DevOps Engineer",
        "description": "Automates software delivery pipelines and manages infrastructure for reliable deployments.",
        "core_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "Jenkins", "Git", "Terraform", "Ansible", "Shell Scripting", "Python", "AWS", "Azure"],
        "tools": ["Jenkins", "GitLab CI", "GitHub Actions", "Docker", "Kubernetes", "Terraform", "Ansible", "Prometheus", "Grafana", "ELK Stack"],
        "project_keywords": ["ci/cd", "pipeline", "docker", "kubernetes", "deployment", "infrastructure", "automation", "devops project", "monitoring"],
        "internship_keywords": ["devops", "linux", "docker", "ci/cd", "cloud", "automation", "jenkins"],
        "experience_keywords": ["devops engineer", "site reliability engineer", "sre", "cloud engineer", "infrastructure engineer"]
    },
    {
        "title": "Cloud Engineer",
        "description": "Designs and manages cloud infrastructure solutions on platforms like AWS, Azure, or GCP.",
        "core_skills": ["AWS", "Azure", "GCP", "Cloud Architecture", "Linux", "Networking", "Security", "Docker", "Kubernetes", "Terraform", "Python"],
        "tools": ["AWS EC2", "S3", "Lambda", "RDS", "Azure VM", "GCP", "Terraform", "Ansible", "CloudFormation", "VPC", "IAM"],
        "project_keywords": ["cloud migration", "aws", "azure", "infrastructure", "serverless", "cloud architecture", "lambda", "s3", "azure functions"],
        "internship_keywords": ["cloud", "aws", "azure", "gcp", "linux", "networking", "cloud services"],
        "experience_keywords": ["cloud engineer", "aws architect", "azure engineer", "cloud architect", "infrastructure engineer"]
    },
    {
        "title": "Software Tester",
        "description": "Ensures software quality through manual and automated testing methodologies.",
        "core_skills": ["Manual Testing", "Automation Testing", "Selenium", "TestNG", "JUnit", "JIRA", "SQL", "API Testing", "Postman", "SDLC", "STLC", "Test Cases"],
        "tools": ["Selenium WebDriver", "Postman", "JIRA", "TestNG", "Maven", "Jenkins", "Git", "Appium", "JMeter", "Cucumber", "BDD"],
        "project_keywords": ["test automation", "selenium", "manual testing", "api testing", "performance testing", "test cases", "bug report", "qa project"],
        "internship_keywords": ["testing", "qa", "quality assurance", "selenium", "manual testing", "automation"],
        "experience_keywords": ["software tester", "qa engineer", "test engineer", "automation tester", "quality analyst"]
    },
    {
        "title": "Database Developer",
        "description": "Designs, develops, and optimizes databases for scalability and performance.",
        "core_skills": ["SQL", "MySQL", "PostgreSQL", "Oracle", "MongoDB", "Database Design", "Normalization", "Stored Procedures", "Indexing", "PL/SQL", "Query Optimization"],
        "tools": ["MySQL Workbench", "pgAdmin", "Oracle SQL Developer", "MongoDB Compass", "Redis", "Cassandra", "Git", "DBeaver", "SSMS"],
        "project_keywords": ["database design", "schema", "sql queries", "normalization", "stored procedure", "data warehouse", "etl", "database project"],
        "internship_keywords": ["sql", "database", "mysql", "postgresql", "mongodb", "pl/sql", "data management"],
        "experience_keywords": ["database developer", "dba", "database administrator", "sql developer", "data engineer"]
    }
]

# ─────────────────────────────────────────────
# RESUME PARSING
# ─────────────────────────────────────────────

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        print(f"PDF error: {e}")
    return text

def extract_text_from_docx(filepath):
    text = ""
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"DOCX error: {e}")
    return text

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(filepath)
    elif ext == "docx":
        return extract_text_from_docx(filepath)
    return ""

def extract_name(text):
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:5]:
        if re.search(r"[@:/.()]|\d{5,}", line):
            continue
        if re.search(r"resume|curriculum|vitae|cv|objective|summary|profile", line, re.IGNORECASE):
            continue
        words = line.split()
        if 1 < len(words) <= 5:
            return line.strip()
    return "Unknown"

def extract_email(text):
    match = re.search(r"[\w.\-+]+@[\w.\-]+\.[a-zA-Z]{2,}", text)
    return match.group() if match else ""

def extract_phone(text):
    match = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
    return match.group().strip() if match else ""

def extract_internship_years(text):
    patterns = [
        r'internship[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
        r'(\d+\.?\d*)\s*(?:year|yr)[^.]*?internship',
        r'internship[^.]*?(\d+)\s*month',
        r'(\d+)\s*month[^.]*?internship',
        r'intern[^.]*?(\d+\.?\d*)\s*(?:year|yr|month)',
    ]
    text_lower = text.lower()
    for pattern in patterns:
        match = re.search(pattern, text_lower)
        if match:
            val = match.group(1)
            if 'month' in pattern:
                months = float(val)
                years = round(months / 12, 1)
                return f"{val} months ({years} years)"
            return f"{val} years"
    intern_section = re.search(r'internship.*?(\d{4})\s*[-–to]+\s*(\d{4}|present)', text_lower)
    if intern_section:
        try:
            start = int(intern_section.group(1))
            end_str = intern_section.group(2)
            end = datetime.now().year if end_str == 'present' else int(end_str)
            years = end - start
            return f"{years} year(s)"
        except:
            pass
    return ""

def extract_experience_years(text):
    patterns = [
        r'(\d+\.?\d*)\+?\s*(?:year|yr)s?\s*(?:of\s+)?experience',
        r'experience[^.]*?(\d+\.?\d*)\s*(?:year|yr)',
        r'worked\s+for\s+(\d+\.?\d*)\s*(?:year|yr)',
        r'(\d+\.?\d*)\s*(?:year|yr)s?\s+(?:of\s+)?(?:work|industry|professional)',
    ]
    text_lower = text.lower()
    for pattern in patterns:
        match = re.search(pattern, text_lower)
        if match:
            return f"{match.group(1)} years"
    return ""

def find_keywords_in_text(text, keywords):
    text_lower = text.lower()
    found = []
    for kw in keywords:
        pattern = r'\b' + re.escape(kw.lower()) + r'\b'
        if re.search(pattern, text_lower):
            found.append(kw)
    return found

def has_internship_in_resume(text):
    patterns = [
        r'\binternship\b', r'\bintern\b', r'\btrainee\b',
        r'\bindustry training\b', r'\bsummer training\b', r'\bproject trainee\b',
    ]
    text_lower = text.lower()
    for p in patterns:
        if re.search(p, text_lower):
            return True
    return False

# ─────────────────────────────────────────────
# SCORING ALGORITHM
# ─────────────────────────────────────────────

def calculate_score_fresher(candidate_data, role, screening_options):
    require_internship = screening_options.get('require_internship', True)
    core_total   = len(json.loads(role["core_skills"]))
    tools_total  = len(json.loads(role["tools"]))
    proj_total   = len(json.loads(role["project_keywords"]))
    intern_total = len(json.loads(role["internship_keywords"]))

    def section_pct(found, total):
        if not total: return 0
        return min((len(found) / total) * 100, 100)

    core_score  = section_pct(candidate_data["skills_found"],  core_total)
    tools_score = section_pct(candidate_data["tools_found"],   tools_total)
    proj_score  = section_pct(candidate_data["projects_found"], proj_total)

    if require_internship:
        intern_score = section_pct(candidate_data["internship_found"], intern_total)
        total = core_score * 0.45 + tools_score * 0.20 + proj_score * 0.20 + intern_score * 0.15
    else:
        intern_score = 0
        total = core_score * 0.50 + tools_score * 0.25 + proj_score * 0.25

    return {
        "core_score": round(core_score, 1),
        "tools_score": round(tools_score, 1),
        "projects_score": round(proj_score, 1),
        "internship_score": round(intern_score, 1),
        "experience_score": 0,
        "total_score": round(total, 1)
    }

def calculate_score_experience(candidate_data, role, min_years=None, max_years=None):
    weights = {
        "core": role["core_weight"],
        "tools": role["tools_weight"],
        "projects": role["projects_weight"],
        "internship": role["internship_weight"],
        "experience": role["experience_weight"]
    }

    def section_pct(found, total):
        if not total: return 0
        return min((len(found) / total) * 100, 100)

    core_score   = section_pct(candidate_data["skills_found"],      len(json.loads(role["core_skills"])))
    tools_score  = section_pct(candidate_data["tools_found"],       len(json.loads(role["tools"])))
    proj_score   = section_pct(candidate_data["projects_found"],    len(json.loads(role["project_keywords"])))
    intern_score = section_pct(candidate_data["internship_found"],  len(json.loads(role["internship_keywords"])))
    exp_score    = section_pct(candidate_data["experience_found"],  len(json.loads(role["experience_keywords"])))

    total = (
        core_score   * weights["core"] +
        tools_score  * weights["tools"] +
        proj_score   * weights["projects"] +
        intern_score * weights["internship"] +
        exp_score    * weights["experience"]
    )

    return {
        "core_score": round(core_score, 1),
        "tools_score": round(tools_score, 1),
        "projects_score": round(proj_score, 1),
        "internship_score": round(intern_score, 1),
        "experience_score": round(exp_score, 1),
        "total_score": round(total, 1)
    }

def calculate_score(candidate_data, role):
    return calculate_score_experience(candidate_data, role)

def check_experience_years_range(exp_years_str, min_years, max_years):
    if not exp_years_str: return False
    match = re.search(r'(\d+\.?\d*)', exp_years_str)
    if not match: return False
    years = float(match.group(1))
    if min_years is not None and years < min_years: return False
    if max_years is not None and years > max_years: return False
    return True

def generate_rejection_reason(scores, threshold, candidate_type='experience'):
    reasons = []
    if scores["core_score"] < 20:
        reasons.append("insufficient core technical skills")
    if candidate_type == 'fresher':
        if scores["projects_score"] < 10:
            reasons.append("no relevant project exposure")
    else:
        if scores["tools_score"] < 20:
            reasons.append("limited relevant tool experience")
        if scores["projects_score"] < 10:
            reasons.append("no relevant project exposure")
        if scores["internship_score"] < 10 and scores["experience_score"] < 10:
            reasons.append("no relevant internship or work experience")
    if not reasons:
        reasons.append(f"overall profile score ({scores['total_score']:.1f}%) below minimum threshold ({threshold}%)")
    return f"Profile score {scores['total_score']:.1f}% is below threshold {threshold}%. Reasons: {'; '.join(reasons)}."

def screen_resume(filepath, role, candidate_type='experience', screening_options=None):
    if screening_options is None:
        screening_options = {}

    text = extract_text(filepath)
    if not text:
        return None

    name  = extract_name(text)
    email = extract_email(text)
    phone = extract_phone(text)

    core_skills = json.loads(role["core_skills"])
    tools       = json.loads(role["tools"])
    proj_kw     = json.loads(role["project_keywords"])
    intern_kw   = json.loads(role["internship_keywords"])
    exp_kw      = json.loads(role["experience_keywords"])

    skills_found      = find_keywords_in_text(text, core_skills)
    tools_found       = find_keywords_in_text(text, tools)
    projects_found    = find_keywords_in_text(text, proj_kw)
    internship_found  = find_keywords_in_text(text, intern_kw)
    experience_found  = find_keywords_in_text(text, exp_kw)

    internship_years  = extract_internship_years(text)
    experience_years  = extract_experience_years(text)

    candidate_data = {
        "skills_found": skills_found,
        "tools_found": tools_found,
        "projects_found": projects_found,
        "internship_found": internship_found,
        "experience_found": experience_found
    }

    threshold = role["min_threshold"]

    if candidate_type == 'fresher':
        scores = calculate_score_fresher(candidate_data, role, screening_options)
        status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
        rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold, 'fresher')
    else:
        min_years = screening_options.get('min_years')
        max_years = screening_options.get('max_years')
        scores = calculate_score_experience(candidate_data, role, min_years, max_years)

        if (min_years is not None or max_years is not None) and experience_years:
            if not check_experience_years_range(experience_years, min_years, max_years):
                status = "rejected"
                if min_years and max_years:
                    yr_range = f"{min_years}-{max_years} years"
                elif min_years:
                    yr_range = f"{min_years}+ years"
                else:
                    yr_range = f"up to {max_years} years"
                rejection_reason = f"Experience ({experience_years}) does not match required range ({yr_range})."
            else:
                status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
                rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)
        else:
            status = "shortlisted" if scores["total_score"] >= threshold else "rejected"
            rejection_reason = "" if status == "shortlisted" else generate_rejection_reason(scores, threshold)

    return {
        "name": name, "email": email, "phone": phone,
        "raw_text": text[:2000],
        "skills_found":     json.dumps(skills_found),
        "tools_found":      json.dumps(tools_found),
        "projects_found":   json.dumps(projects_found),
        "internship_found": json.dumps(internship_found),
        "experience_found": json.dumps(experience_found),
        "internship_years": internship_years,
        "experience_years": experience_years,
        "status": status,
        "rejection_reason": rejection_reason,
        **scores
    }

# ─────────────────────────────────────────────
# EMAIL INTEGRATION
# ─────────────────────────────────────────────

def send_interview_email(settings, candidate_name, candidate_email, job_role_title):
    subject = settings["email_subject"].replace("{job_role}", job_role_title).replace("{name}", candidate_name)
    body    = settings["email_body"].replace("{name}", candidate_name).replace("{job_role}", job_role_title)

    msg = MIMEMultipart("alternative")
    msg["Subject"] = subject
    msg["From"]    = settings["sender_email"]
    msg["To"]      = candidate_email

    html_body = f"""
    <html><body>
    <div style="font-family:Arial,sans-serif; max-width:600px; margin:0 auto; padding:20px;">
      <div style="background:#2563eb; padding:20px; border-radius:8px 8px 0 0;">
        <h2 style="color:white; margin:0;">Interview Invitation</h2>
      </div>
      <div style="background:#f8fafc; padding:30px; border:1px solid #e2e8f0; border-radius:0 0 8px 8px;">
        {body.replace(chr(10), '<br>')}
      </div>
    </div>
    </body></html>
    """

    msg.attach(MIMEText(body, "plain"))
    msg.attach(MIMEText(html_body, "html"))

    try:
        server = smtplib.SMTP(settings["smtp_host"], settings["smtp_port"])
        server.ehlo()
        server.starttls()
        server.login(settings["sender_email"], settings["sender_password"])
        server.sendmail(settings["sender_email"], candidate_email, msg.as_string())
        server.quit()
        return True, "Email sent successfully"
    except Exception as e:
        return False, str(e)

# ─────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────

@app.route("/")
def index():
    # ── Redirect to login if not logged in ──
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]
    conn = get_db()
    roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()

    # ── Only count THIS user's candidates ──
    total_candidates = conn.execute(
        "SELECT COUNT(*) as c FROM candidates WHERE user_id=?", (user_id,)
    ).fetchone()["c"]
    shortlisted = conn.execute(
        "SELECT COUNT(*) as c FROM candidates WHERE status='shortlisted' AND user_id=?", (user_id,)
    ).fetchone()["c"]
    rejected = conn.execute(
        "SELECT COUNT(*) as c FROM candidates WHERE status='rejected' AND user_id=?", (user_id,)
    ).fetchone()["c"]
    conn.close()

    is_new_user = session.pop("is_new_user", False)
    popup_notifications = session.pop('popup_notifications', [])

    return render_template("index.html", roles=roles,
                           total=total_candidates, shortlisted=shortlisted, rejected=rejected,
                           user_name=session.get("user_name", ""),
                           is_new_user=is_new_user,
                           popup_notifications=popup_notifications)

# ── JOB ROLES ──

@app.route("/roles")
def roles():
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn  = get_db()
    roles = conn.execute("SELECT * FROM job_roles ORDER BY created_at DESC").fetchall()
    conn.close()
    return render_template("roles.html", roles=roles)

@app.route("/roles/new", methods=["GET", "POST"])
def new_role():
    if "user_id" not in session:
        return redirect(url_for("login"))
    if request.method == "POST":
        def parse_list(field):
            raw   = request.form.get(field, "")
            items = [x.strip() for x in raw.split(",") if x.strip()]
            return json.dumps(items)
        conn = get_db()
        conn.execute("""
            INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
                internship_keywords, experience_keywords, core_weight, tools_weight,
                projects_weight, internship_weight, experience_weight, min_threshold)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            request.form["title"], request.form.get("description", ""),
            parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
            parse_list("internship_keywords"), parse_list("experience_keywords"),
            float(request.form.get("core_weight", 0.40)),
            float(request.form.get("tools_weight", 0.25)),
            float(request.form.get("projects_weight", 0.15)),
            float(request.form.get("internship_weight", 0.10)),
            float(request.form.get("experience_weight", 0.10)),
            int(request.form.get("min_threshold", 50))
        ))
        conn.commit(); conn.close()
        flash("Job role created successfully!", "success")
        return redirect(url_for("roles"))
    return render_template("new_role.html")

@app.route("/roles/<int:role_id>/edit", methods=["GET", "POST"])
def edit_role(role_id):
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn = get_db()
    role = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()
    if request.method == "POST":
        def parse_list(field):
            raw   = request.form.get(field, "")
            items = [x.strip() for x in raw.split(",") if x.strip()]
            return json.dumps(items)
        conn.execute("""
            UPDATE job_roles SET title=?, description=?, core_skills=?, tools=?,
                project_keywords=?, internship_keywords=?, experience_keywords=?,
                core_weight=?, tools_weight=?, projects_weight=?, internship_weight=?,
                experience_weight=?, min_threshold=?
            WHERE id=?
        """, (
            request.form["title"], request.form.get("description", ""),
            parse_list("core_skills"), parse_list("tools"), parse_list("project_keywords"),
            parse_list("internship_keywords"), parse_list("experience_keywords"),
            float(request.form.get("core_weight", 0.40)),
            float(request.form.get("tools_weight", 0.25)),
            float(request.form.get("projects_weight", 0.15)),
            float(request.form.get("internship_weight", 0.10)),
            float(request.form.get("experience_weight", 0.10)),
            int(request.form.get("min_threshold", 50)),
            role_id
        ))
        conn.commit()
        flash("Job role updated!", "success")
        return redirect(url_for("roles"))
    conn.close()
    return render_template("edit_role.html", role=role)

@app.route("/roles/<int:role_id>/delete", methods=["POST"])
def delete_role(role_id):
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn = get_db()
    conn.execute("DELETE FROM job_roles WHERE id=?", (role_id,))
    conn.commit(); conn.close()
    flash("Role deleted.", "info")
    return redirect(url_for("roles"))

# ── UPLOAD & SCREEN ──

@app.route("/upload", methods=["GET", "POST"])
def upload():
    if "user_id" not in session:
        return redirect(url_for("login"))

    conn  = get_db()
    roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()
    conn.close()

    if request.method == "POST":
        # ── Get logged-in user ──
        user_id = session["user_id"]

        role_id        = int(request.form.get("role_id", 0))
        files          = request.files.getlist("resumes")
        candidate_type = request.form.get("candidate_type", "experience")

        require_internship = request.form.get("require_internship") == "on"
        require_projects   = request.form.get("require_projects")   == "on"
        min_years_raw      = request.form.get("min_years", "").strip()
        max_years_raw      = request.form.get("max_years", "").strip()

        screening_options = {
            'require_internship': require_internship,
            'require_projects':   require_projects,
            'min_years': float(min_years_raw) if min_years_raw else None,
            'max_years': float(max_years_raw) if max_years_raw else None,
        }

        if not role_id:
            flash("Please select a job role.", "danger")
            return redirect(url_for("upload"))
        if not files or all(f.filename == "" for f in files):
            flash("Please upload at least one file.", "danger")
            return redirect(url_for("upload"))

        conn  = get_db()
        role  = conn.execute("SELECT * FROM job_roles WHERE id=?", (role_id,)).fetchone()

        # ── Batch ID includes user_id to keep batches user-specific ──
        batch_id = (datetime.now().strftime("%Y%m%d%H%M%S")
                    + "_u" + str(user_id)
                    + "_" + str(role_id)
                    + "_" + candidate_type)

        batch_label_parts = [role['title'], candidate_type.title()]
        if candidate_type == 'experience' and (min_years_raw or max_years_raw):
            batch_label_parts.append(f"{min_years_raw or '0'}-{max_years_raw or '∞'} yrs")
        batch_label = " | ".join(batch_label_parts)

        processed = 0
        errors    = 0

        for f in files:
            if f and allowed_file(f.filename):
                filename  = secure_filename(f.filename)
                ts        = datetime.now().strftime("%Y%m%d%H%M%S%f")
                save_name = f"{ts}_{filename}"
                filepath  = os.path.join(app.config["UPLOAD_FOLDER"], save_name)
                f.save(filepath)

                result = screen_resume(filepath, role, candidate_type, screening_options)
                if result:
                    # ── Save user_id with every candidate record ──
                    conn.execute("""
                        INSERT INTO candidates (
                            user_id, job_role_id, candidate_type, batch_id,
                            name, email, phone, raw_text,
                            skills_found, tools_found, projects_found,
                            internship_found, experience_found,
                            internship_years, experience_years,
                            core_score, tools_score, projects_score,
                            internship_score, experience_score,
                            total_score, status, rejection_reason,
                            filename, uploaded_at)
                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                    """, (
                        user_id, role_id, candidate_type, batch_id,
                        result["name"], result["email"], result["phone"], result["raw_text"],
                        result["skills_found"], result["tools_found"], result["projects_found"],
                        result["internship_found"], result["experience_found"],
                        result.get("internship_years", ""), result.get("experience_years", ""),
                        result["core_score"], result["tools_score"], result["projects_score"],
                        result["internship_score"], result["experience_score"],
                        result["total_score"], result["status"], result["rejection_reason"],
                        save_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    ))
                    processed += 1
                else:
                    errors += 1
            else:
                errors += 1

        if processed > 0:
            now_local = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            # ── Save user_id with batch too ──
            conn.execute("""
                INSERT OR REPLACE INTO upload_batches
                    (id, user_id, job_role_id, candidate_type, label, total_resumes, uploaded_at)
                VALUES (?,?,?,?,?,?,?)
            """, (batch_id, user_id, role_id, candidate_type, batch_label, processed, now_local))

        conn.commit(); conn.close()

        flash(f"Processed {processed} resumes. {errors} errors.", "success" if processed else "danger")
        if processed > 0:
            notif_type = "Fresher" if candidate_type == "fresher" else "Experience"
            session['popup_notifications'] = [{
                'type': 'success',
                'title': 'Screening Complete!',
                'message': f'{processed} {notif_type} resume(s) screened successfully for {role["title"]}.'
            }]
        return redirect(url_for("results", batch_id=batch_id))

    return render_template("upload.html", roles=roles)

# ── RESULTS ──

@app.route("/results")
def results():
    # ── Redirect if not logged in ──
    if "user_id" not in session:
        return redirect(url_for("login"))

    user_id = session["user_id"]

    role_id       = request.args.get("role_id", type=int)
    status_filter = request.args.get("status", "all")
    batch_id      = request.args.get("batch_id")
    type_filter   = request.args.get("candidate_type", "all")

    conn = get_db()
    roles = conn.execute("SELECT id, title FROM job_roles ORDER BY title").fetchall()

    # ── Only show THIS user's batches in sidebar ──
    batches = conn.execute("""
        SELECT ub.*, j.title as role_title
        FROM upload_batches ub
        JOIN job_roles j ON ub.job_role_id=j.id
        WHERE ub.user_id=?
        ORDER BY ub.uploaded_at DESC
    """, (user_id,)).fetchall()

    # ── Always filter candidates by current user ──
    query  = """SELECT c.*, j.title as role_title, j.min_threshold
                FROM candidates c
                JOIN job_roles j ON c.job_role_id=j.id"""
    params = []
    where  = ["c.user_id=?"]     # KEY LINE — never show other users' data
    params.append(user_id)

    if batch_id:
        where.append("c.batch_id=?")
        params.append(batch_id)
    if role_id:
        where.append("c.job_role_id=?")
        params.append(role_id)
    if type_filter != "all":
        where.append("c.candidate_type=?")
        params.append(type_filter)
    if status_filter != "all":
        where.append("c.status=?")
        params.append(status_filter)

    query += " WHERE " + " AND ".join(where)
    query += " ORDER BY c.total_score DESC"

    candidates = conn.execute(query, params).fetchall()
    conn.close()

    popup_notifications = session.pop('popup_notifications', [])

    return render_template("results.html", candidates=candidates, roles=roles,
                           selected_role=role_id, status_filter=status_filter,
                           batches=batches, selected_batch=batch_id,
                           type_filter=type_filter,
                           popup_notifications=popup_notifications)

@app.route("/candidate/<int:cid>")
def candidate_detail(cid):
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn = get_db()
    c = conn.execute("""
        SELECT c.*, j.title as role_title, j.min_threshold, j.core_skills, j.tools,
               j.project_keywords, j.internship_keywords, j.experience_keywords
        FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
    """, (cid,)).fetchone()
    conn.close()

    if not c:
        flash("Candidate not found.", "danger")
        return redirect(url_for("results"))

    parsed = dict(c)
    for field in ["skills_found", "tools_found", "projects_found", "internship_found", "experience_found",
                  "core_skills", "tools", "project_keywords", "internship_keywords", "experience_keywords"]:
        try:
            parsed[field] = json.loads(c[field] or "[]")
        except:
            parsed[field] = []

    return render_template("candidate_detail.html", c=parsed)

@app.route("/candidate/<int:cid>/delete", methods=["POST"])
def delete_candidate(cid):
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn = get_db()
    conn.execute("DELETE FROM candidates WHERE id=?", (cid,))
    conn.commit(); conn.close()
    flash("Candidate deleted.", "info")
    return redirect(url_for("results"))

# ── VIEW RESUME (PDF/DOCX) ──

@app.route("/candidate/<int:cid>/resume")
def view_resume(cid):
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn = get_db()
    c = conn.execute("SELECT filename FROM candidates WHERE id=?", (cid,)).fetchone()
    conn.close()
    if not c or not c["filename"]:
        flash("Resume file not found.", "danger")
        return redirect(url_for("results"))
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], c["filename"])
    if not os.path.exists(filepath):
        flash("Resume file no longer exists on server.", "danger")
        return redirect(url_for("results"))
    ext = c["filename"].rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        # Serve PDF inline so browser opens it directly
        return send_file(filepath, mimetype="application/pdf", as_attachment=False)
    elif ext == "docx":
        # DOCX: download it (browsers can't display docx inline)
        return send_file(filepath,
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         as_attachment=True,
                         download_name=c["filename"])
    else:
        flash("Unsupported file type.", "danger")
        return redirect(url_for("results"))

# ── EMAIL ──

@app.route("/send_emails", methods=["POST"])
def send_emails():
    if "user_id" not in session:
        return redirect(url_for("login"))

    candidate_ids = request.form.getlist("candidate_ids")
    if not candidate_ids:
        flash("No candidates selected.", "danger")
        return redirect(url_for("results"))

    conn     = get_db()
    settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

    if not settings or not settings["sender_email"]:
        flash("Email settings not configured. Please configure SMTP settings first.", "danger")
        conn.close()
        return redirect(url_for("email_settings_page"))

    sent_count = 0
    fail_count = 0

    for cid in candidate_ids:
        c = conn.execute("""
            SELECT c.name, c.email, j.title as role_title
            FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
        """, (cid,)).fetchone()
        if c and c["email"]:
            success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
            if success:
                conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
                             (datetime.now(), cid))
                sent_count += 1
            else:
                fail_count += 1

    conn.commit(); conn.close()

    if sent_count:
        session['popup_notifications'] = [{
            'type': 'success', 'title': 'Emails Sent!',
            'message': f'Successfully sent {sent_count} interview invitation email(s).'
        }]
    if fail_count:
        existing = session.get('popup_notifications', [])
        existing.append({'type': 'error', 'title': 'Email Failed',
                         'message': f'Failed to send {fail_count} email(s). Check SMTP settings.'})
        session['popup_notifications'] = existing

    return redirect(url_for("results"))

@app.route("/send_email_single/<int:cid>", methods=["POST"])
def send_email_single(cid):
    if "user_id" not in session:
        return redirect(url_for("login"))

    conn     = get_db()
    settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()
    c        = conn.execute("""
        SELECT c.name, c.email, j.title as role_title
        FROM candidates c JOIN job_roles j ON c.job_role_id=j.id WHERE c.id=?
    """, (cid,)).fetchone()

    if not settings or not settings["sender_email"]:
        flash("Email settings not configured.", "danger")
    elif not c or not c["email"]:
        flash("Candidate email not found.", "danger")
    else:
        success, msg = send_interview_email(dict(settings), c["name"], c["email"], c["role_title"])
        if success:
            conn.execute("UPDATE candidates SET email_sent=1, email_sent_at=? WHERE id=?",
                         (datetime.now(), cid))
            conn.commit()
            session['popup_notifications'] = [{
                'type': 'success', 'title': 'Email Sent!',
                'message': f'Interview invitation sent to {c["name"]}.'
            }]
        else:
            session['popup_notifications'] = [{
                'type': 'error', 'title': 'Email Failed',
                'message': f'Could not send email: {msg}'
            }]

    conn.close()
    return redirect(url_for("candidate_detail", cid=cid))

@app.route("/email-settings", methods=["GET", "POST"])
def email_settings_page():
    if "user_id" not in session:
        return redirect(url_for("login"))

    conn     = get_db()
    settings = conn.execute("SELECT * FROM email_settings WHERE id=1").fetchone()

    if request.method == "POST":
        conn.execute("""
            UPDATE email_settings SET smtp_host=?, smtp_port=?, sender_email=?,
                sender_password=?, email_subject=?, email_body=? WHERE id=1
        """, (
            request.form["smtp_host"], int(request.form["smtp_port"]),
            request.form["sender_email"], request.form["sender_password"],
            request.form["email_subject"], request.form["email_body"]
        ))
        conn.commit()
        session['popup_notifications'] = [{
            'type': 'success', 'title': 'Settings Saved!',
            'message': 'Email settings have been saved successfully.'
        }]
        return redirect(url_for("email_settings_page"))

    conn.close()
    popup_notifications = session.pop('popup_notifications', [])
    return render_template("email_settings.html", settings=settings,
                           popup_notifications=popup_notifications)

# ── SIGNUP ──

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if request.method == "POST":
        name     = request.form["name"]
        email    = request.form["email"]
        password = generate_password_hash(request.form["password"])

        conn = get_db()
        try:
            conn.execute(
                "INSERT INTO users (name, email, password) VALUES (?,?,?)",
                (name, email, password)
            )
            conn.commit()
            return redirect(url_for("login"))
        except Exception:
            flash("Email already registered!", "danger")
        finally:
            conn.close()

    return render_template("signup.html")

# ── LOGIN ──

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email    = request.form["email"]
        password = request.form["password"]

        conn = get_db()
        user = conn.execute(
            "SELECT * FROM users WHERE email=?", (email,)
        ).fetchone()

        if user and check_password_hash(user["password"], password):
            login_count = user["login_count"] if user["login_count"] else 0
            is_new_user = (login_count == 0)

            conn.execute(
                "UPDATE users SET login_count = login_count + 1 WHERE id=?",
                (user["id"],)
            )
            conn.commit()
            conn.close()

            session["user_id"]     = user["id"]
            session["user_name"]   = user["name"]
            session["is_new_user"] = is_new_user
            return redirect(url_for("index"))
        else:
            conn.close()
            flash("Invalid email or password", "danger")

    return render_template("login.html")

# ── LOGOUT ──

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ── SEED DEFAULT ROLES ──

@app.route("/seed-roles")
def seed_roles():
    if "user_id" not in session:
        return redirect(url_for("login"))
    conn     = get_db()
    existing = conn.execute("SELECT COUNT(*) as c FROM job_roles").fetchone()["c"]
    if existing == 0:
        for role in DEFAULT_ROLES:
            conn.execute("""
                INSERT INTO job_roles (title, description, core_skills, tools, project_keywords,
                    internship_keywords, experience_keywords)
                VALUES (?,?,?,?,?,?,?)
            """, (
                role["title"], role["description"],
                json.dumps(role["core_skills"]), json.dumps(role["tools"]),
                json.dumps(role["project_keywords"]), json.dumps(role["internship_keywords"]),
                json.dumps(role["experience_keywords"])
            ))
        conn.commit()
        flash(f"Seeded {len(DEFAULT_ROLES)} default IT job roles!", "success")
    else:
        flash("Job roles already exist.", "info")
    conn.close()
    return redirect(url_for("roles"))

# ── API ──

@app.route("/api/dashboard")
def api_dashboard():
    if "user_id" not in session:
        return jsonify([])
    user_id = session["user_id"]
    conn = get_db()
    # ── Only count this user's candidates per role ──
    rows = conn.execute("""
        SELECT j.title,
               COUNT(c.id) as total,
               SUM(CASE WHEN c.status='shortlisted' THEN 1 ELSE 0 END) as shortlisted,
               SUM(CASE WHEN c.status='rejected'    THEN 1 ELSE 0 END) as rejected
        FROM job_roles j
        LEFT JOIN candidates c ON j.id=c.job_role_id AND c.user_id=?
        GROUP BY j.id
    """, (user_id,)).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ─────────────────────────────────────────────
# CUSTOM JINJA2 FILTERS
# ─────────────────────────────────────────────

@app.template_filter("fromjson")
def fromjson_filter(value):
    try:
        return json.loads(value)
    except:
        return []

# if __name__ == "__main__":
#     init_db()
#     app.run(debug=True, port=5000)

if __name__ == "__main__":
    init_db()

    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
